# -*- coding: utf-8 -*-
"""大雪地区広域連合 第10期介護保険事業計画
策定委員会資料　保険者機能強化推進交付金等の評価結果.

令和8年9月3日のご指示
  「交付金の要件は委員会にはかることが要件となるため
    本文記載ではなく策定委員会資料として整理。
    給付費適正化については本文に記載をお願いします。」

交付金の評価指標は、計画の進捗状況のモニタリング結果を
外部の関係者を含む議論の場で検証すること（推進 目標Ⅰ-2 イ）、
評価結果を関係者間で共有して自立支援等に資する取組を検討すること
（同 目標Ⅰ-4）を求めている。
評価結果を策定委員会（運営協議会）に諮ること自体が要件であるため、
計画本文には掲載せず、毎年度の委員会資料として整理する。
計画本文には、委員会へ報告し検証する手順のみを第1章第6節に定める。

分析の方法は、他業務（北塩原村 第10期計画）で用いた方法によっている。
公表資料の明細列ごとに全国該当率を算定し、
「全国の多くの市町村が得点しているのに当区域が0点である項目」を
洗い出す。取組の水準が低いのか、事業に着手していないのかを分けられる。

出力
  output/第10期計画_策定委員会資料_交付金評価結果.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import data_kofukin_detail as D
import data_kofukin_zenkoku as Z

OUT = ("/home/user/repository/output/"
       "第10期計画_策定委員会資料_交付金評価結果.docx")

FONT = "游ゴシック"
NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAYT = RGBColor(0x59, 0x59, 0x59)
TEXTW = 21.0 - 1.9 * 2

TOWNS = ["東川町", "美瑛町", "東神楽町"]
NENDO = ["令和6年度", "令和7年度", "令和8年度"]

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.0)
sec.left_margin = sec.right_margin = Cm(1.9)

st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.15


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cellmargin(t, left=0.12, right=0.12, top=0.06, bottom=0.06):
    mar = OxmlElement("w:tblCellMar")
    for tag, v in (("top", top), ("left", left),
                   ("bottom", bottom), ("right", right)):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), str(int(v * 567)))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    t._tbl.tblPr.append(mar)


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.bold = True
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def P(text, size=10.5, bold=False, gray=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    if gray:
        r.font.color.rgb = GRAYT
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def BUL(text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    r = p.add_run("・" + text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def NOTE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("※ " + text)
    r.font.name = FONT
    r.font.size = Pt(9)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def SRC(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("資料：" + text)
    r.font.name = FONT
    r.font.size = Pt(8.5)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


TBLNO = [0]


def CAP(text):
    TBLNO[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("表%d　%s" % (TBLNO[0], text))
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.bold = True
    r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def TBL(head, rows, widths, size=9.5, center=None, first_bold=False):
    center = center or set()
    t = doc.add_table(rows=1 + len(rows), cols=len(head))
    t.style = "Table Grid"
    t.autofit = False
    tot = sum(widths)
    if tot > TEXTW:
        widths = [w * TEXTW / tot for w in widths]
    _cellmargin(t)
    tr = t.rows[0]
    tr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for j, v in enumerate(head):
        c = tr.cells[j]
        c.width = Cm(widths[j])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade(c, "1F3864")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(v)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for i, row in enumerate(rows):
        tr = t.rows[i + 1]
        tr._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for j, v in enumerate(row):
            c = tr.cells[j]
            c.width = Cm(widths[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2 == 1:
                _shade(c, "F2F5FA")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if j in center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v))
            r.font.name = FONT
            r.font.size = Pt(size)
            r.font.bold = first_bold and j == 0
            r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return t


def FIG(name, src=None, width=15.0):
    import os
    p = "/home/user/repository/output/figures/%s.png" % name
    if not os.path.exists(p):
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(p, width=Cm(width))
    if src:
        SRC(src)


# ================================================================== 表紙
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("大雪地区広域連合介護保険事業計画策定委員会　資料")
r.font.name = FONT
r.font.size = Pt(10.5)
r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
r = p.add_run("保険者機能強化推進交付金及び\n介護保険保険者努力支援交付金の評価結果")
r.font.name = FONT
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = NAVY
r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("令和6年度〜令和8年度")
r.font.name = FONT
r.font.size = Pt(12)
r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# ================================================================== 1
H1("1　この資料の位置づけ")
P("保険者機能強化推進交付金及び介護保険保険者努力支援交付金は、"
  "国が全国のすべての市町村を共通の様式で採点するものです。"
  "当広域連合の取組状況を、全国・北海道と同じ物差しで比較できる資料です。")
P("交付金の評価指標は、次の事項を求めています。"
  "評価結果を策定委員会にお諮りすること自体が要件に当たるため、"
  "計画本文には掲載せず、毎年度の委員会資料として整理します。")
CAP("評価結果を委員会にお諮りする理由（令和8年度評価指標）")
TBL(["交付金・目標", "評価指標", "求められている事項", "配点"],
    [["推進　目標Ⅰ", "2 介護保険事業計画の進捗状況（計画値と実績値の乖離状況）を"
      "分析しているか　イ",
      "モニタリングの結果を外部の関係者と共有し、"
      "乖離の要因やその対応策について、"
      "外部の関係者を含む議論の場で検証を行っている", "4点"],
     ["推進　目標Ⅰ", "同　エ", "モニタリングの結果を公表している", "4点"],
     ["推進　目標Ⅰ", "3・4 施策の実施状況の把握・改善、評価結果の活用",
      "評価結果を関係者間で共有し、"
      "自立支援等に資する取組を検討している", "16点"],
     ["推進　目標Ⅱ", "1 介護給付費の適正化に向けた方策　エ",
      "指標の達成状況を含む取組の成果を公表している。"
      "公表は会議での資料配付又はホームページへの掲載が想定される", "8点"]],
    [2.8, 5.0, 7.0, 1.6])
SRC("厚生労働省「令和8年度保険者機能強化推進交付金及び"
    "介護保険保険者努力支援交付金に係る評価指標（市町村分）」")
NOTE("「議論の場」は、地域ケア会議や計画策定委員会等、"
     "庁外の地域の関係者が参加しているものが想定されています。")

P("")
H2("交付金の仕組み")
BUL("交付金は市町村に交付されます。保険者は大雪地区広域連合ですが、"
    "評価は東川町・美瑛町・東神楽町の3町ごとに行われ、"
    "広域連合としての得点はありません。本資料の「大雪」は3町の単純平均です。")
BUL("推進交付金400点・支援交付金400点の計800点満点です。"
    "目標Ⅰ〜Ⅳがそれぞれ100点で、目標Ⅳ（成果指標群）は"
    "厚生労働省が介護データベースから算定するため、市町村の申告によりません。")
BUL("交付額は得点に応じて配分されます。得点の差はそのまま財源の差になります。")
BUL("令和8年度の交付金は、令和7年度（2025年度）に実施した取組を評価しています。"
    "第9期の最終年度（令和8年度）の取組は令和9年度交付金で評価されるため、"
    "本計画の策定時には結果が出ません。")

doc.add_page_break()

# ================================================================== 2
H1("2　3か年の得点の推移")
_h6 = Z.hikaku("令和6年度", "推進・支援合計")
_h8 = Z.hikaku("令和8年度", "推進・支援合計")
CAP("推進・支援合計の推移（800点満点）")
TBL(["年度", "東川町", "美瑛町", "東神楽町", "大雪（3町平均）", "北海道平均",
     "全国平均", "対全国"],
    [[y] + ["%d" % Z.MACHI[y][t]["推進・支援合計"] for t in TOWNS]
     + ["%.1f" % Z.hikaku(y, "推進・支援合計")["3町平均"],
        "%.1f" % Z.HOK[y]["推進・支援合計"],
        "%.1f" % Z.ZEN[y]["推進・支援合計"],
        "%.1f％" % Z.hikaku(y, "推進・支援合計")["対全国"]]
     for y in NENDO],
    [2.2, 1.8, 1.8, 1.8, 2.4, 2.0, 2.0, 1.8], center={1, 2, 3, 4, 5, 6, 7})
SRC("厚生労働省「保険者機能強化推進交付金及び介護保険保険者努力支援交付金"
    "（市町村分）に係る全国集計結果」令和6〜8年度。"
    "全国は1,745〜1,746市町村、北海道は179市町村の平均")
FIG("fig26_1_kofukin")
P("全国平均は3年で%.1f点（%.1f点→%.1f点）上がりました。"
  "3町平均は%.1f点（%.1f点→%.1f点）の変化にとどまり、"
  "対全国比は%.1f％から%.1f％へ低下しています。"
  % (_h8["全国"] - _h6["全国"], _h6["全国"], _h8["全国"],
     _h8["3町平均"] - _h6["3町平均"], _h6["3町平均"], _h8["3町平均"],
     _h6["対全国"], _h8["対全国"]))
P("取組が後退したのではなく、全国が上がる速さに追いついていない状態です。",
  bold=True)

CAP("全国順位（下位から数えた位置・％）")
TBL(["年度", "東川町", "美瑛町", "東神楽町", "集計対象の市町村数", "読み方"],
    [[y] + ["%.1f％" % Z.PCT[y][t] for t in TOWNS]
     + ["{:,}".format(Z.N[y]),
        "小さいほど全国の下位に位置します"] for y in NENDO],
    [2.2, 2.2, 2.2, 2.2, 3.0, 5.4], center={1, 2, 3, 4})
P("令和8年度は東川町%.1f％・美瑛町%.1f％で、いずれも全国の下位1割以内です。"
  "東神楽町は%.1f％で下位2割の水準です。"
  % (Z.PCT["令和8年度"]["東川町"], Z.PCT["令和8年度"]["美瑛町"],
     Z.PCT["令和8年度"]["東神楽町"]), size=10)

doc.add_page_break()

# ================================================================== 3
H1("3　どこで差がついているか")
H2("（1）指標群別")
CAP("指標群別の得点（令和8年度）")
TBL(["交付金・指標群", "満点", "大雪", "北海道", "全国", "対全国", "達成度"],
    [["推進　体制・取組指標群", "196", "94.7", "134.5", "145.5", "65.1％", "48.3％"],
     ["推進　活動指標群", "104", "23.0", "32.9", "36.9", "62.3％", "22.1％"],
     ["推進　成果指標群", "100", "55.0", "54.4", "47.7", "115.3％", "55.0％"],
     ["支援　体制・取組指標群", "184", "71.7", "116.7", "130.7", "54.8％", "39.0％"],
     ["支援　活動指標群", "116", "34.7", "39.4", "46.5", "74.6％", "29.9％"],
     ["支援　成果指標群", "100", "55.0", "54.4", "47.7", "115.3％", "55.0％"]],
    [4.4, 1.6, 2.0, 2.0, 2.0, 2.0, 2.0], center={1, 2, 3, 4, 5, 6})
FIG("fig26_2_shihyogun")
P("成果（アウトカム）は全国平均の115.3％で、3町とも全国を上回ります。"
  "要介護度の変化と認定率という結果は良好です。"
  "一方、体制・取組と活動は全国の54.8〜74.6％にとどまります。"
  "結果は出ているが、取組の実施と報告の側に課題がある構造です。")

H2("（2）目標別")
CAP("目標別の得点（令和8年度・各100点満点）")
TBL(["目標", "大雪", "全国", "対全国", "差"],
    [[nm, "%.1f" % Z.hikaku("令和8年度", k)["3町平均"],
      "%.1f" % Z.ZEN["令和8年度"][k],
      "%.1f％" % Z.hikaku("令和8年度", k)["対全国"],
      "%+.1f" % (Z.hikaku("令和8年度", k)["3町平均"] - Z.ZEN["令和8年度"][k])]
     for nm, k in [
        ("共通Ⅳ 高齢者が可能な限り自立した日常生活を営む", "推進Ⅳ合計"),
        ("支援Ⅰ 介護予防／日常生活支援を推進する", "支援Ⅰ合計"),
        ("推進Ⅲ 介護人材の確保その他のサービス提供基盤の整備", "推進Ⅲ合計"),
        ("推進Ⅰ 持続可能な地域のあるべき姿をかたちにする", "推進Ⅰ合計"),
        ("支援Ⅱ 認知症総合支援を推進する", "支援Ⅱ合計"),
        ("推進Ⅱ 公正・公平な給付を行う体制を構築する", "推進Ⅱ合計"),
        ("支援Ⅲ 在宅医療・在宅介護連携の体制を構築する", "支援Ⅲ合計")]],
    [7.4, 2.0, 2.0, 2.2, 2.0], center={1, 2, 3, 4})
FIG("fig26_3_mokuhyo")
P("全国を最も下回るのは支援Ⅲ（在宅医療・在宅介護連携）で38.0点、"
  "次いで推進Ⅱ（公正・公平な給付）が31.1点、"
  "支援Ⅱ（認知症総合支援）が21.4点下回ります。"
  "第10期ではこの3目標を重点とします。")

doc.add_page_break()

# ================================================================== 4
H1("4　全国の多くが取得し当区域が取得できていない項目")
P("公表資料は評価指標の明細（ア・イ・ウ・エ及び①〜④）まで得点が入っています。"
  "明細ごとに、全国で得点している市町村の割合（全国該当率）を算定し、"
  "3町とも0点でありながら全国の半数以上が得点している項目を抽出しました。")
P("取組の水準が低いのではなく、事業に着手していないことによる項目を"
  "見分けるためのものです。", size=10)
_tk = D.torikoboshi("令和8年度", 50.0)
P("該当は%d件・配点合計%d点です。"
  "上位20件は次のとおりです。" % (len(_tk), sum(r[5] for r in _tk)))
CAP("3町とも0点かつ全国該当率50％以上の項目（令和8年度・上位20件）")
TBL(["交付金", "目標", "評価指標", "枝番", "配点", "全国該当率"],
    [[r[0], r[1].split(" ")[0], r[3], r[4], "%d点" % r[5], "%.1f％" % r[6]]
     for r in _tk[:20]],
    [1.4, 1.8, 7.6, 1.4, 1.4, 2.2], size=9, center={0, 1, 3, 4, 5})
SRC("同上。全国該当率は当該項目に得点のある市町村の割合"
    "（集計対象1,746市町村・受託者算定）")
NOTE("配点合計%d点は、これらをすべて取得した場合の理論上の増加分です。"
     "成果指標群（厚生労働省が統計から算定するもの）を含むため、"
     "取組により埋められるものはこれより少なくなります。"
     % sum(r[5] for r in _tk))

P("")
P("最も全国該当率が高いのは、推進交付金・目標Ⅱの"
  "「介護給付費適正化事業の取組状況　ア（3事業の全てを実施している）」で、"
  "全国の96.9％が得点しています。"
  "3町とも0点であり、全国で取得できていないのは3.1％です。", bold=True)

doc.add_page_break()

# ================================================================== 5
H1("5　介護給付費の適正化（計画本文 第5章 基本目標5に対応）")
P("給付費の適正化は計画本文に記載します。"
  "本資料では、交付金の評価から見た実施状況を委員会にお示しします。")

H2("（1）主要3事業の実施状況")
CAP("介護給付適正化の主要3事業")
TBL(["主要3事業", "実施状況", "国の評価による裏づけ"],
    [["要介護認定の適正化", "［要確認］",
      "国の評価に単独の項目がないため確認できません"],
     ["ケアプラン等の点検", "3町とも実績なし",
      "「ケアプラン点検の実施状況」（16点）が令和6〜8年度とも3町とも0点"],
     ["縦覧点検・医療情報との突合", "3町とも実施",
      "「医療情報との突合の実施状況」（16点）が令和7年度から3町とも満点。"
      "縦覧点検は効果の高い4帳票すべてを点検"]],
    [3.6, 3.0, 10.6])
P("北海道第9期高齢者保健福祉計画・介護保険事業支援計画の評価指標⑧は"
  "「介護給付適正化の主要3事業の実施率 全市町村100％」です。"
  "縦覧点検・医療情報との突合は令和7年度に着手し満点となりましたが、"
  "ケアプランの点検が3か年とも実績がないため、目標に達していません。")

H2("（2）年度別の推移")
CAP("給付適正化に関する得点の推移（3町計・明細列の合計）")
_rows = []
for nm, man in [("給付費適正化方策の策定状況", 32),
                ("給付費適正化事業の取組状況", 36),
                ("ケアプラン点検の実施状況", 16),
                ("医療情報との突合の実施状況", 16)]:
    row = [nm, "%d点" % man]
    for y in NENDO:
        it = D.item(y, nm)
        row += ["／".join(str(sum(r[7][i] or 0 for r in it))
                          for i in range(3))]
    _rows.append(row)
TBL(["評価指標", "満点", "令和6年度", "令和7年度", "令和8年度"],
    _rows, [5.4, 1.6, 3.2, 3.2, 3.2], center={1, 2, 3, 4})
NOTE("各年度の欄は「東川町／美瑛町／東神楽町」の順です。"
     "満点は1町あたりの配点です。")

H2("（3）取得できていない項目の内訳")
CAP("給付適正化に関する評価項目（令和8年度）")
_gy = []
for nm in ["給付費適正化方策の策定状況", "給付費適正化事業の取組状況",
           "ケアプラン点検の実施状況", "医療情報との突合の実施状況"]:
    for r in D.item("令和8年度", nm):
        _gy.append([nm, r[4], "%d点" % r[5], "%.1f％" % r[6]]
                   + [str(v) for v in r[7]])
TBL(["評価指標", "枝番", "配点", "全国該当率", "東川町", "美瑛町", "東神楽町"],
    _gy, [5.0, 1.4, 1.4, 2.2, 1.6, 1.6, 1.6], size=9,
    center={1, 2, 3, 4, 5, 6})
SRC("同上")
P("3町とも0点で全国の半数以上が得点している項目は46点分です。"
  "内訳は、適正化方策のPDCA（評価指標の設定8点・毎年度の分析改善8点・"
  "成果の公表8点）24点、主要3事業の全てを実施6点、"
  "有料老人ホーム等を含むケアプラン点検8点、"
  "ケアプラン点検の実施割合8点です。", size=10)
P("適正化方策そのものは3町とも策定済みです（8点満点）。"
  "策定はしているが、指標を定めて検証し公表する段階に至っていない状態です。",
  size=10)

doc.add_page_break()

# ================================================================== 6
H1("6　成果指向型配分枠（令和8年度に新設）")
P("令和8年度から、推進交付金の400点とは別に、"
  "成果指向型の介護予防・健康づくりの取組を100点満点で評価する枠が"
  "新設されました。")
CAP("成果指向型配分枠の要件")
TBL(["要件", "内容"],
    [["評価の視点",
      "地域のデータ分析に基づき、自らターゲットとなる対象者、"
      "成果目標及び評価指標を設定した上で、"
      "当該成果を達成するために成果指向型の"
      "介護予防・健康づくりの取組を行っているか"],
     ["データ分析",
      "地域の介護給付費の動向や地域資源、"
      "医療・介護の健康づくりに関するデータ等、"
      "十分なデータに基づき地域分析を行うこと"],
     ["対象層", "年齢・状態・性別等の具体的な住民層を特定して設定すること"],
     ["アウトプット指標",
      "データ等に基づく客観的な評価が行えるものになっていること"],
     ["アウトカム指標",
      "取組の実施により、期待されていた対象層や社会に現れた変化を"
      "客観的に効果検証できるものとなっていること"],
     ["目標値", "アウトプット指標及びアウトカム指標に具体的な目標値を記載すること"]],
    [3.4, 13.8], first_bold=True)
SRC("厚生労働省「令和8年度保険者機能強化推進交付金及び"
    "介護保険保険者努力支援交付金に係る評価指標（市町村分）」")
P("「事業を実施したか」ではなく「対象を絞り、目標を立て、効果を検証したか」を"
  "問うものです。第10期の施策設計に直接関わります。")
P("本業務で作成した地域差の分析、認定率の年齢調整分析、"
  "健康とくらしの調査の集計、在宅生活改善調査は、"
  "要件の「十分なデータに基づく地域分析」の材料になります。", size=10)
CAP("対象層の候補（受託者の案）")
TBL(["候補", "現状", "アウトカム指標の例"],
    [["通いの場に参加していない層", "通いの場の参加率8.8％（3町計25箇所・330人）",
      "参加率、参加者の心身機能の維持改善割合"],
     ["認定を受けながらサービスを利用していない層",
      "未利用の認定者504人（認定者の25.7％）",
      "サービス利用率、要介護度の変化"],
     ["85歳以上", "第10期の3年間で162人増（1,990人→2,152人）",
      "85歳以上の認定率、在宅生活の継続期間"],
     ["フレイル該当者", "フレイル該当割合19.1％（令和7年度）",
      "フレイル該当割合、要支援認定の新規発生率"]],
    [4.0, 6.4, 6.8])
NOTE("令和8年度の全国集計には成果指向型配分枠の得点欄がありません。"
     "3町が該当状況調査でどのように回答したかは、3町へのご照会を要します。")

doc.add_page_break()

# ================================================================== 7
H1("7　委員会にお諮りする事項")
CAP("ご意見をいただきたい事項")
TBL(["No.", "事項", "内容", "受託者の案"],
    [["1", "得点の低い項目の要因",
      "3町とも0点で全国の半数以上が得点している項目が%d件・%d点あります。"
      "未実施によるものか、要件を満たしていないのか、"
      "記録・報告の問題かによって対応が変わります。"
      % (len(_tk), sum(r[5] for r in _tk)),
      "3町の事務担当と項目ごとに突き合わせ、要因を分けたうえで"
      "第10期のアウトプット指標に位置づけます。"],
     ["2", "ケアプランの点検の着手",
      "主要3事業のうちケアプランの点検が3か年とも実績がありません。"
      "北海道の評価指標⑧の目標（実施率100％）に達していません。",
      "令和9年度当初に実施要領を定め、同年度から実施します。"
      "初年度は件数を絞り、居宅介護支援事業所と手順を共有したうえで"
      "段階的に広げます（計画本文 第5章 基本目標5）。"],
     ["3", "適正化方策のPDCA",
      "方策は策定済みですが、効果を検証する評価指標の設定、"
      "毎年度の分析・改善、成果の公表の3つが揃っていません（計24点）。",
      "評価指標を令和9年度当初に定め、"
      "毎年度の評価結果を本委員会に報告し、"
      "公表することで一連の手順とします。"],
     ["4", "評価結果の毎年度の報告",
      "評価結果を外部の関係者を含む議論の場で検証することが"
      "交付金の要件です。",
      "毎年度7〜9月に本委員会へ報告し、ご意見をいただく手順を"
      "計画本文（第1章第6節）に定めます。"],
     ["5", "成果指向型配分枠への対応",
      "令和8年度に新設された100点の枠です。"
      "対象層・アウトプット指標・アウトカム指標と目標値の設定が要件です。",
      "第10期の介護予防・健康づくりの施策に織り込みます。"
      "対象層の候補は6の表のとおりです。"],
     ["6", "全国順位の扱い",
      "令和8年度は東川町が全国の下位6.7％、美瑛町が6.0％、"
      "東神楽町が20.5％です。",
      "計画本文には掲載せず、本委員会資料にとどめることを案とします。"
      "交付額は得点に応じて配分されるため、"
      "順位は財源の差として実質的な意味を持ちます。"]],
    [1.0, 3.6, 6.6, 6.0], size=9.5, center={0})

P("")
NOTE("本資料の得点はすべて厚生労働省の公表資料によるものです。"
     "3町×3か年の198値を原本と突合しており、収録値との不一致はありません。"
     "分析の方法は他業務（北塩原村 第10期計画）で用いた方法によっています。")
NOTE("交付金の評価結果は計画本文には掲載していません。"
     "評価結果を本委員会にお諮りし、ご意見を踏まえて対応することが"
     "交付金の評価指標の求める事項であるためです。"
     "計画本文には、毎年度本委員会へ報告し検証する手順のみを"
     "第1章第6節に定めています。")

doc.save(OUT)
print("saved:", OUT)
print("段落 %d / 表 %d" % (len(doc.paragraphs), len(doc.tables)))
print("3町とも0点かつ全国該当率50％以上　%d件・配点計%d点"
      % (len(_tk), sum(r[5] for r in _tk)))
for y in NENDO:
    h = Z.hikaku(y, "推進・支援合計")
    print("%s 3町平均%.1f 全国%.1f 対全国%.1f％"
          % (y, h["3町平均"], h["全国"], h["対全国"]))
