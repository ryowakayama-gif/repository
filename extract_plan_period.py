# -*- coding: utf-8 -*-
"""
環境基本計画PDFから計画期間（開始年度・終期年度）を抽出する

背景:
  営業先リストの未確認先は、計画期間がHTMLに出ておらずPDF本文にしか
  記載がない団体が多い。1団体あたり「PDFを開いて表紙・第1章の1行を読む」
  だけの作業なので、そこを機械化する。

処理:
  1. PDFを取得（ローカルパス、またはURL）
  2. テキストレイヤの有無を判定（画像のみのPDFはOCRが必要と報告して中断）
  3. PDF → Word(.docx) へ変換（pdf2docx）。表の中に計画期間が書かれている
     ケースを拾うため、テキスト抽出だけでなくWord変換も行う
  4. テキスト源を2系統（pdftotext -layout / docxの段落＋表）用意して突合
  5. 和暦・西暦の期間表現を正規表現で抽出し、西暦へ正規化して候補を提示

使い方:
  python3 extract_plan_period.py <PDFパス または URL> [--docx 出力.docx] [--pages 8]

注意:
  自治体サイトへの直接アクセスが遮断されている環境では、URL指定は失敗する。
  その場合はPDFを手元に置いてローカルパスで渡すこと。
"""

import argparse
import os
import re
import subprocess
import sys
import unicodedata

ERA_BASE = {"明治": 1867, "大正": 1911, "昭和": 1925, "平成": 1988, "令和": 2018}
DASH = "からまで〜～~－ー–—-至"
# 「|」は PDF→Word 変換で表のセル区切りとして現れるため含める。
# ノイズが増えるが、スコア（計画期間などのキーワード近傍）で順位付けする。
DASH_RE = r"(?:から|〜|～|~|－|ー|–|—|-|‐|to|\||｜)"


def to_seireki(era, num):
    """和暦→西暦。num は int または '元'。"""
    if num == "元":
        num = 1
    return ERA_BASE[era] + int(num)


def normalize(text):
    """全角英数・記号のゆれを吸収する。"""
    text = unicodedata.normalize("NFKC", text)
    return text.replace("　", " ")


# ------------------------------------------------------------
# テキスト抽出
# ------------------------------------------------------------
def has_text_layer(pdf_path):
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf
    doc = pymupdf.open(pdf_path)
    chars = 0
    for page in doc[: min(10, doc.page_count)]:
        chars += len(page.get_text().strip())
    doc.close()
    return chars, chars >= 200


def text_from_pdftotext(pdf_path, pages):
    if not _which("pdftotext"):
        return ""
    out = subprocess.run(
        ["pdftotext", "-layout", "-f", "1", "-l", str(pages), pdf_path, "-"],
        capture_output=True, text=True)
    return out.stdout


def pdf_to_docx(pdf_path, docx_path, pages):
    from pdf2docx import Converter
    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path, start=0, end=pages)
    finally:
        cv.close()
    return docx_path


def text_from_docx(docx_path):
    from docx import Document
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.replace("\n", " ") for c in row.cells))
    return "\n".join(parts)


def _which(cmd):
    return subprocess.run(["which", cmd], capture_output=True).returncode == 0


# ------------------------------------------------------------
# 期間の抽出
# ------------------------------------------------------------
WAREKI_RANGE = re.compile(
    r"(明治|大正|昭和|平成|令和)\s*(\d{1,2}|元)\s*年度?"
    r"[^\n]{0,24}?" + DASH_RE +
    r"[^\n]{0,24}?"
    r"(?:(明治|大正|昭和|平成|令和)\s*)?(\d{1,2}|元)\s*年度?"
)
SEIREKI_RANGE = re.compile(
    r"(\d{4})\s*年度?[^\n]{0,20}?" + DASH_RE + r"[^\n]{0,20}?(\d{4})\s*年度?"
)
SEIREKI_PAREN = re.compile(r"(\d{4})\s*(?:年度)?\s*[)）]")
KEYWORDS = ("計画期間", "計画の期間", "期間", "目標年度", "計画年次", "対象期間")


def find_periods(text):
    """(開始西暦, 終期西暦, 元表記, 文脈) の候補を返す。"""
    text = normalize(text)
    results = []
    for line_no, raw in enumerate(text.split("\n")):
        line = raw.strip()
        if not line:
            continue
        for m in WAREKI_RANGE.finditer(line):
            e1, n1, e2, n2 = m.group(1), m.group(2), m.group(3), m.group(4)
            e2 = e2 or e1  # 「平成30年度から9年度」のように後半の元号が略される場合
            try:
                y1, y2 = to_seireki(e1, n1), to_seireki(e2, n2)
            except (KeyError, ValueError):
                continue
            if not (1990 <= y1 <= 2060 and 1990 <= y2 <= 2060 and y2 > y1):
                continue
            results.append((y1, y2, m.group(0), line, line_no))
        for m in SEIREKI_RANGE.finditer(line):
            y1, y2 = int(m.group(1)), int(m.group(2))
            if not (1990 <= y1 <= 2060 and 1990 <= y2 <= 2060 and y2 > y1):
                continue
            results.append((y1, y2, m.group(0), line, line_no))
    return results


def score(cand):
    """計画期間らしさ。キーワード近傍・年数の妥当性で重み付けする。"""
    y1, y2, raw, line, _ = cand
    s = 0
    if any(k in line for k in KEYWORDS):
        s += 10
    if "計画期間" in line:
        s += 10
    span = y2 - y1 + 1
    if 3 <= span <= 20:
        s += 5
    if span in (5, 6, 8, 10, 12):
        s += 3
    if "年度" in raw:
        s += 3
    return s


def seireki_to_wareki(y):
    """年度の西暦→和暦表記（令和優先）。"""
    if y >= 2019:
        return f"令和{y - 2018}年度({y})"
    if y >= 1989:
        return f"平成{y - 1988}年度({y})"
    return f"{y}年度"


# ------------------------------------------------------------
def acquire(src, workdir):
    if not src.lower().startswith(("http://", "https://")):
        if not os.path.exists(src):
            print(f"エラー: ファイルがありません -> {src}")
            sys.exit(1)
        return src
    dst = os.path.join(workdir, "downloaded.pdf")
    print(f"取得中: {src}")
    r = subprocess.run(["curl", "-sSL", "--max-time", "60", "-o", dst, src], capture_output=True, text=True)
    if r.returncode != 0 or not os.path.exists(dst) or os.path.getsize(dst) < 1000:
        print("取得に失敗しました。この環境では自治体サイトへの直接アクセスが")
        print("組織のegressポリシーで遮断されています。PDFを手元に保存し、")
        print("ローカルパスを指定して再実行してください。")
        if r.stderr:
            print(f"  curl: {r.stderr.strip()[:200]}")
        sys.exit(2)
    return dst


def main():
    ap = argparse.ArgumentParser(description="環境基本計画PDFから計画期間を抽出する")
    ap.add_argument("src", help="PDFのパスまたはURL")
    ap.add_argument("--docx", help="Word変換の出力先（既定: 入力と同名の .docx）")
    ap.add_argument("--pages", type=int, default=8, help="先頭から何ページを対象にするか（既定8）")
    ap.add_argument("--keep-docx", action="store_true", help="Wordファイルを残す")
    args = ap.parse_args()

    workdir = os.path.dirname(os.path.abspath(args.docx or args.src)) or "."
    pdf = acquire(args.src, workdir)
    print(f"対象: {pdf}")

    chars, ok = has_text_layer(pdf)
    print(f"テキストレイヤ: 先頭10ページで {chars:,} 文字")
    if not ok:
        print("→ 画像のみのPDFと判断しました。この環境にはOCR(tesseract)が無いため、")
        print("  文字を起こせません。OCR済みPDFを用意するか、手元のWordで開いてください。")
        sys.exit(3)

    docx_path = args.docx or os.path.splitext(pdf)[0] + ".docx"
    print(f"Word変換: {docx_path}")
    try:
        pdf_to_docx(pdf, docx_path, args.pages)
        docx_text = text_from_docx(docx_path)
    except Exception as exc:  # 変換失敗でもテキスト抽出は続ける
        print(f"  ! Word変換に失敗しました: {exc}")
        docx_text = ""

    plain = text_from_pdftotext(pdf, args.pages)
    print(f"テキスト源: pdftotext {len(plain):,}文字 / docx {len(docx_text):,}文字")

    cands = {}
    for label, text in (("docx", docx_text), ("pdftotext", plain)):
        for c in find_periods(text):
            key = (c[0], c[1])
            entry = cands.setdefault(key, {"score": 0, "raw": c[2], "lines": set(), "src": set()})
            entry["score"] = max(entry["score"], score(c))
            entry["lines"].add(c[3][:110])
            entry["src"].add(label)

    if not cands:
        print("\n計画期間らしい記述を抽出できませんでした。対象ページ数を増やして再実行してください。")
        sys.exit(4)

    print("\n=== 計画期間の候補（スコア順） ===")
    ranked = sorted(cands.items(), key=lambda kv: (-kv[1]["score"], kv[0]))
    for (y1, y2), info in ranked[:8]:
        span = y2 - y1 + 1
        mark = "★" if info["score"] >= 20 else "  "
        print(f"{mark} {seireki_to_wareki(y1)} 〜 {seireki_to_wareki(y2)}  （{span}年間） "
              f"score={info['score']} 出典={'+'.join(sorted(info['src']))}")
        for line in sorted(info["lines"])[:2]:
            print(f"     {line}")

    best = ranked[0]
    (y1, y2), info = best
    print("\n=== 判定 ===")
    print(f"  計画期間: {seireki_to_wareki(y1)} 〜 {seireki_to_wareki(y2)}（{y2 - y1 + 1}年間）")
    print(f"  終期年度: {seireki_to_wareki(y2)}")
    print(f"  R9(2027)満了判定: {'★該当' if y2 == 2027 else '非該当'}")
    if info["score"] < 20:
        print("  ! スコアが低いため、原文の該当行を必ず目視で確認してください。")

    if not args.keep_docx and not args.docx and os.path.exists(docx_path):
        print(f"\n（Word: {docx_path} を残しました。原文確認に使えます）")


if __name__ == "__main__":
    main()
