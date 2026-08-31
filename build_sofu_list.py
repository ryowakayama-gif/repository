# -*- coding: utf-8 -*-
"""大雪地区広域連合 第10期介護保険事業計画　送付資料一覧（令和8年8月）.

発注者へ送付する資料だけを掲げる。
受託者内部の作業記録は掲げない（成果品一覧は内部保管とする）。

シート構成
  00_送付資料一覧
  01_ご確認・ご決定をお願いする事項
  02_今後更新する箇所
"""

import os
import sys

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from data_dispatch import DISPATCH, SOFU, JOKEN     # noqa: E402

ODIR = "/home/user/repository/output"
OUT = os.path.join(ODIR, "第10期計画_送付資料一覧.xlsx")

FONT = "游ゴシック"
NAVY, HEAD = "1F3864", "4472C4"
IN_Y, OK_G, NG_O, MID_B, GRAY = "FFF2CC", "E2EFDA", "FCE4D6", "DEEBF7", "F2F2F2"
thin = Side(style="thin", color="BFBFBF")
BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)

wb = Workbook()
wb.remove(wb.active)


def sheet(name, title, subtitle, widths, freeze="A5"):
    ws = wb.create_sheet(name)
    ws.sheet_view.showGridLines = False
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    c = ws.cell(row=1, column=1, value=title)
    c.font = Font(name=FONT, size=14, bold=True, color=NAVY)
    ws.row_dimensions[1].height = 22
    c = ws.cell(row=2, column=1, value=subtitle)
    c.font = Font(name=FONT, size=9)
    c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(widths))
    ws.row_dimensions[2].height = 46
    ws.freeze_panes = freeze
    return ws


def header(ws, row, cols, height=28):
    for i, v in enumerate(cols, start=1):
        c = ws.cell(row=row, column=i, value=v)
        c.font = Font(name=FONT, size=9, bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=HEAD)
        c.alignment = Alignment(horizontal="center", vertical="center",
                                wrap_text=True)
        c.border = BORDER
    ws.row_dimensions[row].height = height
    return row + 1


def body(ws, row, vals, fills=None, height=30, align=None, bold=False):
    for i, v in enumerate(vals, start=1):
        c = ws.cell(row=row, column=i, value=v)
        c.font = Font(name=FONT, size=9, bold=bold)
        c.alignment = Alignment(wrap_text=True, vertical="top",
                                horizontal=(align or {}).get(i, "left"))
        c.border = BORDER
        if fills and i in fills:
            c.fill = PatternFill("solid", fgColor=fills[i])
    ws.row_dimensions[row].height = height
    return row + 1


def lead(ws, row, text, span):
    c = ws.cell(row=row, column=1, value=text)
    c.font = Font(name=FONT, size=10, bold=True, color=NAVY)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)
    ws.row_dimensions[row].height = 18
    return row + 1


def note(ws, row, text, span, height=80):
    c = ws.cell(row=row, column=1, value=text)
    c.font = Font(name=FONT, size=8.5)
    c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)
    ws.row_dimensions[row].height = height
    return row + 1


def scale_of(path):
    if path.endswith(".xlsx"):
        w = load_workbook(path, read_only=True)
        n = len(w.sheetnames)
        w.close()
        return "%dシート" % n
    if path.endswith(".odt"):
        return "ODF形式"
    d = Document(path)
    return "%d段落 %d表" % (len(d.paragraphs), len(d.tables))


# ------------------------------------------------ 分類（送付資料の並び順）
GROUP = [
    ("1　計画本体", [
        "第10期介護保険事業計画_協議用素案_令和8年8月.docx",
        "第10期介護保険事業計画_骨子案.docx",
        "第10期計画_図表集_白黒.xlsx",
        "第10期計画_計画素案の別管理表.xlsx",
    ]),
    ("2　第9期計画の評価・検証（令和8年8月の提出物）", [
        "第10期計画_第9期計画の評価・検証_中間報告.docx",
        "第10期計画_中間報告の根拠対照表.xlsx",
        "第10期計画_妥当性検証報告書.xlsx",
        "第10期計画_第9期施策と調査・KPIの紐付けレビュー.xlsx",
        "第10期計画_第9期施策別評価表と暫定評価ルール.xlsx",
        "第10期計画_施策体系新旧対照表.xlsx",
    ]),
    ("3　調査の集計と分析", [
        "第10期計画_アンケート調査の集計分析報告書.xlsx",
        "第10期計画_調査クロス集計・分析.xlsx",
        "第10期計画_実施済み調査_結果報告書.docx",
        "第10期計画_実施済み3調査の受領点検と集計.xlsx",
        "第10期計画_追加調査報告書.xlsx",
        "第10期計画_事業所調査の照会票と確定値管理表.xlsx",
    ]),
    ("4　推計と分析", [
        "第10期計画_将来推計_人口と認定者数.xlsx",
        "第10期計画_将来推計_第2段階_サービス見込量.xlsx",
        "第10期計画_将来推計_需要3シナリオの感度表.xlsx",
        "第10期計画_将来推計_第3段階_給付費と保険料.xlsx",
        "第10期計画_人口推計の基礎の検証.xlsx",
        "第10期計画_社人研推計の町別データの受領点検.xlsx",
        "第10期計画_人口推計の補正_65歳以上75歳以上の突合.xlsx",
        "第10期計画_認定率の年齢調整分析.xlsx",
        "第10期計画_地域差の分析.xlsx",
        "第10期計画_保険料と施策評価の他団体比較.xlsx",
        "第10期計画_保険料の所得段階と低所得者軽減の検証.xlsx",
    ]),
    ("5　基礎データの点検と突合", [
        "第10期計画_要介護認定データの確認.xlsx",
        "第10期計画_サービス受給者数データの確認.xlsx",
        "第10期計画_住まいと施設の公表名簿との突合.xlsx",
        "第10期計画_3町の社会資源一覧との突合.xlsx",
        "第10期計画_従業員数の重複計上の整理.xlsx",
        "第10期計画_給付実績データの受領点検.xlsx",
        "第10期計画_令和8年8月28日受領資料の点検結果.xlsx",
        "第10期計画_見える化総括表の受領点検.xlsx",
        "第10期計画_総合事業実施状況調査の受領点検.xlsx",
        "第10期計画_年報月報の受領点検.xlsx",
        "第10期計画_令和6年度決算書の受領点検.xlsx",
        "第10期計画_基金条例の確認.xlsx",
        "第10期計画_世帯構成の突合.xlsx",
    ]),
    ("6　ご確認・ご決定をお願いする資料", [
        "第10期計画_発注者確認事項一覧.xlsx",
        "第10期計画_必要事項の一覧.xlsx",
        "第10期計画_確認依頼書.docx",
        "第10期計画_資料提供依頼_第9期の施策事業実績.xlsx",
        "第10期計画_代表KPIの振替案と確認事項の精査.xlsx",
        "第10期計画_概要版の構成案.xlsx",
    ]),
    ("7　工程", [
        "第10期計画_業務進捗報告書_令和8年8月分.docx",
        "第10期計画_業務工程管理表.xlsx",
        "第10期計画_送付資料一覧.xlsx",
    ]),
    ("8　会議", [
        "第10期計画_キックオフ会議資料_令和8年8月.docx",
        "第10期計画_キックオフ会議議事録_令和8年8月6日.odt",
        "第10期計画_中間報告会議議事録_令和8年8月26日.odt",
        "第10期計画_キックオフ会議ヒアリングシート.docx",
        "第10期計画_3町ヒアリング資料.docx",
        "第10期計画_3町別の論点整理.xlsx",
    ]),
]

# ============================================================ 00
ws = sheet("00_送付資料一覧", "送付資料一覧（令和8年8月）",
           "今回送付する資料の一覧です。"
           "「取扱い」が「お諮りする内容を含む」の資料には、"
           "受託者の判断・推奨・試算が含まれます。確定値ではありません。"
           "「送付可能（留保明記済み）」の資料にも、"
           "重複の取扱いが未確定の数値等が含まれる場合があり、"
           "その旨は各資料の本文に明記しています。",
           [5, 42, 52, 10, 20])

r = header(ws, 4, ["No.", "資料名", "内容", "規模", "取扱い"])
DESC = {k: v[1] for k, v in DISPATCH.items()}
i = 0
miss, cnt = [], {SOFU: 0, JOKEN: 0}
for gname, files in GROUP:
    r = lead(ws, r, "【%s】" % gname, 5)
    for fn in files:
        kb, why = DISPATCH.get(fn, ("未分類", ""))
        if kb not in (SOFU, JOKEN):
            miss.append("%s（区分 %s）" % (fn, kb))
            continue
        p = os.path.join(ODIR, fn)
        if not os.path.exists(p):
            miss.append(fn + "（実体なし）")
            continue
        i += 1
        cnt[kb] += 1
        r = body(ws, r, [i, fn, why, scale_of(p),
                         "送付可能\n（留保明記済み）" if kb == SOFU else "お諮りする内容を含む"],
                 {5: OK_G if kb == SOFU else IN_Y}, height=34,
                 align={1: "center", 4: "center", 5: "center"})

r += 1
r = body(ws, r, ["計", i, "うち 送付可能（留保明記済み）%d件／お諮りする内容を含む %d件"
                 % (cnt[SOFU], cnt[JOKEN]), "", ""],
         {j: GRAY for j in range(1, 6)}, height=20, bold=True,
         align={1: "center", 2: "center"})
ws.merge_cells(start_row=r - 1, start_column=3, end_row=r - 1, end_column=5)

r += 1
r = note(ws, r,
         "注1）「お諮りする内容を含む」の資料は、"
         "受託者が現時点で受領している資料の範囲で作成した案・試算・判定を含みます。"
         "ご決定は発注者及び策定委員会において行っていただくものです。\n"
         "注2）計画素案は「協議用素案（令和8年8月時点）」です。"
         "本文には［要協議］［要確認］［要内訳］を残しています。"
         "所在と決定の単位は「発注者確認事項一覧」をご覧ください。\n"
         "注3）本一覧のほか、計画本文に掲載する図表の画像（PNG）を"
         "送付用ZIPの「03_図表の画像」に収めています。\n"
         "注4）本業務で作成した資料のうち、受託者の作業過程の記録"
         "（校正の記録、自己点検の記録、修正指示書、進行者の手元資料）は"
         "本一覧に含めていません。"
         "業務仕様書４（10）「その他業務で作成した資料　一式」として保管しており、"
         "納品時又はお求めに応じて提出します。", 5, height=140)

# ============================================================ 01
ws = sheet("01_ご確認・ご決定をお願いする事項",
           "ご確認・ご決定をお願いする事項",
           "詳細は「発注者確認事項一覧」をご覧ください。"
           "本シートは、そのうち令和8年9月から10月までにご決定を要するものの要約です。",
           [5, 34, 46, 16, 14])

r = header(ws, 4, ["No.", "ご決定・ご確認をお願いする事項", "決定しない場合の影響",
                   "関係資料", "期限（案）"])
for a in [
    (1, "人口推計の基礎に何を用いるか"
        "（社人研推計／3町の地方創生総合戦略／補正後の推計）",
     "第6章の採用値、及び第3段階（給付費・保険料）を確定できません。"
     "第1段階（人口・認定者数）と第2段階（サービス見込量）の試算は"
     "本送付に含めています。",
     "人口推計の基礎の検証／人口推計の補正", "令和8年9月"),
    (2, "認定率の将来シナリオをどれとするか（3シナリオ）",
     "サービス見込量の幅が決まりません。", "将来推計_需要3シナリオの感度表",
     "令和8年9月"),
    (3, "第9期代表KPIの中間評価の判定を了とするか",
     "第3章の記述が確定しません。"
     "第9期最終年度の実績確定後に最終評価を行う前提です。",
     "第9期計画の評価・検証 中間報告", "令和8年9月"),
    (4, "代表KPIのうち算定できない4項目の振替案を了とするか",
     "第4章の代表KPIが確定しません。", "代表KPIの振替案と確認事項の精査",
     "令和8年9月"),
    (5, "概要版の構成を3案のいずれとするか",
     "概要版の作成に着手できません。", "概要版の構成案", "令和8年10月"),
    (6, "3町の高齢者福祉計画に共通指標を掲載していただけるか",
     "資料編 資料2（3町との役割分担・共通指標）が機能しません。",
     "施策体系新旧対照表／3町ヒアリング資料", "令和8年10月"),
    (7, "介護サービス自給率を計画本文に掲載するか",
     "町により算定できる区分数が異なるため"
     "（美瑛6区分、東川・東神楽5区分）、"
     "平均値の掲載の可否についてご判断を要します。",
     "地域差の分析", "令和8年10月"),
]:
    r = body(ws, r, list(a), {}, height=54, align={1: "center", 5: "center"})

r += 1
r = lead(ws, r, "【ご提供をお願いする資料】", 5)
r = header(ws, r, ["No.", "資料", "必要な理由", "", "優先"])
for a in [
    (13, "令和6〜8年度の施策・事業実績"
         "（第9期19施策・地域支援事業15区分・基盤整備23サービス・決算）",
     "第9期の評価のうちプロセス評価（施策が計画どおり実施されたか）に"
     "必要です。他の統計では代替できません。記入様式は作成済みです"
     "（資料提供依頼_第9期の施策事業実績）。", "", "最高"),
    (5, "保険者機能強化推進交付金の評価調書",
     "活動指標群の得点が低い理由が、未実施・要件未充足・"
     "記録又は報告の不備のいずれによるかを確認するために必要です。", "", "高"),
    (15, "第9期の代表KPIに関する年度ごとの評価・共有の記録",
     "年度ごとの進捗確認の実施の有無を確定するために必要です。", "", "高"),
]:
    r = body(ws, r, list(a), {5: NG_O if a[4] == "最高" else IN_Y},
             height=54, align={1: "center", 5: "center"})
    ws.merge_cells(start_row=r - 1, start_column=3, end_row=r - 1, end_column=4)

note(ws, r + 1,
     "注）資料の全体は「必要事項の一覧」（不足資料18件・確認事項56件）"
     "及び「確認依頼書」に掲げています。", 5, height=32)

# ============================================================ 02
ws = sheet("02_今後更新する箇所", "今後更新する箇所",
           "今回送付する資料のうち、今後の資料の受領又はご決定により"
           "更新する箇所を示します。",
           [5, 30, 46, 26, 14])

r = header(ws, 4, ["No.", "資料", "更新する箇所", "更新の条件", "時期（見込み）"])
for a in [
    (1, "第9期計画の評価・検証 中間報告",
     "プロセス評価（施策が計画どおり実施されたか）。"
     "現在は成果評価と構造評価のみです。",
     "資料No.13の受領", "受領後1か月"),
    (2, "同上",
     "第9期の最終評価。現在の判定はいずれも中間判定です。",
     "令和8年度の実績の確定（令和9年4月以降）",
     "本業務の期間外"),
    (3, "同上",
     "交付金の活動指標群の得点が低い理由。現在は未確認です。",
     "資料No.5の受領", "受領後2週間"),
    (4, "計画素案（協議用素案）",
     "第6章の見込量・給付費・保険料。［要協議］［要内訳］としています。",
     "人口推計の基礎のご決定、報酬改定率等の確定",
     "令和8年12月"),
    (5, "同上",
     "第2章第3節の供給量。3町の社会資源一覧を反映します。",
     "受託者の作業", "令和8年9月上旬"),
    (6, "将来推計（第1段階・第2段階）",
     "推計値。補正の採否により変わります。",
     "人口推計の基礎のご決定", "令和8年9月"),
    (7, "図表集",
     "令和6・7年度の実績を反映した図表。", "資料No.13の受領",
     "受領後1か月"),
    (8, "業務工程管理表・業務進捗報告書",
     "各業務の進捗率。", "毎月の更新", "毎月"),
]:
    r = body(ws, r, list(a), {}, height=44, align={1: "center", 5: "center"})

note(ws, r + 1,
     "注）本業務の契約期間は令和9年3月26日まで、"
     "第9期計画の最終年度末は令和9年3月31日です。"
     "第9期の最終評価は本業務の期間中には完結しないため、"
     "本業務では取得可能な最新時点までの実績により評価しています。", 5, height=48)

wb.save(OUT)
print("saved:", os.path.basename(OUT), "sheets=%d" % len(wb.sheetnames))
print("送付資料 %d件（そのまま %d／お諮りする内容を含む %d）"
      % (i, cnt[SOFU], cnt[JOKEN]))
print("一覧に載せなかったもの:", "、".join(miss) if miss else "なし")

# 送付区分が「送付」「条件付き」なのに一覧に載っていないものを検出する
listed = {f for _g, fs in GROUP for f in fs}
gap = sorted(f for f, (k, _w) in DISPATCH.items()
             if k in (SOFU, JOKEN) and f not in listed
             and os.path.exists(os.path.join(ODIR, f)))
print("送付区分との突合:",
      "全件が一覧にある" if not gap else "一覧に未掲載 " + "、".join(gap))
