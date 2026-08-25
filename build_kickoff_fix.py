# -*- coding: utf-8 -*-
"""発注者ご作成のキックオフ会議資料（更新版）に校正結果を反映する.

発注者からご提示のあった原本
  40c7a514-_2026.8.5__________________1.docx（令和8年8月6日現在）
に対し、校正結果30件のうち、令和8年8月6日にご指示のあった方針で
修正を反映した版を作成する。

ご指示
  ・3町の担当範囲の「同上」はそのまま
  ・担当窓口からの「受託者」の削除はそのまま
  ・表紙の「会議日時・出席者」の行の削除はそのまま
  ・5. の表及び11. の表の「差額0円」の記載も削除する
  ・事業所実態調査は実施済みとして記載を統一する

原本の書式（フォント・段落・表の体裁）を保つため、
python-docx により実行単位（run）で文字列を置換する。
段落・表の追加は行わない。

出力
  output/第10期計画_キックオフ会議資料（更新版）_校正反映.docx
"""

import os

from docx import Document

SRC = ("/root/.claude/uploads/54f527c9-842b-534d-b822-e5a6c91f837c/"
       "40c7a514-_2026.8.5__________________1.docx")
OUT = ("/home/user/repository/output/"
       "第10期計画_キックオフ会議資料（更新版）_校正反映.docx")


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    yield p


def replace_in_paragraph(p, old, new):
    """段落内の文字列を置換する。実行（run）をまたぐ場合にも対応する。

    最初に該当した実行に置換後の文字列を入れ、
    以降の該当部分の実行は空にする。書式は最初の実行のものが残る。
    """
    runs = p.runs
    if not runs:
        return 0
    full = "".join(r.text for r in runs)
    if old not in full:
        return 0
    n = 0
    guard = 0
    while old in full:
        guard += 1
        if guard > 50:
            raise RuntimeError("置換が収束しません: %r → %r" % (old, new))
        start = full.index(old)
        end = start + len(old)
        pos, first = 0, None
        for r in runs:
            rs, re_ = pos, pos + len(r.text)
            if re_ > start and rs < end:
                head = r.text[:max(0, start - rs)]
                tail = r.text[max(0, end - rs):] if re_ > end else ""
                if first is None:
                    r.text = head + new + tail
                    first = r
                else:
                    r.text = head + tail
            pos = re_
        n += 1
        full = "".join(r.text for r in runs)
    return n


REPS = [
    # ---------------------------------------------------------- A 誤字
    ("大切地区広域連合", "大雪地区広域連合", 1),
    ("大雪広域連合様と弊社において", "大雪地区広域連合様と弊社において", 1),
    ("居宅変更実態調査", "居所変更実態調査", 1),
    ("推計（案、骨子・KPI目標", "推計（案）、骨子・KPI目標", 1),

    # ------------------------------------------ B 事業所調査を実施済みに
    ("ニーズ調査実施済／事業所実態調査は設計案作成済・実施前",
     "ニーズ調査・事業所実態調査とも実施済（回答受領済）", 1),
    ("稼働・待機・受入不可は事業所調査で補完",
     "稼働・待機・受入不可は事業所調査により補完", 1),
    ("欠員・採用・定着・縮小意向等は事業所調査で把握予定",
     "採用・定着等は事業所調査で把握済。"
     "欠員・縮小意向等は受領した回答に含まれない", 1),
    ("欠員、採用・定着、受入困難、縮小意向、ICT、BCP等",
     "実施済（回答受領済）。介護職員数、常勤・非常勤、外国人、派遣、"
     "採用・離職", 1),
    ("集計仕様、調査開始、評価根拠", "集計仕様、評価根拠", 1),
    ("対象・日程・公表範囲を確定し、供給・人材・経営・ICT・BCPを集計",
     "公表範囲を確定し、受領した回答を集計", 1),
    ("標準項目＋供給・人材・経営・ICT・BCP等を確認",
     "介護職員数、常勤・非常勤、外国人、派遣、採用・離職を確認", 1),
    ("対象事業所一覧、送付・回収方法、照会先、実施・督促日程",
     "回答データ又は集計表、対象事業所一覧、コード表、公表単位", 1),
    ("① 調査の実施", "① 調査データの受渡し", 1),
    ("事業所実態調査票、対象、調査期間、督促、回収・公表単位",
     "事業所実態調査の回答データの受渡し方法、集計単位、公表単位", 1),
    ("実施済みニーズ調査は町別・年齢階級別等の集計へ進めます。",
     "実施済みニーズ調査は集計・分析を完了しました。"
     "事業所実態調査も回答を受領済みであり、"
     "介護職員数、常勤・非常勤、採用・離職の集計を完了しています。", 1),

    # ------------------------------------------ C 調査の状況
    ("実施済・分析待ち", "実施済・集計分析済", 1),
    ("データ受領、町別・属性別集計、KPI基準値への接続",
     "KPI基準値への接続、計画本文への反映", 1),
    ("・ニーズ調査の提供データ・集計単位を確認し、"
     "町別・属性別集計を開始します。",
     "・実施済み調査の集計結果を計画本文へ反映します。", 1),

    # ------------------------------------------ D 記載の不足・不整合
    ("計画骨子案、KPI定義書、3町役割分担表、見える化整合まで作業済み",
     "計画骨子案、計画素案、KPI定義書、3町役割分担表、"
     "見える化整合まで作業済み", 1),
    ("委員会意見、実績及び調査結果を反映した計画第9稿",
     "委員会意見、実績及び調査結果を反映した計画素案の次稿", 1),
    ("第1回委員会資料・工程を準備中", "第1回委員会資料・工程を準備済", 1),
    ("第1回委員会資料・キックオフ\t", "第1回委員会資料・キックオフ資料 作成済\t",
     0),                      # セル内に単独で存在するため下で個別に処理

    # ------------------------------------------ E 表記の統一
    ("ＪＡＧＥＳ調査", "JAGES調査", 1),
    ("広域連合：【要確認】　　各町：【要確認】",
     "広域連合：【要確認】　各町：【要確認】", 1),
    ("スケジュール・レビュー方法等を（令和8年8月6日現在）共有し",
     "スケジュール・レビュー方法等について、"
     "令和8年8月6日現在の状況を共有し", 1),
    ("計画素案、及びKPI定義書まで作業を行い",
     "計画素案及びKPI定義書まで作業を行い", 1),

    # ------------------------------------------ F ご指示（5.・11.の削除）
    ("自然体推計は見える化6時点と差額0円。\n基金等は未反映・要協議",
     "基金等は未反映・要協議", 0),
    ("第10期平均6,238円、長期5時点を含む公表値6点と差額0円",
     "第10期平均6,238円", 1),
]

doc = Document(SRC)

done = {}
for old, new, must in REPS:
    n = 0
    for p in _iter_paragraphs(doc):
        n += replace_in_paragraph(p, old, new)
    done[old] = n

# 表10「工程・会議資料」の「現在の状況」セル（セル内で完結するため個別に処理）
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            if c.text.strip() == "第1回委員会資料・キックオフ":
                for p in c.paragraphs:
                    if p.text.strip() != "第1回委員会資料・キックオフ":
                        continue
                    # 置換後の文字列が置換前を含むため、末尾に追記する
                    p.runs[-1].text = (p.runs[-1].text.rstrip()
                                       + "資料 作成済")
                    done["工程・会議資料の状況"] = 1

# 表4「保険料・財政」（セル内で改行しているため段落単位で処理）
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in list(c.paragraphs):
                if p.text.strip() == "自然体推計は見える化6時点と差額0円。":
                    for r in p.runs:
                        r.text = ""
                    p._element.getparent().remove(p._element)
                    done["5.の差額0円の削除"] = 1

# 目次の頁番号を半角に統一
ZEN = "０１２３４５６７８９"
for p in doc.paragraphs:
    if "・・・" not in p.text:
        continue
    for r in p.runs:
        if any(ch in r.text for ch in ZEN):
            r.text = r.text.translate(str.maketrans(ZEN, "0123456789"))
            done["目次の頁番号"] = done.get("目次の頁番号", 0) + 1

doc.save(OUT)

# 原本から引き継いだ文書プロパティ（題名の稿番号・担当者名）を整える
from docmeta import clean_docx                    # noqa: E402
clean_docx(OUT,
           title="大雪地区広域連合　第10期介護保険事業計画"
                 "　キックオフ会議資料（令和8年8月）",
           subject="サービス見込量・保険料の検討手順及び工程")

print("saved:", OUT)
print()
ng = [k for (k, _n, must) in REPS if must and not done.get(k)]
for old, new, must in REPS:
    n = done.get(old, 0)
    mark = "OK " if n else ("未置換" if must else "―  ")
    print("%s %d件  %s" % (mark, n, old[:48].replace("\n", "/")))
for k in ["工程・会議資料の状況", "5.の差額0円の削除", "目次の頁番号"]:
    print("OK  %d件  %s" % (done.get(k, 0), k))
if ng:
    print("\n置換できなかった項目があります:", ng)
