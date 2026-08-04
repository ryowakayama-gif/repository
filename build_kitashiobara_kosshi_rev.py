# -*- coding: utf-8 -*-
"""
北塩原村 第8期障がい福祉計画・第4期障がい児福祉計画
骨子案 修正版ジェネレータ

入力: source/北塩原村_骨子案_原本_20260731.docx（他担当者作成の骨子案）
出力: output/北塩原村_骨子案_修正版.docx

比較レビュー（docs/北塩原村_骨子案_計画素案第9稿_比較レビュー.md）で
「骨子案を正本とし、計画素案第9稿から移植する」と整理した項目のうち、
村資料の受領を待たずに着手できるものを反映する。

本スクリプトが行う修正
  M-1 障がい者手帳所持者数の将来推計（仕様書4-Ⅱ(3)①・4-Ⅱ(5)③の必須事項。原本に欠落）
  M-2 人口データの更新（令和5年→令和6年実績、令和11年の村公式推計を追加）
  M-3 財源構成（令和8年7月7日打合せ「国・県・村の財政負担の推移を可視化」）
  M-6 こども・子育て計画／こども家庭センターとの接続

原本は書き換えず、常に output 側に新規生成する。書式（フォント・罫線・
見出し・注記ボックス）は原本の体裁に合わせて生成する。
"""

import os

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC_FILE = "/home/user/repository/source/北塩原村_骨子案_原本_20260731.docx"
OUT_DIR = "/home/user/repository/output"
OUT_FILE = f"{OUT_DIR}/北塩原村_骨子案_修正版.docx"

FONT = "BIZ UDPゴシック"
TABLE_W = 9638          # 原本の表幅（dxa）
BORDER_COLOR = "AAAAAA"
HEADER_FILL = "1F3864"
NOTE_BORDER = "C9A03B"
NOTE_FILL = "FFF2CC"

changes = []


# ============================================================
# XML ヘルパー
# ============================================================
def _el(tag, **attrs):
    e = OxmlElement(tag if ":" in tag else f"w:{tag}")
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), str(v))
    return e


def _set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = _el("rFonts")
        rpr.insert(0, rf)
    for a in ("ascii", "eastAsia", "hAnsi"):
        rf.set(qn(f"w:{a}"), FONT)


def _tbl_borders(color):
    b = _el("tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b.append(_el(side, val="single", sz="4", space="0", color=color))
    return b


def _cell_shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(_el("shd", val="clear", color="auto", fill=fill))


def _cell_margins(cell, top=60, left=100, bottom=60, right=100):
    mar = _el("tcMar")
    for name, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el(name, w=v, type="dxa"))
    tcpr = cell._tc.get_or_add_tcPr()
    tcpr.append(mar)
    tcpr.append(_el("vAlign", val="center"))


def _detach(doc, element):
    """python-docx が本文末尾に追加した要素を切り離して返す。"""
    doc.element.body.remove(element)
    return element


def make_para(doc, text, style=None, size=10.5, bold=False, indent=True):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    if style is None and indent and text:
        p.paragraph_format.first_line_indent = Pt(10.5)
    return _detach(doc, p._element)


def make_table(doc, rows, widths=None, header=True, border=BORDER_COLOR):
    """原本のデータ表と同じ体裁の表を作る。rows は文字列の2次元リスト。"""
    ncol = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=ncol)
    tblpr = t._tbl.tblPr
    for old in tblpr.findall(qn("w:tblBorders")):
        tblpr.remove(old)
    tblpr.append(_tbl_borders(border))
    t._tbl.tblPr.append(_el("tblLook", val="0000", firstRow="0", lastRow="0",
                            firstColumn="0", lastColumn="0", noHBand="0", noVBand="0"))
    if widths is None:
        widths = [TABLE_W // ncol] * ncol

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            tcpr = cell._tc.get_or_add_tcPr()
            for old in tcpr.findall(qn("w:tcW")):
                tcpr.remove(old)
            tcpr.append(_el("tcW", w=widths[ci], type="dxa"))
            _cell_margins(cell)
            para = cell.paragraphs[0]
            is_head = header and ri == 0
            if is_head or ci == 0 and ncol > 2:
                para.alignment = 1 if is_head else 0
            if is_head:
                para.alignment = 1
            run = para.add_run(str(val))
            _set_font(run, size=9.5, bold=is_head)
            if is_head:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _cell_shade(cell, HEADER_FILL)
            elif ri > 0 and ci > 0 and str(val).replace(",", "").replace("人", "").replace(
                    "千円", "").replace("％", "").replace(".", "").replace("+", "").replace("－", "").isdigit():
                para.alignment = 2  # 数値は右寄せ
    return _detach(doc, t._tbl)


def make_note(doc, text):
    """【要更新・要確認】注記ボックス（1x1表）を原本と同じ体裁で作る。"""
    t = doc.add_table(rows=1, cols=1)
    tblpr = t._tbl.tblPr
    for old in tblpr.findall(qn("w:tblBorders")):
        tblpr.remove(old)
    tblpr.append(_tbl_borders(NOTE_BORDER))
    cell = t.cell(0, 0)
    cell._tc.get_or_add_tcPr().append(_el("tcW", w=TABLE_W, type="dxa"))
    _cell_shade(cell, NOTE_FILL)
    _cell_margins(cell, 120, 160, 120, 160)
    run = cell.paragraphs[0].add_run(text)
    _set_font(run, size=9.5)
    return _detach(doc, t._tbl)


def make_empty(doc):
    p = doc.add_paragraph()
    return _detach(doc, p._element)


# ============================================================
# 位置検索・挿入
# ============================================================
def find_para(doc, prefix, style=None):
    """本文中から、指定の文字列で始まる段落の XML 要素を返す。"""
    for child in doc.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        p = Paragraph(child, doc)
        if p.text.strip().startswith(prefix) and (style is None or p.style.name == style):
            return child
    raise LookupError(f"段落が見つかりません: {prefix!r}")


def find_block_after(doc, anchor, tag):
    """anchor 以降で最初に現れる指定タグの要素を返す。"""
    seen = False
    for child in doc.element.body.iterchildren():
        if child is anchor:
            seen = True
            continue
        if seen and child.tag == qn(f"w:{tag}"):
            return child
    raise LookupError(f"{tag} が見つかりません")


def insert_after(anchor, elements):
    """anchor の直後に elements を順番どおり挿入する。"""
    cur = anchor
    for e in elements:
        cur.addnext(e)
        cur = e
    return cur


def insert_before(anchor, elements):
    for e in elements:
        anchor.addprevious(e)
    return anchor


def replace_note_text(doc, tbl_element, text):
    """既存の【要更新・要確認】ボックスの本文を差し替える。"""
    t = Table(tbl_element, doc)
    para = t.cell(0, 0).paragraphs[0]
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    run = para.add_run(text)
    _set_font(run, size=9.5)


def replace_para_text(doc, element, paragraphs):
    """段落の本文を差し替える。1段落を複数段落に分ける場合はリストで渡す。"""
    p = Paragraph(element, doc)
    for run in list(p.runs):
        run._r.getparent().remove(run._r)
    run = p.add_run(paragraphs[0])
    _set_font(run)
    p.paragraph_format.first_line_indent = Pt(10.5)
    extra = [make_para(doc, t) for t in paragraphs[1:]]
    insert_after(element, extra)
    return extra[-1] if extra else element


# ============================================================
# M-2 人口データの更新
# ============================================================
def revise_population(doc):
    anchor = find_para(doc, "本村の人口（住民基本台帳）は、令和５年４月１日現在2,443人")
    replace_para_text(doc, anchor, [
        "本村の人口（住民基本台帳）は、令和６年４月１日現在2,394人であり、平成31年（2,743人）と比較すると"
        "349人の減少となっています。年少人口（０〜14歳）は令和６年に180人まで減少し、令和３年以降も"
        "毎年20人以上減少するペースが続いています。老年人口（65歳以上）は令和３年の1,006人をピークに"
        "緩やかな減少に転じていますが、年少人口・生産年齢人口の減少幅がそれを上回るため、高齢化率は"
        "上昇を続け、令和６年は41.4％（平成31年比＋6.3ポイント）に達しています。",
        "また、「北塩原村こども・子育て計画」（令和７年３月策定）が示すコーホート変化率法による将来推計では、"
        "令和11年度の総人口は2,185人、年少人口は111人（構成比5.1％）、老年人口は1,051人（構成比48.1％）に"
        "なると見込まれています。令和８年３月に告示された国の基本指針が重視する"
        "「人口減少地域におけるサービスの維持・確保」の視点は、本村においても本計画期間を通じた重要な課題となります。",
    ])

    note = find_block_after(doc, anchor, "tbl")  # 【要更新・要確認】ボックス
    rows = [
        ["区分", "平成31年", "令和５年", "令和６年", "令和11年（推計）"],
        ["総人口", "2,743人", "2,443人", "2,394人", "2,185人"],
        ["年少人口（０〜14歳）", "284人", "205人", "180人", "111人（5.1％）"],
        ["生産年齢人口（15〜64歳）", "1,495人", "1,241人", "1,223人", "1,023人（46.8％）"],
        ["老年人口（65歳以上）", "964人", "997人", "991人", "1,051人（48.1％）"],
        ["高齢化率", "35.1％", "40.8％", "41.4％", "48.1％"],
    ]
    widths = [2600, 1750, 1750, 1750, 1788]
    insert_before(note, [
        make_para(doc, "■　人口の推移と将来推計", style="Heading 3"),
        make_table(doc, rows, widths),
        make_para(doc, "資料：住民基本台帳（各年４月１日現在）。令和11年の推計は「北塩原村こども・子育て計画」"
                       "（コーホート変化率法）による村の公式推計。", size=9, indent=False),
        make_empty(doc),
    ])
    replace_note_text(doc, note,
                      "【要更新・要確認】本文及び上表は令和６年４月１日現在の実績まで更新済みです。"
                      "ただし上部のグラフは第７期計画（令和５年度時点）のデータのままであるため、"
                      "令和６〜８年の実績を反映して作成し直してください。"
                      "あわせて、令和８年４月１日現在の人口データを受領のうえ確定し、"
                      "村独自の人口ビジョン・将来人口推計がある場合は、"
                      "こども・子育て計画の推計に代えてそちらを用いて再計算してください。")
    changes.append("M-2 第2章1：人口データを令和6年実績に更新し、令和11年の村公式推計を追加"
                   "（グラフは未更新のため注記を差し替え）")


# ============================================================
# M-1 障がい者手帳所持者数の将来推計（新設）
# ============================================================
def add_tegata_projection(doc):
    # 「６　障がい児の就学状況」の直前に新設し、以降の節番号を繰り下げる
    anchor = find_para(doc, "６　障がい児の就学状況", style="Heading 2")

    blocks = [
        make_para(doc, "６　障がい者手帳所持者数の将来推計", style="Heading 2"),
        make_para(doc,
                  "本節は、業務委託仕様書が計画案の記載事項として定める"
                  "「障がい者手帳所持者数（身体・療育・精神）の将来推計」に対応するものです。"
                  "本村は対象者数が少なく年による振れ幅が大きいため、性質の異なる２つの方法を併用し、"
                  "幅を持たせて提示します。"),
        make_para(doc, "■　推計に用いた実績", style="Heading 3"),
        make_table(doc, [
            ["区分", "平成31年", "令和２年", "令和３年", "令和４年", "令和５年", "５年間の傾向"],
            ["身体障害者手帳", "126人", "121人", "128人", "122人", "114人", "緩やかな減少傾向"],
            ["療育手帳", "13人", "13人", "14人", "12人", "12人", "横ばい"],
            ["精神障害者保健福祉手帳", "24人", "28人", "27人", "26人", "30人", "緩やかな増加傾向"],
            ["合計", "163人", "162人", "169人", "160人", "156人", "－"],
        ], [2100, 1050, 1050, 1050, 1050, 1050, 2288]),
        make_para(doc, "資料：現行計画本編（各年４月１日現在）。", size=9, indent=False),
        make_empty(doc),
        make_para(doc, "■　推計の方法", style="Heading 3"),
        make_table(doc, [
            ["方法", "考え方", "特性"],
            ["方法Ａ　トレンド延長法",
             "平成31年から令和５年までの５か年について、手帳所持者数（実数）の推移を直線回帰し、"
             "令和９〜11年度まで延長する。",
             "直近の実数の動き（精神の増加、身体の減少）をそのまま反映するが、単年の変動に引きずられやすい。"],
            ["方法Ｂ　所持率法",
             "令和５年時点の人口に対する手帳所持率が今後も一定と仮定し、"
             "村の公式な将来人口推計（本章１）に乗じる。",
             "人口減少の影響を反映するが、所持率自体の変化（精神の増加傾向等）は捉えない。"],
        ], [1900, 3900, 3838], header=True),
        make_empty(doc),
        make_para(doc, "■　推計結果（暫定）", style="Heading 3"),
        make_table(doc, [
            ["区分", "手法", "令和９年度末", "令和10年度末", "令和11年度末"],
            ["身体障害者手帳", "方法Ａ（トレンド延長）", "108人", "106人", "104人"],
            ["", "方法Ｂ（所持率法）", "106人", "104人", "102人"],
            ["", "中位（Ａ・Ｂの平均）", "107人", "105人", "103人"],
            ["療育手帳", "方法Ａ（トレンド延長）", "11人", "11人", "10人"],
            ["", "方法Ｂ（所持率法）", "11人", "11人", "11人"],
            ["", "中位（Ａ・Ｂの平均）", "11人", "11人", "11人"],
            ["精神障害者保健福祉手帳", "方法Ａ（トレンド延長）", "33人", "34人", "35人"],
            ["", "方法Ｂ（所持率法）", "28人", "27人", "27人"],
            ["", "中位（Ａ・Ｂの平均）", "31人", "31人", "31人"],
            ["合計", "中位推計", "149人", "147人", "145人"],
        ], [2100, 2338, 1700, 1750, 1750]),
        make_empty(doc),
        make_para(doc, "■　読み取りのポイント", style="Heading 3"),
        make_para(doc,
                  "・中位推計では、手帳所持者数の合計は令和５年の156人から令和11年度末には145人程度まで"
                  "緩やかに減少する見通しです。ただし方法Ａ・方法Ｂの差が示すとおり幅のある暫定推計であり、"
                  "単一の数値として確定表記することは避け、計画本文では幅を併記します。", indent=False),
        make_para(doc,
                  "・老年人口が令和３年をピークに横ばい〜微減で推移する一方、身体障害者手帳所持者数は"
                  "平成31年の126人から令和５年の114人へ減少しており、一見すると整合しない動きになっています。"
                  "手帳の更新時期、制度改正、転出入等の影響である可能性があるため、令和６〜８年度の実績で"
                  "継続的な傾向か一時的な変動かを確認します。", indent=False),
        make_para(doc,
                  "・精神障害者保健福祉手帳は、人口減少が続く中でも所持者数・所持率とも増加傾向にあります。"
                  "相談・受診機会の広がりを踏まえると次期計画期間も同様の傾向が続く可能性があり、"
                  "第４章（２）の精神障害にも対応した地域包括ケアシステムの構築及び相談支援体制の強化と"
                  "整合させて記載します。", indent=False),
        make_empty(doc),
        make_note(doc,
                  "【要更新・要確認】本推計は、現行計画本編に記載された平成31年〜令和５年の実績のみを用いた"
                  "暫定値です。村が保有する令和６〜８年度の手帳所持者数実績及び年齢階層別内訳を受領のうえ、"
                  "確定値に更新してください。あわせて、業務委託仕様書に記載されたアンケート調査対象者数"
                  "（身体約130人・精神約40人・療育約20人＝計約190人）と、本表の令和５年時点156人との差について、"
                  "年数経過による自然増か、集計基準・対象年齢・サービス利用者の重複計上の違いによるものかを"
                  "確認してください。"),
        make_empty(doc),
    ]
    insert_before(anchor, blocks)

    # 節番号の繰り下げ
    for old, new in (("６　障がい児の就学状況", "７　障がい児の就学状況"),
                     ("７　アンケート調査の実施について", "８　アンケート調査の実施について")):
        el = find_para(doc, old, style="Heading 2")
        p = Paragraph(el, doc)
        for run in p.runs:
            if old[:2] in run.text or run.text.strip().startswith(old[0]):
                run.text = run.text.replace(old[0], new[0], 1)
                break
        else:
            p.runs[0].text = new
    changes.append("M-1 第2章：「６　障がい者手帳所持者数の将来推計」を新設（節番号６〜７を７〜８へ繰り下げ）")


# ============================================================
# M-6 こども・子育て計画／こども家庭センター
# ============================================================
def add_kodomo_katei_center(doc):
    # (1) 第1章2 計画の位置づけ
    anchor = find_para(doc, "また、「第３次健康21・北塩原村グッドヘルスプラン」の計画位置づけ図")
    insert_after(anchor, [
        make_para(doc,
                  "また、令和７年３月に策定された「北塩原村こども・子育て計画」（計画期間：令和７年度〜令和11年度、"
                  "令和９年度に中間見直し）は、こども基本法第10条に基づく市町村こども計画であり、"
                  "施策体系「Ⅰ(8)　援助を必要とするこどもや家庭への支援」に基本施策「①障がい児支援・"
                  "医療的ケア児等への支援」を掲げています。同計画は、障がい児への支援に特化した施設が"
                  "遠方に立地しており保護者やこどもの身体的負担が課題であることを明記しており、"
                  "本計画が前提とする圏域連携の方針と一致しています。さらに同計画は、令和８年度に"
                  "母子保健と児童福祉を一体的に担う「こども家庭センター」を新設する方針を示しています。"
                  "本計画では、同センターを障がいや発達の特性の早期発見・早期療育に係る相談の入口として"
                  "位置づけ、第４章（４）及び第７章で連携方法を整理します。なお、両計画の基本理念は"
                  "同一であり、整合が図られていることを確認しています。"),
    ])
    changes.append("M-6 第1章2：こども・子育て計画及びこども家庭センターとの関係を追記")

    # (2) 第4章（４）障がい児支援の提供体制の整備等
    anchor = find_para(doc, "本村には児童発達支援センターの設置がなく（０か所）")
    insert_after(anchor, [
        make_para(doc,
                  "また、令和８年度に設置予定の「こども家庭センター」は、母子保健と児童福祉を一体的に担う"
                  "窓口であり、障がいや発達の特性の早期発見・早期療育に係る相談の入口となることが"
                  "見込まれます。乳幼児健診・母子保健事業から障がい児通所支援・障がい児相談支援への"
                  "相談導線を明確にし、こども家庭センター、保健センター、教育委員会、圏域の"
                  "児童発達支援センター及び相談支援事業所の役割分担を整理します。基本指針が新たに求める"
                  "「障がい児等への伴走的な相談支援体制の構築（のぞまないセルフプランの解消）」についても、"
                  "同センターを起点とした相談体制の中で対応します。医療的ケア児等コーディネーターの"
                  "配置検討にあたっても、同センターとの連携を前提とします。"),
        make_note(doc,
                  "【要更新・要確認】こども家庭センターの設置時期、所管課、配置人員及び障がい児に係る"
                  "相談の受付フロー（乳幼児健診からの引継ぎ、圏域の児童発達支援センターへのつなぎ方）を"
                  "村に確認し、記載を確定してください。"),
    ])
    changes.append("M-6 第4章（４）：こども家庭センターとの相談導線・役割分担を追記")

    # (3) 第7章 他計画との連携
    anchor = find_para(doc, "第8章　資料編", style="Heading 1")
    insert_before(anchor, [
        make_para(doc, "３　北塩原村こども・子育て計画との連携", style="Heading 2"),
        make_para(doc,
                  "「北塩原村こども・子育て計画」（令和７年３月策定、計画期間：令和７年度〜令和11年度）は、"
                  "こども基本法第10条に基づく市町村こども計画です。計画期間が本計画（令和９〜11年度）と"
                  "重なり、令和９年度に中間見直しが予定されているため、本計画の策定内容を中間見直しに"
                  "反映できるよう、庁内で情報を共有します。"),
        make_table(doc, [
            ["項目", "こども・子育て計画", "本計画（障がい福祉・障がい児福祉計画）"],
            ["計画期間", "令和７年度〜令和11年度（令和９年度に中間見直し）", "令和９年度〜令和11年度"],
            ["基本理念",
             "障がいのあるなしに関わらず、お互いの人格や個性を尊重し、多様な価値観を認め合い、"
             "誰もが自分らしく輝くむら",
             "（第４次障がい者計画）同左。両計画で理念が一致していることを確認済み"],
            ["障がい児支援の位置づけ",
             "施策体系「Ⅰ(8)　援助を必要とするこどもや家庭への支援」の基本施策"
             "「①障がい児支援・医療的ケア児等への支援」",
             "第４章（４）障がい児支援の提供体制の整備等、第５章５　障がい児支援"],
            ["相談窓口",
             "令和８年度に「こども家庭センター」を新設（母子保健・児童福祉の一体窓口）",
             "こども家庭センターを障がい児の早期把握・相談の入口として位置づけ"],
        ], [1700, 3900, 4038]),
        make_empty(doc),
        make_para(doc, "具体的な連携事項", style="Heading 3"),
        make_para(doc,
                  "・早期発見・早期療育：こども家庭センター及び保健センターの乳幼児健診・母子保健事業から、"
                  "障がい児通所支援・障がい児相談支援へつなぐ相談導線を整理します。", indent=False),
        make_para(doc,
                  "・圏域連携：こども・子育て計画が課題として挙げる「障がい児への支援に特化した施設が"
                  "遠方に立地していること」への対応として、会津北部圏域（猪苗代町・磐梯町・湯川村・北塩原村）"
                  "での児童発達支援センターの４つの中核機能へのアクセス確保を進めます。", indent=False),
        make_para(doc,
                  "・医療的ケア児支援：医療的ケア児等コーディネーターの配置検討を、こども家庭センターの"
                  "母子保健機能と一体で進めます。", indent=False),
        make_para(doc,
                  "・18歳以降への移行：特別支援学校卒業後の進路（就労選択支援、就労継続支援、生活介護等）"
                  "について、在学中から教育委員会・相談支援事業所と連携して準備を進めます。", indent=False),
        make_para(doc,
                  "・インクルージョンの推進：基本指針が新たに求める「インクルージョン推進のための協議の場」"
                  "について、こども・子育て計画の保育・教育施策と連動させて検討します。", indent=False),
        make_empty(doc),
    ])
    changes.append("M-6 第7章：「３　北塩原村こども・子育て計画との連携」を新設")


# ============================================================
# M-3 財源構成（新設）
# ============================================================
def add_funding_section(doc):
    anchor = find_para(doc, "第6章　成年後見制度の利用促進", style="Heading 1")
    blocks = [
        make_para(doc, "９　障がい福祉サービス等に係る費用の推移と財源構成", style="Heading 2"),
        make_para(doc,
                  "これまでの計画では、障がい福祉サービスの利用件数・利用量の推移を中心に整理してきましたが、"
                  "本計画では、サービスの提供に要する費用と、その財源が国・県・村にどのように配分されているかを"
                  "あわせて示します。サービス見込量の確保は財源の確保と一体であり、"
                  "特に人口規模の小さい本村では、給付費の増減が村財政に与える影響を把握しておく必要があるためです。"),
        make_para(doc, "■　給付費の推移", style="Heading 3"),
        make_table(doc, [
            ["年度", "介護給付費等", "障害児給付費", "計", "令和２年度比"],
            ["令和２年度", "36,002千円", "746千円", "36,748千円", "－"],
            ["令和３年度", "45,992千円", "3,255千円", "49,246千円", "＋34.0％"],
            ["令和４年度", "44,234千円", "4,030千円", "48,264千円", "＋31.3％"],
            ["令和５年度", "44,420千円", "3,578千円", "47,998千円", "＋30.6％"],
            ["令和６年度", "52,699千円", "1,641千円", "54,340千円", "＋47.9％"],
            ["令和７年度", "53,877千円", "1,636千円", "55,513千円", "＋51.1％"],
        ], [1700, 2100, 1900, 2000, 1938]),
        make_para(doc, "資料：村提供の障がいサービス給付実績（千円未満四捨五入）。", size=9, indent=False),
        make_empty(doc),
        make_para(doc,
                  "介護給付費等は令和２年度の36,002千円から令和７年度の53,877千円へ、"
                  "約49.6％増加しています。増加の中心は就労継続支援（Ｂ型）であり、"
                  "令和２年度の9,459千円から令和７年度の19,707千円へ約2.08倍となっています。"
                  "一方、障害児給付費は令和４年度の4,030千円をピークに減少し、"
                  "令和６・７年度は1,600千円台で推移しています。"),
        make_para(doc, "■　財源構成（令和７年度・法定負担割合による試算）", style="Heading 3"),
        make_table(doc, [
            ["区分", "金額", "構成比", "根拠"],
            ["自立支援給付費・障害児給付費　計", "55,513千円", "100.0％", "－"],
            ["　国庫負担", "27,756千円", "50.0％", "障害者総合支援法第92条〜第95条、児童福祉法第57条の２等"],
            ["　県負担", "13,878千円", "25.0％", "同上"],
            ["　村負担", "13,878千円", "25.0％", "同上"],
        ], [3000, 1700, 1400, 3538]),
        make_empty(doc),
        make_para(doc, "■　サービス別の構成（令和７年度）", style="Heading 3"),
        make_table(doc, [
            ["サービス", "給付費", "構成比", "計画上の位置づけ"],
            ["就労継続支援（Ｂ型）", "19,707千円", "35.5％", "村内に事業所がなく圏域事業所を利用。増加の主因"],
            ["生活介護", "13,210千円", "23.8％", "重度者の日中活動。卒業後の進路と接続"],
            ["共同生活援助", "9,663千円", "17.4％", "居住系の中心。地域移行・親亡き後支援と接続"],
            ["施設入所支援", "5,994千円", "10.8％", "４人で横ばい。第４章（１）の成果目標と整合"],
            ["計画相談支援", "1,864千円", "3.4％", "圏域の相談支援事業所が実施"],
            ["居宅介護", "1,713千円", "3.1％", "訪問可能地区・冬季の提供体制が制約"],
            ["放課後等デイサービス", "1,076千円", "1.9％", "令和７年度に利用が増加"],
            ["その他", "2,286千円", "4.1％", "短期入所、自立訓練、児童発達支援、障害児相談支援等"],
        ], [2500, 1600, 1200, 4338]),
        make_empty(doc),
        make_para(doc, "■　地域生活支援事業及び村単独事業の財源", style="Heading 3"),
        make_para(doc,
                  "自立支援給付及び障害児通所給付が義務的経費であるのに対し、"
                  "地域生活支援事業（第５章７）は国２分の１以内・県４分の１以内の統合補助金による裁量的経費です。"
                  "補助基準額を超える部分は村の一般財源で賄うこととなるため、"
                  "事業の評価にあたっては、実施の有無だけでなく、事業費、国県補助額、村負担額、"
                  "超過負担の有無を分けて把握する必要があります。"),
        make_para(doc,
                  "また、重度心身障害者医療費助成、在宅重度障害者対策事業、人工透析患者通院交通費助成等の"
                  "村単独事業は、障がいのある方の生活継続を支える基盤であり、"
                  "対象者数と村負担額を継続的に把握します。"),
        make_note(doc,
                  "【要更新・要確認】上表の財源構成は、給付費に法定負担割合を乗じた試算です。"
                  "訪問系サービス等で国庫負担基準額を超過している場合、超過分は村負担となるため、"
                  "実際の決算額とは一致しません。障害者自立支援給付費負担金・障害児施設給付費負担金・"
                  "地域生活支援事業費補助金の令和６・７年度交付額と村の一般財源負担額を確認のうえ、"
                  "数値を確定してください。あわせて、地域生活支援事業の事業別の事業費・補助基準額・"
                  "村負担額、村単独事業の対象者数・事業費・県補助額を確認してください。"),
        make_empty(doc),
    ]
    insert_before(anchor, blocks)
    changes.append("M-3 第5章：「９　障がい福祉サービス等に係る費用の推移と財源構成」を新設")


# ============================================================
# 付随する確認事項
# ============================================================
def add_survey_design(doc):
    """第2章8 アンケート調査の実施について：7月7日打合せで決定した実施設計を反映する。"""
    anchor = find_para(doc, "・医療的ケア児・強度行動障害を有する方等、重度障がい者の個別ニーズ")
    insert_after(anchor, [
        make_empty(doc),
        make_para(doc,
                  "なお、本計画の策定に係るアンケート調査については、"
                  "令和８年７月７日の打合せにおいて、次の実施設計を決定しています。"),
        make_table(doc, [
            ["項目", "内容"],
            ["調査対象",
             "身体障害者手帳所持者 約130人、精神障害者保健福祉手帳所持者 約40人、"
             "療育手帳所持者 約20人（計 約190人）"],
            ["調査票の構成",
             "回答者の意向をより明確に把握するため、前回の一体型設計を見直し、"
             "「障がい者用」と「障がい児用」の２種類に分けて作成する"],
            ["回答方法",
             "紙の調査票に加え、二次元コードから回答できるウェブ回答を併用する。"
             "ウェブ回答では設問の条件分岐に対応する。管理番号により紙とウェブの重複回答を除外する"],
            ["障がい特性への配慮",
             "文字サイズは14ポイント程度とし、必要なルビを付与する。色覚特性に配慮し白黒印刷を基本とする。"
             "専門用語には、回答者が迷わないよう機能に即した短い解説を付記する"],
            ["発送・回収",
             "村の封筒を使用し、差出人及び返送先を北塩原村役場とする。"
             "回収した調査票は村役場で一括して受領・保管する"],
            ["集計上の扱い",
             "重複回答・選択件数超過の回答は集計から除外せず、"
             "報告書の欄外にその件数を注記として明記する"],
            ["実施時期", "令和８年８月上旬に発送し、約１か月の回収期間を設定する"],
        ], [2200, 7438]),
        make_empty(doc),
        make_para(doc,
                  "また、本村は対象者数が少なく、１〜２人の増減で割合が大きく変動するため、"
                  "集計・分析にあたっては割合のみで判断せず実人数を併記します。"
                  "対象者が極めて少ないサービスや地区については、クロス集計により個人が特定される"
                  "おそれがあるため、公表する報告書では属性を統合して扱い、個別の情報は"
                  "内部管理用の資料でのみ扱います。また、回答が０件であった項目についても、"
                  "制度の未周知、移動の制約など実態から想定される理由を付記し、"
                  "過小な推計とならないようにします。"),
    ])
    changes.append("C-10 第2章8：7月7日打合せで決定したアンケート実施設計と少数母数の取扱いを反映")


def add_misc_notes(doc):
    # 短期入所・就労継続支援B型の実績を見込量表の注記として追加
    anchor = find_para(doc, "第6章　成年後見制度の利用促進", style="Heading 1")
    insert_before(anchor, [
        make_note(doc,
                  "【要更新・要確認】給付実績との突合が必要な事項があります。"
                  "①短期入所は現行計画で見込量０人としていますが、令和６年度に２件、令和７年度に３件の"
                  "給付実績があります。令和９年度からの計上を検討してください。"
                  "②就労継続支援（Ｂ型）は、給付費が令和２年度比で約2.08倍に増加している一方、"
                  "本章２の見込量は11人分・200人日分で据置きとしています。"
                  "利用者数の増、利用日数の増、報酬改定、重度化のいずれによるものかを確認し、"
                  "見込量の考え方を整理してください。"
                  "③本計画の成果目標（第４章（１））における施設入所者数の基準は令和７年度末です。"
                  "令和５年度末の４人ではなく、令和７年度末実績を基準として確定してください。"),
        make_empty(doc),
    ])
    changes.append("付随：短期入所・就労継続支援B型・施設入所者基準年の確認事項を第5章末に追記")

    # 奥付の所管課
    anchor = find_para(doc, "編集：北塩原村")
    insert_after(anchor, [
        make_note(doc,
                  "【要更新・要確認】本計画書の編集所管課の表記を確認してください。"
                  "業務委託仕様書の事務担当は保健福祉課福祉係、現行計画における障がい者相談支援事業の"
                  "実施主体は住民課と記載されています。"),
    ])
    changes.append("付随：奥付の編集所管課に確認注記を追加")

    # 冒頭の改訂履歴
    anchor = find_para(doc, "最終的な計画の確定にあたっては")
    insert_after(anchor, [
        make_empty(doc),
        make_para(doc, "本版での主な追加・修正", style="Heading 3"),
        make_table(doc, [
            ["追加・修正内容", "反映箇所"],
            ["障がい者手帳所持者数の将来推計を新設（業務委託仕様書の記載事項）", "第２章６"],
            ["人口データを令和６年実績に更新し、令和11年の村公式推計を追加", "第２章１"],
            ["障がい福祉サービス等に係る費用の推移と財源構成を新設", "第５章９"],
            ["こども・子育て計画及びこども家庭センターとの接続を追加", "第１章２、第４章（４）、第７章３"],
        ], [6500, 3138]),
        make_empty(doc),
    ])
    changes.append("付随：冒頭に本版での追加・修正の一覧を追加")


def main():
    if not os.path.exists(SRC_FILE):
        raise SystemExit(f"原本が見つかりません: {SRC_FILE}")
    os.makedirs(OUT_DIR, exist_ok=True)

    doc = docx.Document(SRC_FILE)
    revise_population(doc)
    add_tegata_projection(doc)
    add_survey_design(doc)
    add_kodomo_katei_center(doc)
    add_funding_section(doc)
    add_misc_notes(doc)
    doc.save(OUT_FILE)

    print(f"作成: {OUT_FILE}")
    print("反映した修正:")
    for c in changes:
        print(f"  - {c}")


if __name__ == "__main__":
    main()
