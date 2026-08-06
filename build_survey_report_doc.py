# -*- coding: utf-8 -*-
"""大雪地区広域連合 第10期介護保険事業計画　実施済み調査 結果報告書（Word）.

発注者から提示のあった様式（令和8年度 介護人材実態調査 結果報告書）の
スタイル・構成を基に、仕様書４（3）が定める4調査の集計・分析の結果を
報告書としてまとめる。

  ① 在宅生活改善調査　　　　事業所票15件・利用者票99票
  ② 居所変更実態調査　　　　施設等票18件
  ③ 介護人材実態調査　　　　事業所票27件・職員個票317人・職員票26件
  ④ 健康とくらしの調査　　　個票4,729票

4調査を横断したクロス集計は、利用者票の所在地区の記入形式が統一されておらず
個票を地区に割り付けられないため、発注者のご意向により行わない。
これに代えて、公表データによる供給構造の分析により調査相互を接続する。

様式（提示のあった報告書に合わせる）
  用紙　　21.6×27.9cm　余白 上下2.5cm・左右1.9cm
  本文　　BIZ UDPゴシック 10.5pt
  見出1　16pt 太字　　見出2　14pt 太字
  注記　　ＭＳ Ｐゴシック 11pt（灰色）

改ページの制御
  ・図表の表題は次の段落（図）と同じページに固定する（keepNext）。
  ・表は行の途中で改ページしない（cantSplit）。
    14行以下の表は表全体を同じページに収める。
  ・表とその図が1ページに収まる場合は、両者を同じページに固定する。
  ・見出しは直後の本文と同じページに置く。
"""

import collections
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import rcParams

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import data_survey2025 as S
import data_survey_cross as C
import data_mieruka_km as MK
import data_hokkaido_roster as R
import data_hokkaido_shitei as H

OUT = ("/home/user/repository/output/"
       "第10期計画_実施済み調査_結果報告書.docx")
FIGDIR = "/home/user/repository/output/figures_report"
os.makedirs(FIGDIR, exist_ok=True)

# 白黒印刷を前提としたグレースケール。前回計画及び図表集と同じ配色。
rcParams["font.family"] = "IPAGothic"
rcParams["axes.unicode_minus"] = False
rcParams["figure.dpi"] = 220
rcParams["savefig.dpi"] = 220
rcParams["savefig.bbox"] = "tight"
rcParams["savefig.pad_inches"] = 0.04
rcParams["font.size"] = 8.5
rcParams["axes.edgecolor"] = "#000000"
rcParams["axes.linewidth"] = 0.8

GRAYS = ["#595959", "#A6A6A6", "#D9D9D9", "#F2F2F2", "#7F7F7F", "#BFBFBF",
         "#4D4D4D", "#E6E6E6"]
HATCH = ["", "", "", "", "///", "...", "\\\\", "xxx"]
DARK = {0, 4, 6}          # 白文字にする濃い色の位置


def _ramp(n):
    """順序のある区分（要介護度等）に用いる淡→濃の階調。"""
    lo, hi = 0xF2, 0x4D
    out = []
    for i in range(n):
        v = int(lo + (hi - lo) * (i / max(1, n - 1)))
        out.append("#%02X%02X%02X" % (v, v, v))
    return out
FIGN = [0]
# 直前に描いた表。表とその図を同じページに収めるために用いる。
LAST = {"table": None, "spacer": None, "cm": 0.0}

BODY = "BIZ UDPゴシック"
NOTEF = "ＭＳ Ｐゴシック"
GRAY = RGBColor(0x66, 0x66, 0x66)
HEADFILL = "F2F2F2"

# ---------------------------------------------------------------- 集計
SIS = [j for j in S.JIN if j["区分"] == "施設・通所系"]
HOU = [j for j in S.JIN if j["区分"] != "施設・通所系"]
N_SIS = sum(j["職員"] or 0 for j in SIS)
N_HOU = sum(j["職員"] or 0 for j in HOU)
DUP = 13
N_ALL = N_SIS + N_HOU - DUP
SAIYO = sum(j["採用"] or 0 for j in S.JIN) - 7
RISHOKU = sum(j["離職"] or 0 for j in S.JIN) - 2
GAIKOKU = sum(j["外国人"] or 0 for j in S.JIN)
N_GAI_JIG = sum(1 for j in S.JIN if (j["外国人"] or 0) > 0)
SK = S.SHOKU

CAT = [
    "GH", "GH", "住宅型有料", "住宅型有料", "GH", "老健", "通所リハ", "GH",
    "サ高住", "通所介護等", "特定施設", "地密特養", "特定施設", "老健",
    "通所リハ", "通所介護等", "通所介護等", "特養", "地密特養", "特養",
    "通所リハ", "老健", "住宅型有料", "GH", "訪問介護", "訪問介護", "訪問介護",
]
BYCAT = collections.Counter()
for _c, _j in zip(CAT, S.JIN):
    BYCAT[_c] += _j["職員"] or 0
BYCAT["住宅型有料"] -= DUP
NOSER = sum(BYCAT[k] for k in ["GH", "特定施設", "住宅型有料", "サ高住"])

CAP_TOKUTEI = sum(y["定員"] for y in R.YU if y["類型"] == "介護付")
CAP_JUTAKU = sum(y["定員"] for y in R.YU if y["類型"] == "住宅型")
CAP_SAKO = sum(y["戸数"] for y in R.SA)
CAP_KEIHI = sum(y["定員"] for y in R.KE)
CAP_TOKUYO = sum(t["定員"] for t in H.TOKUYO if not t["地域密着型"])
CAP_CHITOKU = sum(t["定員"] for t in H.TOKUYO if t["地域密着型"])
CAP_ROKEN = sum(s["定員"] or 0 for s in S.SHI if s["種別"] == 7)
CAP_GH = sum(s["定員"] or 0 for s in S.SHI if s["種別"] == 4) + 18

SW = [k for k in H.KOHYO if k["事業所名"] == "さわやか東神楽館"][0]
SW_KD = SW["要介護度別入居者数"]
SW_N = sum(SW_KD.values())

UNIQ = {}
for _sv, _rows in H.SHITEI.items():
    for _r in _rows:
        UNIQ.setdefault((_r["法人"], _r["事業所名"]),
                        {"町": _r["町"]})
BYH = collections.Counter(k[0] for k in UNIQ)
N_JIG = len(UNIQ)
N_HOJIN = len(BYH)
TOP6 = BYH.most_common(6)
TOP6N = sum(v for _, v in TOP6)

TOWNS = ["東川町", "美瑛町", "東神楽町"]


def klast(code):
    v = MK.K[code]["値"]
    ks = [k for k in v if v[k] is not None]
    return int(v[ks[-1]])


# ---------------------------------------------------------------- 文書
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)   # A4判
sec.top_margin = sec.bottom_margin = Cm(2.2)
sec.left_margin = sec.right_margin = Cm(1.9)

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.font.color.rgb = RGBColor(0, 0, 0)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)

h1 = doc.styles["Heading 1"]
h1.font.name = BODY
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(0)

h2 = doc.styles["Heading 2"]
h2.font.name = BODY
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
h2.paragraph_format.space_before = Pt(10)
h2.paragraph_format.space_after = Pt(0)
h2.paragraph_format.left_indent = Cm(0.32)

nt = doc.styles.add_style("note", 1)
nt.base_style = doc.styles["Normal"]
nt.font.name = NOTEF
nt.font.size = Pt(11)
nt.font.color.rgb = GRAY
nt.element.rPr.rFonts.set(qn("w:eastAsia"), NOTEF)
nt.paragraph_format.space_before = Pt(11.25)
nt.paragraph_format.space_after = Pt(11.25)


def P(text="", size=None, bold=False, align=None, space_after=6,
      style=None, color=None):
    LAST["table"] = None
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    if size:
        r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if style != "note":
        r.font.name = BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
    return p


def H1(text):
    LAST["table"] = None
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def H2(text):
    LAST["table"] = None
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True
    return p


def SUB(text):
    """（１）〜 の太字小見出し。"""
    return P(text, bold=True, space_after=3)


def NOTE(text):
    return P(text, style="note", space_after=4)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def cell_text(cell, text, size=9, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)


def _keep_table(t, limit=14):
    """行の途中で改ページさせない。短い表は表全体を同じページに収める。"""
    whole = len(t.rows) <= limit
    for i, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if whole and i < len(t.rows) - 1:
            for c in row.cells:
                for p_ in c.paragraphs:
                    p_.paragraph_format.keep_with_next = True


TEXTW = 21.0 - 1.9 * 2                   # 本文幅 17.2cm


def TBL(head, rows, widths=None, size=9, num_from=1):
    t = doc.add_table(rows=0, cols=len(head))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.add_row().cells
    for i, h in enumerate(head):
        cell_text(hr[i], h, size=size, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(hr[i], HEADFILL)
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.RIGHT if i >= num_from else None
            cell_text(c[i], v, size=size, align=al)
    if widths:
        tot = sum(widths)
        if tot > TEXTW:                  # 本文幅に収まるよう比例配分する
            widths = [w * TEXTW / tot for w in widths]
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    _keep_table(t)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.space_before = Pt(0)
    LAST["table"] = t
    LAST["spacer"] = sp
    LAST["cm"] = 0.72 + 0.52 * (len(t.rows) - 1)
    return t


def _fin(fig, name):
    path = os.path.join(FIGDIR, name + ".png")
    fig.savefig(path, facecolor="white")
    plt.close(fig)
    return path


def _bind_table(fig_cm):
    """直前の表と図が1ページに収まる場合、同じページに固定する。"""
    t = LAST["table"]
    if t is None:
        return
    if LAST["cm"] + fig_cm + 1.6 > 21.0:
        return
    last = t.rows[-1]
    for c in last.cells:
        for p_ in c.paragraphs:
            p_.paragraph_format.keep_with_next = True
    LAST["spacer"].paragraph_format.keep_with_next = True


def _emit(fig, name, title, width):
    """図を保存し、表題と図を同じページに置いて挿入する。"""
    from PIL import Image
    path = _fin(fig, name)
    with Image.open(path) as im:
        cm = width * im.height / im.width
    _bind_table(cm)
    _cap(title)
    _put(path, width)


def _cap(title):
    """図表の表題。次の段落（図）と必ず同じページに置く。"""
    FIGN[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    r = p.add_run("【図表%d】　%s" % (FIGN[0], title))
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.name = BODY
    r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)


def _put(path, width=16.6):
    LAST["table"] = None
    doc.add_picture(path, width=Cm(width))
    pp = doc.paragraphs[-1]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_after = Pt(8)
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.keep_together = True


def BAND(name, title, pairs, width=16.6, ncol=4, height=0.95, ramp=False):
    """単数回答の帯グラフ（100％積み上げ横棒）。"""
    labs = [k for k, _ in pairs]
    vals = [float(v) for _, v in pairs]
    tot = sum(vals) or 1.0
    cols = _ramp(len(labs)) if ramp else \
        [GRAYS[i % len(GRAYS)] for i in range(len(labs))]
    hats = [""] * len(labs) if ramp else \
        [HATCH[i % len(HATCH)] for i in range(len(labs))]
    dark = set(range(len(labs))[len(labs) // 2 + 1:]) if ramp else DARK
    fig, ax = plt.subplots(figsize=(6.5, height))
    left = 0.0
    for i, (lab, v) in enumerate(zip(labs, vals)):
        pct = v / tot * 100
        ax.barh([0], [pct], left=left, height=0.55, label=lab,
                color=cols[i], hatch=hats[i],
                edgecolor="black", linewidth=0.7)
        if pct >= 6:
            ax.text(left + pct / 2, 0, "%.1f%%" % pct, ha="center",
                    va="center", fontsize=7.5,
                    color="white" if i in dark else "black")
        left += pct
    ax.set_xlim(0, 100)
    ax.set_ylim(-0.5, 0.5)
    ax.set_yticks([])
    ax.set_xticks([0, 20, 40, 60, 80, 100])
    ax.set_xticklabels(["0%", "20%", "40%", "60%", "80%", "100%"], fontsize=7.5)
    for sp in ["top", "right", "left"]:
        ax.spines[sp].set_visible(False)
    ax.legend(fontsize=7.5, ncol=ncol, frameon=False, loc="upper center",
              bbox_to_anchor=(0.5, -0.35), handlelength=1.4,
              handletextpad=0.5, columnspacing=1.2)
    _emit(fig, name, title, width)


def HBAR(name, title, pairs, denom=None, unit="件", width=16.6,
         xlabel=None, height=None):
    """複数回答・度数分布の横棒グラフ。上位から並べる。"""
    labs = [k for k, _ in pairs]
    vals = [float(v) for _, v in pairs]
    n = len(labs)
    h = height or max(1.5, 0.32 * n + 0.6)
    fig, ax = plt.subplots(figsize=(6.5, h))
    ys = list(range(n))[::-1]
    ax.barh(ys, vals, height=0.62, color=GRAYS[1], edgecolor="black",
            linewidth=0.7)
    mx = max(vals) if vals else 1
    for y, v in zip(ys, vals):
        t = "%d%s" % (v, unit)
        if denom:
            t += "（%.1f%%）" % (v / denom * 100)
        ax.text(v + mx * 0.015, y, t, va="center", fontsize=7.5)
    ax.set_yticks(ys)
    ax.set_yticklabels(labs, fontsize=8)
    ax.set_xlim(0, mx * 1.28)
    ax.grid(axis="x", color="#BFBFBF", linewidth=0.5)
    ax.set_axisbelow(True)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=8)
    ax.tick_params(labelsize=7.5)
    _emit(fig, name, title, width)


def MBAND(name, title, rows, cats, width=16.6, ncol=4, ramp=False):
    """複数系列の帯グラフ。rows は [(行名, [値, ...]), ...]。"""
    cols = _ramp(len(cats)) if ramp else \
        [GRAYS[i % len(GRAYS)] for i in range(len(cats))]
    hats = [""] * len(cats) if ramp else \
        [HATCH[i % len(HATCH)] for i in range(len(cats))]
    dark = set(range(len(cats))[len(cats) // 2 + 1:]) if ramp else DARK
    fig, ax = plt.subplots(figsize=(6.5, max(1.3, 0.52 * len(rows) + 0.8)))
    ys = list(range(len(rows)))[::-1]
    for y, (_lab, vals) in zip(ys, rows):
        tot = sum(vals) or 1.0
        left = 0.0
        for i, v in enumerate(vals):
            pct = v / tot * 100
            ax.barh([y], [pct], left=left, height=0.58,
                    color=cols[i], hatch=hats[i],
                    edgecolor="black", linewidth=0.7,
                    label=cats[i] if y == ys[0] else None)
            if pct >= 7:
                ax.text(left + pct / 2, y, "%.1f" % pct, ha="center",
                        va="center", fontsize=7,
                        color="white" if i in dark else "black")
            left += pct
    ax.set_xlim(0, 100)
    ax.set_yticks(ys)
    ax.set_yticklabels([r[0] for r in rows], fontsize=8)
    ax.set_xticks([0, 20, 40, 60, 80, 100])
    ax.set_xticklabels(["0%", "20%", "40%", "60%", "80%", "100%"], fontsize=7.5)
    for sp in ["top", "right", "left"]:
        ax.spines[sp].set_visible(False)
    ax.legend(fontsize=7.5, ncol=ncol, frameon=False, loc="upper center",
              bbox_to_anchor=(0.5, -0.12 - 0.06 * len(rows)),
              handlelength=1.4, handletextpad=0.5, columnspacing=1.2)
    _emit(fig, name, title, width)


def GBAR(name, title, cats, series, width=16.6, unit="", ncol=3,
         ylabel=None):
    """系列比較の縦棒グラフ。series は [(系列名, [値, ...]), ...]。"""
    import numpy as np
    fig, ax = plt.subplots(figsize=(6.5, 2.8))
    x = np.arange(len(cats))
    w = 0.8 / len(series)
    for i, (lab, vals) in enumerate(series):
        ax.bar(x + i * w - 0.4 + w / 2, vals, width=w * 0.9, label=lab,
               color=GRAYS[i % len(GRAYS)], hatch=HATCH[i % len(HATCH)],
               edgecolor="black", linewidth=0.7)
        for xi, v in zip(x + i * w - 0.4 + w / 2, vals):
            ax.annotate("%g%s" % (v, unit), (xi, v), textcoords="offset points",
                        xytext=(0, 3), ha="center", fontsize=7)
    ax.set_xticks(x)
    ax.set_xticklabels(cats, fontsize=8)
    ax.grid(axis="y", color="#BFBFBF", linewidth=0.5)
    ax.set_axisbelow(True)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=8)
    ax.tick_params(labelsize=7.5)
    ax.legend(fontsize=7.5, ncol=ncol, frameon=False, loc="upper center",
              bbox_to_anchor=(0.5, -0.14), handlelength=1.4,
              handletextpad=0.5, columnspacing=1.2)
    _emit(fig, name, title, width)


def dist_rows(d, labels, total=None):
    tot = total or sum(v for k, v in d.items() if k != "-")
    rows = []
    for code, lab in labels.items():
        v = d.get(code, 0)
        rows.append([lab, "%d" % v, "%.1f%%" % (v / tot * 100)])
    mu = d.get("-", 0)
    if mu:
        rows.append(["無回答", "%d" % mu, "―"])
    rows.append(["合計（n=%d）" % tot, "%d" % tot, "100.0%"])
    return rows, tot


# ================================================================== 表紙
P("令和8年度", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
P("第10期介護保険事業計画", size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
  space_after=6)
P("実施済み調査　結果報告書", size=26, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
P("在宅生活改善調査／居所変更実態調査／介護人材実態調査／"
  "健康とくらしの調査", size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
  space_after=40)
P("令和8年8月", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
P("大雪地区広域連合", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
  space_after=6)
P("（東川町・美瑛町・東神楽町）", size=12,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

doc.add_page_break()

# ================================================================== 第1章
H1("第１章　調査の概要")

H2("1. 調査の目的")
P("本報告書は、第10期介護保険事業計画（令和9年度～令和11年度）の策定にあたり、"
  "大雪地区広域連合が実施した４つの調査の結果を集計・分析したものです。"
  "在宅生活の継続を困難にしている要因、施設・居住系への入退所の実態、"
  "介護人材の確保・定着の状況、及び高齢者の健康と暮らしの状況を把握し、"
  "施策の必要性、サービス見込量の感度検討及びKPIの設定の"
  "参考資料とすることを目的とします。")
NOTE("本報告書の集計値は、サービス見込量の算定式に直接入力するものではありません"
     "（第１章第２節（５）、第８章第２節）。")

H2("2. 調査の設計")
SUB("（１） 調査対象")
P("・在宅生活改善調査：居宅介護支援事業所、小規模多機能型居宅介護事業所及び"
  "地域包括支援センターと、その利用者（要介護者本人・家族）")
P("・居所変更実態調査：施設・居住系サービスを提供する事業所")
P("・介護人材実態調査：介護サービス事業所とその職員")
P("・健康とくらしの調査：65歳以上の一般高齢者及び総合事業対象者"
  "（要支援・要介護認定者を含みません）")
P("")
SUB("（２） 調査時期")
P("・在宅生活改善調査、居所変更実態調査、介護人材実態調査：令和7年4月1日現在")
P("・健康とくらしの調査：令和7年11月17日～12月8日")
P("")
SUB("（３） 調査方法")
P("いずれも大雪地区広域連合が実施したもので、受託者は貸与を受けた資料に基づき"
  "集計・分析を行っています。"
  "調査票の企画、設計、配布、回収及び入力には関与していません。")
P("")
SUB("（４） 配布・回収状況")
TBL(["調査・票種", "受領件数", "母数（公表データ等）", "把握率"],
    [["在宅生活改善調査　事業所票", "15件",
      "居宅介護支援13・小規模多機能5・地域包括2", "―"],
     ["在宅生活改善調査　利用者票", "99票", "―", "―"],
     ["居所変更実態調査　施設等票", "18件",
      "区域内31施設（指定を受けるもの19、受けないもの12）", "58.1%"],
     ["介護人材実態調査　事業所票（施設・通所系）", "24件", "―", "―"],
     ["介護人材実態調査　事業所票（訪問系）", "3件", "訪問介護13事業所",
      "23.1%"],
     ["介護人材実態調査　職員個票", "317人",
      "施設・通所系24事業所の介護職員319人", "99.4%"],
     ["介護人材実態調査　職員票（訪問系）", "26件",
      "訪問系3事業所の介護職員42人", "61.9%"],
     ["健康とくらしの調査", "4,798票", "65歳以上7,121人", "67.4%"]],
    [6.4, 2.0, 5.4, 2.0], num_from=1)
NOTE("在宅生活改善調査、居所変更実態調査及び介護人材実態調査は"
     "配布名簿・配布数の記録がないため、回収率（回収数÷配布数）は"
     "算定できません。"
     "表の「把握率」は、母数を公表データで置き換えられるものに限り、"
     "受領件数÷公表データによる母数として算定したものであり、"
     "回収率とは異なります。")
NOTE("居所変更実態調査の58.1%は、区域内31施設のうち"
     "18施設から回答を得たという施設数ベースの割合"
     "（回答施設割合）です。"
     "定員ベースの把握率は第３章の冒頭に示します。")
NOTE("健康とくらしの調査は回収4,798票（回収率67.4%）で、"
     "このうち分析対象は4,729票です。"
     "本報告書第６章の集計はこの4,729票によります。")

H2("3. 集計・分析の方法")
SUB("（１） 集計方法")
P("単純集計及び各調査内のクロス集計を行いました。"
  "４つの調査を横断したクロス集計は、"
  "在宅生活改善調査の利用者票の所在地区の記入形式が統一されておらず、"
  "個票を地区に割り付けられないため行っていません（第８章第３節）。"
  "これに代えて、公表データによる供給構造の分析を行い、"
  "集計値の水準で調査相互を接続しています（第７章）。")
P("")
SUB("（２） 集計結果の表記")
NOTE("１．nは質問に対する回答者数（無回答者を含まない）であり、"
     "集計対象総数として割合算出の基準となります。")
NOTE("２．Nは回答いただいた事業所が雇用している介護職員の人数であり、"
     "割合算出の基準となります。"
     "なお、母体を同一とする事業所が同一の職員を重複して計上している例が"
     "１件あり、取扱いを確定していません（第８章第３節）。")
NOTE("３．割合は、nまたはNに対する各回答数の百分率（%）で算出しています。"
     "小数点以下第２位を四捨五入し、小数点以下第１位までを表記しているため、"
     "合計が100.0%とならない場合があります。")
NOTE("４．複数回答可の設問では、各選択肢の割合の合計が100.0%を"
     "超える場合があります。"
     "在宅生活改善調査の複数回答の設問は、"
     "選択肢ごとの無回答を判別できないため、"
     "回答票数99票を分母としています。"
     "単数回答の設問は無回答を除いた有効回答数（n）を分母とするため、"
     "同じ調査でも設問により分母が異なります。")
NOTE("５．グラフや表の選択肢（カテゴリー）は、"
     "文字数の制約により簡略化して表記している場合があります。")
NOTE("６．図表の表題に付す分母は、"
     "単数回答を「n=○」、複数回答を「回答票n=○・複数回答」、"
     "事業所・施設単位の集計を「N=○施設」又は「N=○事業所」、"
     "介護職員を単位とする集計を「N=○人」、"
     "時間の集計を「n=○職員票・合計○分」と表記します。")
NOTE("７．表記は次の基準によります。"
     "本文中の１桁の数は全角数字、２桁以上の数、割合及び年月日は半角数字、"
     "表・グラフ内の数値は半角数字とします。"
     "元号は「令和7年度」のように半角数字で表記します。"
     "サービスの名称は本文・表では正式名称を用い、"
     "グラフの軸ラベルに限り略記します。")
P("")
SUB("（３） グラフ・集計表の見方")
P("・単数回答（帯グラフ）")
P("単数回答形式の設問は、原則として帯グラフで表示します。"
  "構成比が6%に満たない区分は、グラフ内の数値表記を省略しています。")
P("・複数回答（横棒グラフ）")
P("複数回答形式の設問は、回答の多い順に横棒グラフで表示します。"
  "件数と、回答者数に対する割合を併記します。")
P("・図表はいずれも白黒印刷を前提とし、"
  "濃淡とハッチング（網掛け）で区分を表しています。")
P("")
SUB("（４） 用語とサービスの範囲")
P("本報告書では、サービスの範囲を次のとおり用います。")
TBL(["用語", "含まれるサービス", "介護保険の指定"],
    [["施設サービス",
      "介護老人福祉施設（特別養護老人ホーム）、"
      "地域密着型介護老人福祉施設、介護老人保健施設", "受ける"],
     ["居住系サービス",
      "認知症対応型共同生活介護（グループホーム）、"
      "特定施設入居者生活介護（介護付有料老人ホーム）", "受ける"],
     ["指定を受けない住まい",
      "住宅型有料老人ホーム、サービス付き高齢者向け住宅、"
      "軽費老人ホーム（ケアハウス）", "受けない"],
     ["施設・住まい", "上記3区分の全体（区域内31施設）", "―"],
     ["在宅サービス",
      "訪問系、通所系、短期入所、小規模多機能型居宅介護等。"
      "指定を受けない住まいの入居者もこれを利用する", "受ける"]],
    [4.0, 9.6, 3.4], num_from=99)
NOTE("「サービス付き高齢者向け住宅」は本文・表で略記しません。"
     "グラフの軸ラベルに限り、文字数の制約により"
     "「サ高住」と表記する場合があります。")
P("")
SUB("（５） 調査結果とサービス見込量の関係")
P("サービス見込量は、要介護度別の認定者数に給付実績から求めた利用率と"
  "受給者1人当たりの利用日数・回数を乗じて算定します"
  "（別冊「将来推計 第2段階　サービス見込量」）。"
  "本報告書の調査結果は、この算定の直接の入力値ではありません。")
P("調査結果は、①施策の必要性の根拠、②見込量の上振れ・下振れを"
  "検討する際の材料、③代表KPIの基準値（第８章第２節（２））の3つに用います。"
  "調査で得られた件数（たとえば「より適切と思われるサービス」の件数）を"
  "そのまま需要人数や必要整備量に読み替えることはしません。"
  "読み替えるには、重複、回答事業所の構成、利用意向、"
  "現在の給付実績、町別の供給状況を併せて補正する必要があり、"
  "本報告書ではその補正を行っていません。")
NOTE("上記の補正と、低位・標準・高位の需要シナリオの設定は、"
     "点検事項の取扱いが決まった後に、"
     "将来推計の資料として別途作成します（第8章第3節）。")
P("")
SUB("（６） 割合の表記")
NOTE("割合は小数第2位を四捨五入し、小数第1位まで表記しています。"
     "端数処理により、内訳の合計が100.0%とならない場合があります。"
     "また、設問により無回答の数が異なるため、"
     "分母（n）は設問ごとに異なります。"
     "各図表には「n=○」を付しています。")
P("")
SUB("（７） 本報告書を読む際の留意点")
P("４つの調査は対象者が異なります。"
  "在宅生活改善調査は事業所が課題があると判断した利用者、"
  "居所変更実態調査は施設入所者、介護人材実態調査は事業所の職員、"
  "健康とくらしの調査は一般高齢者（要支援・要介護認定者を含まない）が"
  "対象です。"
  "割合を相互に比較したり、人数を足し合わせたりすることはできません。"
  "本報告書では、各数値に母集団を明記しています。")

# ================================================================== 第2章
doc.add_page_break()
H1("第２章　在宅生活改善調査")
NOTE("本章では、居宅介護支援事業所等が在宅生活の継続に課題があると判断した"
     "利用者99人について、事業所の方にご回答いただいた内容の"
     "集計・分析を行います。"
     "本調査は課題のある利用者を抽出して回答する設計であるため、"
     "割合は区域内の在宅利用者全体を表すものではありません。")

H2("問１　利用者の状況について")
SUB("（１） 現在の居所")
rows, tot = dist_rows(S.RIYO["現在の居所"],
                      {"1": "自宅等", "2": "住宅型有料老人ホーム",
                       "3": "サービス付き高齢者向け住宅",
                       "4": "軽費老人ホーム"})
TBL(["区分", "件数", "割合"], rows, [8.0, 3.0, 3.0])
BAND("f_kyosho", "現在の居所（n=98）",
     [("自宅等", 81), ("住宅型有料老人ホーム", 11), ("軽費老人ホーム", 6)],
     ncol=3)
NOTE("該当のない区分（サービス付き高齢者向け住宅）は"
     "グラフから省略しています。")
P("回答のあった98人のうち、自宅等で生活している方が81人（82.7%）と"
  "大半を占めます。"
  "住宅型有料老人ホームが11人（11.2%）、軽費老人ホームが６人（6.1%）で、"
  "介護保険の指定を受けない住まいに居住している方が17人（17.3%）です。"
  "サービス付き高齢者向け住宅に該当する回答はありません。")
P("")
SUB("（２） 要支援・要介護度")
rows, tot = dist_rows(S.RIYO["要介護度"],
                      {"1": "要支援1", "2": "要支援2", "3": "要介護1",
                       "4": "要介護2", "5": "要介護3", "6": "要介護4",
                       "7": "要介護5", "8": "新規申請中"})
TBL(["区分", "件数", "割合"], rows, [8.0, 3.0, 3.0])
HBAR("f_yokaigo", "利用者の要支援・要介護度（n=98）",
     [("要支援1", 3), ("要支援2", 13), ("要介護1", 33), ("要介護2", 22),
      ("要介護3", 19), ("要介護4", 5), ("要介護5", 3)], denom=98, unit="人")
P("要介護1が33人（33.7%）と最も多く、要介護2が22人（22.4%）、"
  "要介護3が19人（19.4%）と続きます。"
  "要介護4・5は８人（8.2%）です。"
  "本調査で抽出された課題事例では、中軽度の方が多数を占めています。")
NOTE("本調査は事業所が在宅生活の継続に課題があると判断した利用者を"
     "抽出したものであり、区域内の在宅利用者の要介護度の分布を"
     "示すものではありません。"
     "課題事例に中軽度の方が多いことと、"
     "区域全体で中軽度の方が在宅生活を続けにくいこととは別の事柄です。")
P("")
SUB("（３） 世帯類型")
rows, tot = dist_rows(S.RIYO["世帯類型"],
                      {"1": "独居", "2": "夫婦のみ", "3": "単身の子との同居",
                       "4": "その他の同居"})
TBL(["区分", "件数", "割合"], rows, [8.0, 3.0, 3.0])
BAND("f_setai", "世帯類型（n=98）",
     [("独居", 45), ("夫婦のみ", 17), ("単身の子との同居", 14),
      ("その他の同居", 22)])
P("独居が45人（45.9%）と最も多く、夫婦のみが17人（17.3%）です。"
  "独居と夫婦のみを合わせると62人（63.3%）で、"
  "３分の２を占めます。")
NOTE("同居する介護の担い手の有無や年齢は、本設問からは分かりません。"
     "主な介護者の年代は別に集計しており、"
     "60代34人・50代19人・80歳以上15人・70代9人です"
     "（家族等介護者はいない12人を除く）。")

H2("問２　在宅生活の維持の見通しについて")
rows, tot = dist_rows(S.RIYO["生活の維持"],
                      {"1": "現在の状態では在宅生活の維持が困難",
                       "2": "当面は在宅生活を維持できる"})
TBL(["区分", "件数", "割合"], rows, [10.0, 3.0, 3.0])
MBAND("f_iji", "在宅生活の維持の見通しと就労継続（n=98）",
      [("在宅生活の維持が困難", [72, 26]),
       ("介護者の就労継続が困難", [41, 57])],
      ["該当する", "該当しない"], ncol=2)
P("調査対象99人のうち本設問の有効回答は98人で、"
  "現在の状態では在宅生活の維持が困難とされた方は"
  "72人（有効回答の73.5%）です。"
  "本調査は事業所が課題があると判断した利用者を抽出する設計であるため、"
  "この割合を区域内の在宅利用者全体の割合として読むことはできません。"
  "計画本文でも「調査対象99人のうち本設問の有効回答98人において」と"
  "母集団と分母を明記します。")
P("")
P("主な介護者の就労継続が困難になっているとされた方は41人（41.8%）です。"
  "介護離職の防止は、"
  "在宅生活の継続と就労の両立を支える施策の課題となります。")

H2("問３　在宅生活の継続を困難にしている要因について")
SUB("（１） 本人の状態等")
rows = [[k, "%d" % v, "%.1f%%" % (v / 99 * 100)]
        for k, v in sorted(S.RIYO["本人の状態等"].items(),
                           key=lambda x: -x[1])]
TBL(["要因", "件数", "回答者99人に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_honnin", "在宅生活の継続を困難にしている要因（本人の状態等・回答票n=99・複数回答）",
     sorted(S.RIYO["本人の状態等"].items(), key=lambda x: -x[1]), denom=99)
P("認知症の症状の悪化が37件（37.4%）と最も多く、"
  "必要な生活支援の発生・増大が30件（30.3%）、"
  "必要な身体介護の増大が29件（29.3%）と続きます。"
  "医療的ケア・医療処置の必要性の高まりは15件（15.2%）です。")
P("")
SUB("（２） 家族等介護者の状況")
rows = [[k, "%d" % v, "%.1f%%" % (v / 99 * 100)]
        for k, v in sorted(S.RIYO["家族等介護者"].items(),
                           key=lambda x: -x[1])]
TBL(["要因", "件数", "回答者99人に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_kazoku", "在宅生活の継続を困難にしている要因（家族等介護者・回答票n=99・複数回答）",
     sorted(S.RIYO["家族等介護者"].items(), key=lambda x: -x[1]), denom=99)
P("介護者の介護に係る不安・負担量の増大が50件（50.5%）と最も多く、"
  "回答者の半数を超えます。"
  "家族等の介護等技術では対応が困難が21件（21.2%）、"
  "本人と家族等の関係性に課題があるが16件（16.2%）です。"
  "在宅生活の継続を困難にしているのは本人の状態だけではなく、"
  "支える側の状況が大きく関わっていることがわかります。")
P("")
SUB("（３） 主な介護者の負担となっている介護")
rows = [[k, "%d" % v, "%.1f%%" % (v / 99 * 100)]
        for k, v in sorted(S.RIYO["介護者の負担"].items(),
                           key=lambda x: -x[1])[:10]]
TBL(["介護の内容", "件数", "回答者99人に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_futan", "主な介護者の負担となっている介護（上位10・回答票n=99・複数回答）",
     sorted(S.RIYO["介護者の負担"].items(), key=lambda x: -x[1])[:10],
     denom=99)
P("外出の付き添い、送迎等が30件（30.3%）と最も多く、"
  "認知症状への対応が28件（28.3%）、"
  "日中の排泄と夜間の排泄がそれぞれ20件（20.2%）と続きます。"
  "移動の支援と認知症への対応が、家族の負担の中心にあります。")

H2("問４　必要な生活支援について")
rows = [[k, "%d" % v, "%.1f%%" % (v / 99 * 100)]
        for k, v in sorted(S.RIYO["必要な生活支援"].items(),
                           key=lambda x: -x[1])]
TBL(["生活支援", "件数", "回答者99人に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_seikatsu", "必要な生活支援（回答票n=99・複数回答）",
     sorted(S.RIYO["必要な生活支援"].items(), key=lambda x: -x[1]), denom=99)
P("外出同行（通院、買い物など）が51件（51.5%）と最も多く、"
  "回答者の半数を超えます。"
  "移送サービス19件（19.2%）を含めると、"
  "移動に関する支援が上位を占めます。"
  "次いで見守り・声かけとサロンなどの定期的な通いの場が"
  "それぞれ34件（34.3%）です。")
P("外出同行、見守り・声かけ、通いの場は、"
  "いずれも介護保険の給付では担いにくい支援です。"
  "生活支援体制整備事業及び構成３町の施策で対応する領域として、"
  "第10期計画の基本目標に位置づけます。")

H2("問５　より適切と思われるサービスについて")
SVC = S.RIYO["より適切なサービス"]
rows = [[k, "%d" % v, "%.1f%%" % (v / 99 * 100)]
        for k, v in sorted(SVC.items(), key=lambda x: -x[1]) if v > 0]
TBL(["サービス", "件数", "回答者99人に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_svc", "より適切と思われるサービス（回答票n=99・複数回答）",
     [(k, v) for k, v in sorted(SVC.items(), key=lambda x: -x[1]) if v > 0],
     denom=99)
P("小規模多機能型居宅介護が44件（44.4%）と最も多く、"
  "住宅型有料老人ホーム27件（27.3%）、"
  "認知症対応型共同生活介護（グループホーム）と特別養護老人ホームが"
  "それぞれ20件（20.2%）と続きます。")
NOTE("本設問については、回答者である事業所が自らの提供するサービスを"
     "挙げる構造があります。"
     "小規模多機能型居宅介護を選択した44件のうち31件（70.5%）は、"
     "同一法人の小規模多機能２事業所からの提出票によるものです。"
     "件数をそのまま需要の強さとして読むことはできません（第８章第３節）。")
P("")
P("なお、夜間対応型訪問介護（７件）、定期巡回・随時対応型訪問介護看護（９件）、"
  "看護小規模多機能型居宅介護（10件）の３サービスは、"
  "延べ26件が挙げられていますが、"
  "いずれも区域内に事業所が存在せず、受給率も0.0%です。"
  "24時間対応サービスの確保方策として、"
  "第10期計画の整備方針で検討します。")

H2("問６　施設等への入所の緊急度について")
rows, tot = dist_rows(S.RIYO["緊急度"],
                      {"1": "3か月以内", "2": "1年以内", "3": "1年より先"},
                      total=60)
TBL(["区分", "件数", "割合"], rows, [8.0, 3.0, 3.0])
BAND("f_kinkyu", "施設等への入所の緊急度（n=60）",
     [("3か月以内", 17), ("1年以内", 39), ("1年より先", 4)], ncol=3)
P("より適切と思われるサービスとして施設等を選んだ60人のうち、"
  "３か月以内の入所が必要とされた方は17人（28.3%）です。"
  "１年以内を含めると56人（93.3%）となります。")
P("")
SUB("入所できていない理由")
rows = [["空きがない", "30", "―"], ["費用負担", "8", "―"],
        ["本人が望まない", "3", "―"], ["家族が望まない", "2", "―"],
        ["その他", "10", "―"]]
TBL(["理由", "件数", "備考"], rows, [8.0, 2.6, 4.0])
HBAR("f_riyu", "入所できていない理由（n=53）",
     [("空きがない", 30), ("その他", 10), ("費用負担", 8),
      ("本人が望まない", 3), ("家族が望まない", 2)], denom=53)
P("空きがないが30件と最も多くなっています。"
  "特定施設入居者生活介護は区域内３施設の定員156人に対し入居者154人"
  "（98.7%）でほぼ満室であり（第３章）、"
  "この結果と整合します。")

H2("問７　事業所票の集計")
P("在宅生活改善調査の事業所票15件では、"
  "所属する介護支援専門員34人、"
  "自宅等の利用者550人、"
  "サービス付き高齢者向け住宅等の利用者170人が把握されました。"
  "過去１年間に居場所の変更があった方は171人です。")
Z3 = collections.Counter()
for *_x, q3 in S.ZAI:
    for k, v in q3.items():
        Z3[k] += v or 0
tz3 = sum(Z3.values())
rows = [[k, "%d" % v, "%.1f%%" % (v / tz3 * 100)] for k, v in Z3.most_common()]
rows.append(["合計", "%d" % tz3, "100.0%"])
TBL(["変更先", "人数", "割合"], rows, [8.0, 3.0, 3.0])
HBAR("f_henkou", "過去1年間の居場所の変更先（n=171）",
     Z3.most_common(), denom=tz3, unit="人")
NOTE("事業所票15件のうち1件（居宅介護支援事業所おうか）は、"
     "事業所名の記入はあるものの各設問が0又は記号のみとなっています。"
     "介護サービス情報公表システムの個別公表画面"
     "（記入日 令和8年2月11日）により、"
     "同事業所の事業開始年月日及び指定年月日がいずれも令和7年10月1日で"
     "あることを確認しました。"
     "調査の基準日である令和7年4月1日時点では開設していないため、"
     "0は記入漏れではなく正しい状態です。"
     "本報告書では「回答事業所15件」と"
     "「集計に反映した事業所14件」を区別しています。")
P("")
NOTE("区域内の事業所数は、北海道の介護保険事業所一覧"
     "（令和8年6月30日現在）により、"
     "居宅介護支援13事業所、小規模多機能型居宅介護5事業所、"
     "地域包括支援センター3か所です。"
     "回答は居宅介護支援10事業所、小規模多機能型居宅介護3事業所、"
     "地域包括支援センター2か所です。"
     "居宅介護支援の未回答3事業所は、"
     "ひがしかわ介護相談センター、ほの香居宅介護支援事業所及び"
     "ライフデザイン陽風です。"
     "ただし、これらの事業所の指定年月日を確認していないため、"
     "基準日に開設していたかどうかは確定していません。")

# ================================================================== 第3章
doc.add_page_break()
H1("第３章　居所変更実態調査")
# 回答した18施設の定員（種別ごと）
_GRP = {9: "特養", 10: "地密特養", 7: "老健", 4: "GH", 5: "特定施設",
        1: "住宅型有料", 3: "サ高住"}
ANSCAP = collections.Counter()
for s_ in S.SHI:
    ANSCAP[_GRP[s_["種別"]]] += s_["定員"] or 0
KUCAP = {"特養": CAP_TOKUYO, "地密特養": CAP_CHITOKU, "老健": CAP_ROKEN,
         "GH": CAP_GH, "特定施設": CAP_TOKUTEI, "住宅型有料": CAP_JUTAKU,
         "サ高住": CAP_SAKO, "軽費": CAP_KEIHI}
SHITEI_KEYS = ["特養", "地密特養", "老健", "GH", "特定施設"]
CAP_SHITEI_KU = sum(KUCAP[k] for k in SHITEI_KEYS)          # 717
CAP_SHITEI_ANS = sum(ANSCAP[k] for k in SHITEI_KEYS)        # 531

NOTE("本章では、施設・居住系サービスを提供する事業所18件にご回答いただいた"
     "内容の集計・分析を行います。"
     "区域内の施設・住まいは31施設（介護保険の指定を受けるもの19、"
     "受けないもの12）で、回答は18施設、未回答は13施設です。"
     "施設数ベースの把握率は58.1%、"
     "介護保険の指定を受けるものに限れば73.7%（19施設中14施設）です。"
     "本章の数値は回答した18施設の記述値であり、区域全体の値ではありません。"
     "定員は北海道が公表する名簿によります（第７章第１節）。")
TBL(["区分", "区域内\n施設数", "回答\n施設数", "施設数\n把握率",
     "区域内\n定員", "回答施設\nの定員", "定員\n把握率"],
    [[nm, "%d" % ku, "%d" % ans, "%.1f%%" % (ans / ku * 100),
      "%d" % KUCAP[key] if key else "―",
      "%d" % ANSCAP[key] if key else "―",
      "%.1f%%" % (ANSCAP[key] / KUCAP[key] * 100) if key else "―"]
     for nm, ku, ans, key in [
         ("特別養護老人ホーム", 3, 2, "特養"),
         ("地域密着型介護老人福祉施設", 3, 2, "地密特養"),
         ("介護老人保健施設", 3, 3, "老健"),
         ("認知症対応型共同生活介護", 7, 5, "GH"),
         ("特定施設入居者生活介護", 3, 2, "特定施設")]]
    + [["　指定を受けるもの　小計", "19", "14", "73.7%",
        "%d" % CAP_SHITEI_KU, "%d" % CAP_SHITEI_ANS,
        "%.1f%%" % (CAP_SHITEI_ANS / CAP_SHITEI_KU * 100)]]
    + [[nm, "%d" % ku, "%d" % ans, "%.1f%%" % (ans / ku * 100),
        "%d" % KUCAP[key], "%d" % ANSCAP[key],
        "%.1f%%" % (ANSCAP[key] / KUCAP[key] * 100)]
       for nm, ku, ans, key in [
           ("住宅型有料老人ホーム", 9, 3, "住宅型有料"),
           ("サービス付き高齢者向け住宅", 2, 1, "サ高住")]]
    + [["軽費老人ホーム", "1", "0", "0.0%", "%d" % CAP_KEIHI, "0", "0.0%"],
       ["　指定を受けないもの　小計", "12", "4", "33.3%",
        "%d（参考）" % (CAP_JUTAKU + CAP_SAKO + CAP_KEIHI),
        "%d（参考）" % (ANSCAP["住宅型有料"] + ANSCAP["サ高住"]), "―"],
       ["合計", "31", "18", "58.1%", "―", "―", "―"]],
    [4.2, 1.6, 1.6, 1.8, 1.8, 2.0, 1.8])
NOTE("定員の単位は人です。ただしサービス付き高齢者向け住宅は登録戸数（戸）で"
     "あるため、指定を受けないものの小計及び合計の定員は算定していません。")
NOTE("未回答13施設のうち8施設は介護保険の指定を受けない住まいです。"
     "指定を受けるものでは、"
     "東神楽町特別養護老人ホームアゼリアハイツ（広域型・ユニット型の２件）、"
     "グループホームファミリー、グループホームくるみの郷、"
     "さわやか東神楽館（特定施設・区域内最大）が未回答です。")
NOTE("介護保険の指定を受ける施設・居住系の定員ベースの把握率は"
     "%.1f%%（区域内の既知定員%d人に対し回答施設の定員%d人）です。"
     "ただし、グループホーム１事業所（くるみの郷）の定員が"
     "公表資料で確認できないため、区域内定員%d人はこの１事業所を含まず、"
     "定員ベースの把握率は暫定値です（第８章第３節）。"
     % (CAP_SHITEI_ANS / CAP_SHITEI_KU * 100, CAP_SHITEI_KU,
        CAP_SHITEI_ANS, CAP_SHITEI_KU))

H2("問１　施設等の規模について")
CAPD = {}
for s_ in S.SHI:
    g = {1: "住宅型有料・サ高住", 3: "住宅型有料・サ高住", 4: "グループホーム",
         5: "特定施設", 7: "介護老人保健施設", 9: "特養・地域密着型特養",
         10: "特養・地域密着型特養"}.get(s_["種別"], "その他")
    d = CAPD.setdefault(g, {"n": 0, "cap": 0, "res": 0})
    d["n"] += 1
    d["cap"] += s_["定員"] or 0
    d["res"] += s_["入所"] or 0
FULLG = {
    "特養・地域密着型特養":
        "介護老人福祉施設・地域密着型介護老人福祉施設",
    "介護老人保健施設": "介護老人保健施設",
    "グループホーム": "認知症対応型共同生活介護",
    "特定施設": "特定施設入居者生活介護",
    "住宅型有料・サ高住":
        "住宅型有料老人ホーム・サービス付き高齢者向け住宅",
}
rows = []
for g in ["特養・地域密着型特養", "介護老人保健施設", "グループホーム",
          "特定施設", "住宅型有料・サ高住"]:
    d = CAPD[g]
    res = d["res"] + (58 if g == "特養・地域密着型特養" else 0)
    rows.append([FULLG[g], "%d" % d["n"], "%d" % d["cap"], "%d" % res,
                 "%.1f%%" % (res / d["cap"] * 100)])
tc = sum(d["cap"] for d in CAPD.values())
tr = sum(d["res"] for d in CAPD.values()) + 58
rows.append(["合計", "%d" % sum(d["n"] for d in CAPD.values()),
             "%d" % tc, "%d" % tr, "%.1f%%" % (tr / tc * 100)])
TBL(["区分", "施設数", "定員", "入所者", "入所率"], rows,
    [7.4, 1.8, 2.0, 2.2, 2.2])
HBAR("f_nyusho", "種別ごとの入所率（N=18施設）",
     [(g, round(
         (CAPD[g]["res"] + (58 if g == "特養・地域密着型特養" else 0))
         / CAPD[g]["cap"] * 100, 1))
      for g in ["特養・地域密着型特養", "介護老人保健施設", "グループホーム",
                "特定施設", "住宅型有料・サ高住"]], unit="%",
     xlabel="入所率（％）")
NOTE("特別養護老人ホーム美瑛慈光園は入所者数が未記入のため、"
     "要介護度別の内訳58人を加えています。"
     "特定施設は区域内３施設のうち２施設の分です。")
P("回答のあった18施設の定員は652人、入所者は594人で、"
  "入所率は91.1%です。"
  "グループホームは92.6%、特養・地域密着型特養は96.1%と高く、"
  "介護老人保健施設は83.8%です。")

H2("問２　入所者の要介護度について")
rows = []
for g, d in C.CS["種別×入所者の要介護度"].items():
    n = sum(d.values())
    hv = d.get("要介護4", 0) + d.get("要介護5", 0)
    rows.append([FULLG.get(g, g), "%d" % n, "%d" % hv,
                 "%.1f%%" % (hv / n * 100) if n else "―"])
TBL(["種別", "入所者", "うち要介護4・5", "重度者の割合"], rows,
    [7.4, 2.2, 3.0, 3.0])
NOTE("入所者の要介護度別の合計は、問１の入所者数と一致しない種別があります。"
     "介護老人保健施設は要介護度別の合計206人に対し入所者数201人、"
     "特養・地域密着型特養は、特別養護老人ホーム美瑛慈光園が"
     "入所者数を記入していないため、"
     "問１の入所者数146人は要介護度別の内訳58人を加えたものです。"
     "調査票の記入の不一致であり、確認のうえ確定します（第８章第３節）。")
_KC = ["要支援1", "要支援2", "要介護1", "要介護2", "要介護3", "要介護4",
       "要介護5"]
MBAND("f_shubetsu_kaigo", "種別ごとの入所者の要介護度（N=18施設・重いほど濃い）",
      [(g, [C.CS["種別×入所者の要介護度"][g].get(k, 0) for k in _KC])
       for g in ["特養・地域密着型特養", "介護老人保健施設", "グループホーム",
                 "特定施設", "住宅型有料・サ高住"]],
      _KC, ncol=4, ramp=True)
P("特養・地域密着型特養は要介護4・5が73人（50.0%）と重度者の割合が高く、"
  "介護老人保健施設は83人（40.3%）です。"
  "一方、グループホームは22人（29.3%）、特定施設は18人（31.6%）と"
  "重度者の割合が低くなっています。"
  "住宅型有料・サービス付き高齢者向け住宅は47人（40.9%）で、"
  "介護保険の指定を受けない住まいにも重度の方が居住しています。")
P("")
SUB("特定施設入居者生活介護の入居者の要介護度")
P("特定施設は区域内3施設のうち2施設が回答しています。"
  "回答した2施設（57人）の要介護度は次のとおりです。", size=None)
_sur = C.CS["種別×入所者の要介護度"].get("特定施設", {})
_sn = sum(_sur.values())
TBL(["要介護度", "人数", "割合"],
    [[k, "%d" % _sur.get(k, 0), "%.1f%%" % (_sur.get(k, 0) / _sn * 100)]
     for k in ["要支援1", "要支援2", "要介護1", "要介護2", "要介護3",
               "要介護4", "要介護5"]] + [["合計（n=%d）" % _sn, "%d" % _sn,
                                          "100.0%"]],
    [8.0, 3.0, 3.0])
P("")
P("未回答のさわやか東神楽館（定員100人）については、"
  "介護サービス情報公表システムの個別公表画面（記入日 令和7年10月6日）に"
  "97人分の要介護度別の内訳が掲載されています。"
  "調査（令和7年4月1日現在）とは6か月以上の時点差があるため、"
  "次表は参考集計として示します。", size=None)
_all = {k: _sur.get(k, 0) + SW_KD.get(k, 0)
        for k in ["要支援1", "要支援2", "要介護1", "要介護2", "要介護3",
                  "要介護4", "要介護5"]}
_n = sum(_all.values())
rows = [[k, "%d" % v, "%.1f%%" % (v / _n * 100)] for k, v in _all.items()]
rows.append(["合計", "%d" % _n, "100.0%"])
TBL(["要介護度", "人数", "割合"], rows, [8.0, 3.0, 3.0])
HBAR("f_tokutei_kaigo", "特定施設入居者生活介護の入居者の要介護度（n=154人・参考集計）",
     list(_all.items()), denom=_n, unit="人")
NOTE("【参考集計】居所変更実態調査に回答した２施設（57人・令和7年4月1日現在）と、"
     "介護サービス情報公表システムの個別公表画面による１施設"
     "（さわやか東神楽館・97人・令和7年10月6日現在）を"
     "合わせたものです。"
     "時点が異なり、1施設が全体の63.0%を占めるため、"
     "区域全体の要介護度構成を示すものではありません。"
     "計画判断には同一時点のデータへの更新を要します。")
P("参考集計では、要介護1が56人（36.4%）と最も多く、"
  "要支援1・2の30人（19.5%）を合わせると、"
  "要支援1から要介護1までが86人（55.8%）と半数を超えます。"
  "要介護4・5は27人（17.5%）です。"
  "回答した2施設のみでみても、要支援1〜要介護1が18人（31.6%）、"
  "要介護4・5が18人（31.6%）です。"
  "軽度者の割合が高いことは両者に共通しますが、"
  "その程度は資料により異なります。"
  "特定施設が重度者の受け皿として機能しているかどうかの判断には、"
  "同一時点の3施設のデータを要します。")

H2("問３　新規入所者の入所前の居場所について")
IN_ = collections.Counter()
for g, d in C.CS["種別×入所前の居場所"].items():
    for k, v in d.items():
        IN_[k] += v
ti = sum(IN_.values())


def _split(cnt):
    """区域内・区域外を合算した区分に集約する。"""
    agg = collections.Counter()
    inn = collections.Counter()
    out = collections.Counter()
    for k, v in cnt.items():
        base = k.replace("（区域外）", "")
        agg[base] += v
        (out if "（区域外）" in k else inn)[base] += v
    return agg, inn, out


IAG, IIN, IOUT = _split(IN_)
rows = [[k, "%d" % v, "%.1f%%" % (v / ti * 100),
         "%d" % IIN[k], "%d" % IOUT[k]]
        for k, v in IAG.most_common()]
rows.append(["合計", "%d" % ti, "100.0%", "%d" % sum(IIN.values()),
             "%d" % sum(IOUT.values())])
TBL(["入所前の居場所", "人数", "割合", "うち区域内", "うち区域外"], rows,
    [6.0, 2.4, 2.4, 2.4, 2.4])
HBAR("f_nyushomae", "新規入所者の入所前の居場所（n=335人・N=18施設）",
     IAG.most_common(), denom=ti, unit="人")
P("新規入所者335人のうち、病院・診療所からの入所が191人（57.0%）と"
  "最も多く、自宅からが92人（27.5%）です。"
  "回答した18施設では、医療機関を経由する入所が半数を超えており、"
  "医療と介護の連携の重要性が示唆されます。")
NOTE("本調査は入所前の居場所を記述したものであり、"
     "連携の有無や質が入所の流れを左右したかどうかは測定していません。"
     "退院調整の実態、受入れを断った理由、待機期間の把握には"
     "別の調査又は聞き取りを要します。")
P("")
SUB("区域内・区域外の別")
rows = []
_IO = []
for g, d in C.CS["種別×入所前の居場所"].items():
    out = sum(v for k, v in d.items() if "（区域外）" in k)
    tot = sum(d.values())
    inn = tot - out
    if tot == 0:
        continue
    rows.append([g, "%d" % inn, "%d" % out, "%d" % tot,
                 "%.1f%%" % (out / tot * 100)])
    _IO.append((g, [inn, out]))
rows.append(["合計", "%d" % sum(int(r[1]) for r in rows),
             "%d" % sum(int(r[2]) for r in rows), "%d" % ti,
             "%.1f%%" % (sum(int(r[2]) for r in rows) / ti * 100)])
TBL(["種別", "区域内", "区域外", "計", "区域外の割合"], rows,
    [5.0, 2.2, 2.2, 2.2, 3.0])
MBAND("f_kuiki", "新規入所者の区域内・区域外の別（N=18施設）", _IO, ["区域内", "区域外"],
      ncol=2)
P("グループホームは84.6%、特定施設は84.2%が区域外からの入所です。"
  "一方、特養・地域密着型特養は17.8%、介護老人保健施設は32.7%にとどまります。"
  "全体では335人のうち137人（40.9%）が区域外からの入所です。"
  "これらの施設には住所地特例が適用され、"
  "保険者は従前の市町村のままとなります。"
  "区域内の定員が区域内の被保険者のために使われているとは限りません。")
NOTE("調査票の選択肢は「市内／市外」のみであり、"
     "「市内」が町単位か広域連合単位かの定義が示されていません。"
     "介護サービス自給率との接続には確認を要します。")

H2("問４　退去者の退去先と退去理由について")
rows = [["病院・診療所", "164", "47.0%"], ["死亡", "95", "27.2%"],
        ["自宅", "43", "12.3%"], ["その他", "47", "13.5%"],
        ["合計", "349", "100.0%"]]
TBL(["退去先", "人数", "割合"], rows, [8.0, 3.0, 3.0])
HBAR("f_taikyosaki", "退去者の退去先（n=349人・N=18施設）",
     [("病院・診療所", 164), ("死亡", 95), ("自宅", 43), ("その他", 47)],
     denom=349, unit="人")
rr = collections.Counter()
for s_ in S.SHI:
    for k, v in (s_.get("退去理由") or {}).items():
        rr[k] += v or 0
rows = [[k, "%d" % v] for k, v in rr.most_common()]
TBL(["退去理由", "件数"], rows, [9.0, 3.0])
HBAR("f_taikyoriyu", "退去理由（N=18施設・複数回答）", rr.most_common(),
     denom=sum(rr.values()))
P("退去先は病院・診療所が164人（47.0%）と最も多く、"
  "死亡が95人（27.2%）です。"
  "自宅へ戻った方は43人（12.3%）にとどまります。"
  "退去理由では医療的ケアが16件と最も多くなっています。"
  "医療的な対応の必要性と医療機関への退去が"
  "いずれも上位にあることは関連が示唆されますが、"
  "施設別のクロス集計ができていないため、"
  "同一の事例であるかどうかは確かめられていません。")

H2("問５　待機者について")
P("特別養護老人ホームの待機者は、延べ137人と回答されています。"
  "同一法人の２施設で同数が計上されており、"
  "これを重複として除くと82人となります。"
  "個人単位の照合ができていないため、"
  "本報告書では延べ137人・最小82人の範囲で示します。"
  "施設に入所中で特別養護老人ホームを待機している方は16人です。")
NOTE("待機者数は整備の必要量の判断に直結します。"
     "実人数を確定するには、匿名化した申込者の照合キー、申込日、"
     "現在の居所により名寄せを行い、"
     "施設に入所中の待機者と在宅の待機者を分ける必要があります。"
     "申込の継続意思が確認されていない者が含まれている可能性もあります"
     "（第８章第３節）。")

H2("問６　受け入れ可能な医療処置について")
mp = collections.Counter()
for s_ in S.SHI:
    for k, v in (s_.get("医療処置") or {}).items():
        mp[k] += 1 if v else 0
rows = [[k, "%d" % v, "%.1f%%" % (v / 18 * 100)]
        for k, v in mp.most_common() if k != "対応可能な医療処置はない"]
TBL(["医療処置", "施設数", "18施設に対する割合"], rows, [8.6, 2.6, 3.4])
HBAR("f_iryo", "受け入れ可能な医療処置（N=18施設・複数回答）",
     [(k, v) for k, v in mp.most_common() if k != "対応可能な医療処置はない"],
     denom=18, unit="施設")
P("褥瘡の処置が16施設（88.9%）と最も多く、"
  "酸素療法とストーマの処置がそれぞれ13施設（72.2%）と続きます。"
  "一方、モニター測定は１施設のみで、"
  "中心静脈栄養、透析、レスピレーター、疼痛の看護は"
  "それぞれ４施設にとどまります。"
  "対応可能な医療処置がないと回答した施設が２施設あります。")
P("医療依存度の高い方の受け入れ先は限られており、"
  "退去先の47.0%が病院・診療所であることと合わせて、"
  "医療と介護の連携の課題を示しています。")

# ================================================================== 第4章
doc.add_page_break()
H1("第４章　介護人材実態調査　事業所票")
NOTE("本章では、事業所の管理者の方にご回答いただいた内容の"
     "集計・分析を行います。"
     "訪問系は区域内13事業所のうち３事業所からの回答であり、"
     "訪問系を含む数値は区域全体を代表するものではありません。")

H2("問１　提供するサービス種別について")
FULLCAT = {
    "特養": "介護老人福祉施設（特別養護老人ホーム）",
    "地密特養": "地域密着型介護老人福祉施設",
    "老健": "介護老人保健施設",
    "通所リハ": "通所リハビリテーション",
    "通所介護等": "通所介護・地域密着型通所介護等",
    "GH": "認知症対応型共同生活介護（グループホーム）",
    "特定施設": "特定施設入居者生活介護",
    "住宅型有料": "住宅型有料老人ホーム",
    "サ高住": "サービス付き高齢者向け住宅",
    "訪問介護": "訪問介護",
}
rows = []
for cat in ["特養", "地密特養", "老健", "通所リハ", "通所介護等", "GH",
            "特定施設", "住宅型有料", "サ高住", "訪問介護"]:
    n = sum(1 for c, j in zip(CAT, S.JIN) if c == cat)
    rows.append([FULLCAT[cat], "%d" % n, "%d" % BYCAT[cat]])
rows.append(["合計", "%d" % len(S.JIN), "%d" % sum(BYCAT.values())])
TBL(["サービス区分", "事業所数", "介護職員数"], rows, [9.0, 2.4, 2.4])
NOTE("グラフの軸ラベルは、文字数の制約により略記しています。")
HBAR("f_cat_shokuin",
     "サービス区分別の介護職員数（【ケース②】重複を除く案 N=%d人）" % N_ALL,
     [(c, BYCAT[c]) for c in
      ["GH", "老健", "特養", "訪問介護", "通所リハ", "特定施設", "地密特養",
       "通所介護等", "住宅型有料", "サ高住"]], denom=N_ALL, unit="人")
NOTE("サービス区分は、事業所票の種別と事業所名により受託者が付したものです。"
     "本図は重複を除く案（ケース②）によるものです。"
     "重複を除かないケース①では、住宅型有料老人ホームが%d人、"
     "合計が%d人となり、他の区分は変わりません。"
     "重複の取扱いは確定していません（第８章第３節）。"
     % (BYCAT["住宅型有料"] + DUP, N_SIS + N_HOU))

H2("問２－１　介護職員の人数について")
SUB("（１） 全体の人数と常勤・非常勤の割合")
rows = [["施設・通所系（24事業所）", "%d" % N_SIS,
         "%d" % sum(j["常勤"] or 0 for j in SIS),
         "%d" % sum(j["非常勤"] or 0 for j in SIS)],
        ["訪問系（3事業所）", "%d" % N_HOU,
         "%d" % sum(j["常勤"] or 0 for j in HOU),
         "%d" % sum(j["非常勤"] or 0 for j in HOU)],
        ["単純合計", "%d" % (N_SIS + N_HOU), "―", "―"],
        ["重複を除いた場合", "%d" % N_ALL, "―", "―"],
        ["確定していない範囲", "%d〜%d" % (N_ALL, N_SIS + N_HOU), "―", "―"]]
TBL(["区分", "介護職員", "常勤", "非常勤"], rows, [6.0, 3.0, 2.4, 2.4])
NOTE("施設・通所系は、介護職員319人に対し常勤258人・非常勤55人（計313人）で"
     "６人の差があります。"
     "グループホームびえいの郷が介護職員16人・常勤10人と記入し、"
     "非常勤の欄が未記入であることによるものです。"
     "同事業所の公表画面（記入日 令和7年10月4日）では"
     "介護職員は常勤10人・非常勤７人であり、常勤の人数は一致します。"
     "未記入の非常勤は６人程度とみられます。"
     "常勤・非常勤の割合は、記入のあった313人を分母としています。")
MBAND("f_kinmu_kubun",
      "介護職員の常勤・非常勤の別（N=%d人・記入のあった職員）"
      % (sum((j["常勤"] or 0) + (j["非常勤"] or 0) for j in S.JIN)),
      [("施設・通所系", [sum(j["常勤"] or 0 for j in SIS),
                        sum(j["非常勤"] or 0 for j in SIS)]),
       ("訪問系", [sum(j["常勤"] or 0 for j in HOU),
                  sum(j["非常勤"] or 0 for j in HOU)])],
      ["常勤", "非常勤"], ncol=2)
P("施設・通所系24事業所の介護職員は319人、"
  "訪問系３事業所は42人で、単純合計は361人です。"
  "ただし、同一法人が施設・通所系と訪問系の両方の事業所票に"
  "同一の職員13人を記入しているため、"
  "単純に合計すると13人が二重に計上されている可能性があります。"
  "同一人物であるかどうかは職員個票等による照合ができておらず、"
  "確定していません。"
  "本報告書では介護職員数を348〜361人の範囲で示し、"
  "受託者の案（348人）を用いる場合はその旨を明記します。")
NOTE("重複の取扱いは確定していません。"
     "住宅型有料老人ホームは介護保険の指定サービスではなく、"
     "地域包括ケア「見える化」システムにも現れないことから、"
     "受託者は訪問系の13人を実数とし有料老人ホーム側を重複として除いた"
     "348人を用いることを案としています（第８章第３節）。"
     "同じ構造の住宅型有料老人ホームびえいの郷は"
     "「有料１人・訪問介護13人」と分けて記入しており、"
     "公表データでも住宅型有料老人ホーム及び訪問介護事業所（大町2丁目5番28号）と"
     "グループホーム（同5番14号）が別の事業所であることが確認できます。")
NOTE("介護サービス情報公表システムの個別公表画面"
     "（ヘルパーステーションフラワー・記入日 令和7年10月29日）により、"
     "訪問介護員等の実人数13人・常勤8人・非常勤5人、"
     "前年度の採用者数7人・退職者数2人が確認できました。"
     "本調査の訪問系の回答と実人数・常勤・非常勤・採用・離職の"
     "すべてが一致します。"
     "訪問介護事業所の側の13人は実在の値です。"
     "ただし公表制度の対象は介護保険の指定サービスであり、"
     "住宅型有料老人ホームは対象外であるため、"
     "有料老人ホーム側が0であるかどうかは公表データでは確認できません。")
P("")
SUB("（１の２） 公表画面との突合")
P("訪問系の回答３事業所について、"
  "介護サービス情報公表システムの個別公表画面と突合しました。")
TBL(["事業所", "調査（令和7年4月1日）", "公表画面（記入日）", "差"],
    [["花時計訪問介護事業所", "16人（常勤13・非常勤3）",
      "16人（常勤13・非常勤3）令和7年10月16日", "±0"],
     ["ヘルパーステーションフラワー", "13人（常勤8・非常勤5）",
      "13人（常勤8・非常勤5）令和7年10月29日", "±0"],
     ["シルバーハウス訪問介護事業所", "13人（常勤9・非常勤4）",
      "17人（常勤11・非常勤6）令和7年11月6日", "＋4"]],
    [4.6, 4.6, 5.4, 1.8], num_from=99)
NOTE("３事業所のうち２事業所は実人数・常勤・非常勤のすべてが一致します。"
     "本調査の事業所票が、"
     "公表制度の「訪問介護員等の実人数」と同じ定義で記入されていることを"
     "示しています。"
     "シルバーハウス訪問介護事業所の差４人は、"
     "調査の基準日から公表画面の記入日までの約７か月間の増員とみられますが、"
     "確かめたものではありません。"
     "採用者数・離職者数は、調査が「過去1年間」、"
     "公表画面が「前年度」であり、対象期間の解釈が異なる可能性があります。")
P("")
P("本調査に回答のなかった訪問介護事業所についても、"
  "公表画面が入手できたものは訪問介護員等の人数を把握しました。")
TBL(["事業所", "町", "訪問介護員等の実人数", "常勤換算",
     "前年度の採用・退職"],
    [["指定（介護予防）訪問介護ケンセイシャレバレッジ", "東神楽町",
      "10人（常勤6・非常勤4）", "4.0人", "採用1人・退職0人"],
     ["東神楽町ホームヘルプサービスセンター", "東神楽町",
      "6人（常勤4・非常勤2）", "6.0人", "採用0人・退職2人"],
     ["指定訪問介護事業所ひばり", "東神楽町",
      "12人（常勤11・非常勤1）", "6.3人", "採用0人・退職1人"],
     ["東川町社協訪問介護事業所", "東川町",
      "7人（常勤3・非常勤4）", "4.2人", "採用0人・退職0人"],
     ["指定訪問介護（指定介護予防訪問介護）事業所 縁結び", "東川町",
      "31人（常勤11・非常勤20）", "8.28人", "採用6人・退職10人"],
     ["訪問介護事業所 桜華", "東川町",
      "22人（常勤10・非常勤12）", "5.0人", "採用0人・退職0人"],
     ["指定訪問介護事業所 恩送り", "東川町",
      "11人（常勤9・非常勤2）", "記載なし", "採用0人・退職0人"],
     ["美瑛町ホームヘルプサービスセンター", "美瑛町",
      "13人（常勤6・非常勤7）", "6.4人", "採用2人・退職2人"]],
    [5.0, 1.4, 4.4, 2.0, 4.4], num_from=99)
NOTE("これにより、区域内の訪問介護13事業所のうち"
     "11事業所（回答３・未回答８）について"
     "訪問介護員等の実人数を把握しました。合計158人です。"
     "残るのは訪問介護事業所 ほがらか（東神楽町）及び"
     "指定訪問介護ステーション ゆう（東川町）の２事業所です。"
     "８事業所とも本調査の基準日（令和7年4月1日）より前からの指定であり、"
     "基準日に開設していなかったことによる未回答ではありません。")
NOTE("実人数に対する常勤換算の比は事業所により大きく異なります"
     "（東神楽町ホームヘルプサービスセンターは6人に対し6.0人、"
     "縁結びは31人に対し8.28人）。"
     "兼務の多寡によるものです。"
     "訪問系の従事者数を実人数で積み上げる場合はその旨を明記して用います。"
     "なお、指定訪問介護事業所 恩送りは常勤換算人数の欄が0であり、"
     "利用者17人・提供時間20時間との整合がとれないため用いていません。"
     "美瑛町ホームヘルプサービスセンターは常勤の週の勤務すべき時間数が"
     "38.75時間で、他の10事業所（40時間）と異なります。")
NOTE("指定訪問介護事業所ひばりは、"
     "同一敷地に住宅型有料老人ホームひばりの森"
     "（本調査に回答・介護職員9人）及びグループホームひばり"
     "（本調査に回答・介護職員21人）があり、"
     "訪問介護員等12人との重複の有無は確かめられていません。"
     "有料老人ホーム華とヘルパーステーションフラワーと同じ構造ですが、"
     "施設票の9人（常勤7・非常勤2・採用3・離職0）と"
     "公表画面の12人（常勤11・非常勤1・採用0・退職1）は数値が異なり、"
     "同一の値が転記されたものではありません。")
P("")
P("公表画面が揃った11事業所について、"
  "併設する居住系の住まいの有無を整理しました。")
TBL(["区分", "事業所数", "訪問介護員等", "利用者", "事業所"],
    [["併設あり", "７", "122人（77.2%）", "182人（65.7%）",
      "花時計、フラワー、シルバーハウス、ひばり、縁結び、桜華、恩送り"],
     ["併設なし", "４", "36人（22.8%）", "95人（34.3%）",
      "３町の社会福祉協議会の３事業所及びケンセイシャレバレッジ"],
     ["計", "11", "158人", "277人", "―"]],
    [2.0, 1.4, 3.4, 3.2, 7.2], num_from=99)
NOTE("併設の判断は、公表画面の建物名の欄又は特色の欄の記載、"
     "北海道の届出済有料老人ホーム一覧及び住所地特例適用施設一覧との"
     "所在地の一致によります。"
     "把握できた訪問介護員等158人のうち122人（77.2%）が、"
     "住宅型有料老人ホーム又はサービス付き高齢者向け住宅に"
     "併設された事業所に所属しています。"
     "利用者277人のうち182人（65.7%）も同様です。"
     "区域内の訪問介護は、居宅で暮らす方への訪問よりも、"
     "居住系の住まいの入居者への訪問が主となっています。"
     "第４章の在宅サービスの見込量は"
     "地域包括ケア「見える化」システムの受給者数によるため"
     "算定そのものへの影響はありませんが、"
     "本文で訪問介護を「在宅」として記述する際は"
     "この併設の割合を併せて示します。")
P("")
SUB("（２） 外国人職員の状況")
P("外国人職員は%d人で、%d事業所に配置されています。"
  "配置されているのはいずれも施設・居住系及び通所系の事業所で、"
  "訪問系の３事業所には配置されていません。"
  "施設・通所系の介護職員%d人に対する割合は%.1f%%です。"
  % (GAIKOKU, N_GAI_JIG, N_SIS, GAIKOKU / N_SIS * 100))
NOTE("外国人職員はすべて施設・通所系の事業所に配置されているため、"
     "割合の分母は施設・通所系の介護職員%d人としています。"
     "訪問系を含む348人を分母とすると%.1f%%となります。"
     % (N_SIS, GAIKOKU / N_ALL * 100))

H2("問２－３　過去１年間の採用者数と離職者数について")
N_RAW = N_SIS + N_HOU
SAIYO_RAW, RISHOKU_RAW = SAIYO + 7, RISHOKU + 2
rows = [["介護職員数（割合の分母）", "%d人" % N_RAW, "%d人" % N_ALL],
        ["採用者数", "%d人（%.1f%%）" % (SAIYO_RAW, SAIYO_RAW / N_RAW * 100),
         "%d人（%.1f%%）" % (SAIYO, SAIYO / N_ALL * 100)],
        ["離職者数",
         "%d人（%.1f%%）" % (RISHOKU_RAW, RISHOKU_RAW / N_RAW * 100),
         "%d人（%.1f%%）" % (RISHOKU, RISHOKU / N_ALL * 100)],
        ["差（採用－離職）", "%+d人" % (SAIYO_RAW - RISHOKU_RAW),
         "%+d人" % (SAIYO - RISHOKU)]]
TBL(["区分", "【ケース①】重複を除かない", "【ケース②】重複を除く案"], rows,
    [5.2, 5.2, 5.2])
GBAR("f_saiyo", "令和6年度の採用者数と離職者数（2ケースの対照）",
     ["採用者数", "離職者数", "差（採用－離職）"],
     [("【ケース①】重複を除かない（N=%d人）" % N_RAW,
       [SAIYO_RAW, RISHOKU_RAW, SAIYO_RAW - RISHOKU_RAW]),
      ("【ケース②】重複を除く案（N=%d人）" % N_ALL,
       [SAIYO, RISHOKU, SAIYO - RISHOKU])], unit="人", ncol=2, ylabel="人")
P("令和6年度の採用者数は83〜90人、離職者数は65〜67人で、"
  "差は＋18人から＋23人です。"
  "採用が離職を上回っていますが、"
  "介護職員数に対する離職者の割合はケース①18.6%・ケース②18.7%であり、"
  "毎年２割近くが入れ替わっている計算になります。")
NOTE("割合は、それぞれのケースの介護職員数を分母としています。"
     "ケースをまたいで最小値と最大値を組み合わせた割合"
     "（たとえば83人÷361人）は、分子と分母の定義が異なるため用いません。")
NOTE("単純合計では採用90人・離職67人です。"
     "有料老人ホーム華とヘルパーステーションフラワーが"
     "同一の採用７人・離職２人を計上している可能性があり、"
     "これを重複として除くと採用83人・離職65人となります。"
     "同一であるかどうかは確定していないため範囲で示しています。"
     "内訳は、施設・通所系が採用76人・離職60人、"
     "訪問系が採用14人・離職７人です。")
NOTE("北海道の評価指標及び代表KPIが用いる「採用率と離職率の差」は、"
     "各年9月30日現在の在籍者数を分母とします。"
     "本調査は令和7年4月1日現在であり、分母の時点が異なるため、"
     "上記の割合はこれとは別のものです。")


# ================================================================== 第5章
doc.add_page_break()
H1("第５章　介護人材実態調査　職員票")
NOTE("本章では、施設・通所系24事業所の介護職員317人にご回答いただいた"
     "職員個票の集計・分析を行います。"
     "個票そのものは個人情報を含むため本報告書には収録せず、"
     "集計値のみを示します。")

H2("問１　資格の取得・研修の修了状況について")
rows, tot = dist_rows(SK["資格"],
                      {"1": "介護福祉士（認定介護福祉士を含む）",
                       "2": "介護福祉士実務者研修修了又は同等",
                       "3": "介護職員初任者研修修了又は同等",
                       "4": "資格・研修なし"})
TBL(["区分", "人数", "割合"], rows, [9.0, 2.6, 2.6])
BAND("f_shikaku", "資格の取得・研修の修了状況（n=316）",
     [("介護福祉士", 212), ("実務者研修修了", 11), ("初任者研修修了", 38),
      ("資格・研修なし", 55)])
P("介護福祉士が212人（67.1%）を占めます。"
  "実務者研修修了11人（3.5%）、初任者研修修了38人（12.0%）を合わせると、"
  "何らかの資格・研修を有する方が261人（82.6%）となります。"
  "資格・研修なしは55人（17.4%）です。")
P("介護福祉士の割合は、"
  "北海道の評価指標及び第10期介護保険事業計画の基本指針案が求める"
  "介護職員の質に関する指標に対応します。")

H2("問２　雇用形態、性別、年齢、勤務年数について")
SUB("（１） 雇用形態")
rows, tot = dist_rows(SK["雇用形態"], {"1": "常勤職員", "2": "非常勤職員"})
TBL(["区分", "人数", "割合"], rows, [9.0, 2.6, 2.6])
BAND("f_koyou", "雇用形態（n=316）", [("常勤職員", 253), ("非常勤職員", 63)],
     ncol=2)
P("常勤職員が253人（80.1%）を占めます。")
P("")
SUB("（２） 性別")
rows, tot = dist_rows(SK["性別"], {"1": "男性", "2": "女性"})
TBL(["区分", "人数", "割合"], rows, [9.0, 2.6, 2.6])
BAND("f_seibetsu", "性別（n=316）", [("男性", 134), ("女性", 182)], ncol=2)
P("女性が182人（57.6%）、男性が134人（42.4%）です。")
P("")
SUB("（３） 年齢")
rows, tot = dist_rows(SK["年齢"],
                      {"1": "20歳未満", "2": "20代", "3": "30代",
                       "4": "40代", "5": "50代", "6": "60代",
                       "7": "70代以上"})
TBL(["区分", "人数", "割合"], rows, [9.0, 2.6, 2.6])
HBAR("f_nenrei", "年齢（n=316）",
     [("20歳未満", 11), ("20代", 66), ("30代", 59), ("40代", 77),
      ("50代", 61), ("60代", 36), ("70代以上", 6)], denom=316, unit="人")
P("40代が77人（24.4%）と最も多く、20代66人（20.9%）、"
  "50代61人（19.3%）、30代59人（18.7%）と続きます。"
  "60代以上は42人（13.3%）です。")
NOTE("本調査では勤務継続の意向及び定年の定めを把握していないため、"
     "年齢構成から退職の時期を推定することはできません。"
     "将来の需給を検討するには、勤務継続意向と定年後の再雇用の状況を"
     "別途把握する必要があります。")
P("")
SUB("（４） 現在の事業所での勤務年数")
rows, tot = dist_rows(SK["勤務年数"], {"1": "1年以上", "2": "1年未満"})
TBL(["区分", "人数", "割合"], rows, [9.0, 2.6, 2.6])
BAND("f_kinzoku", "現在の事業所での勤務年数（n=316）",
     [("1年以上", 253), ("1年未満", 63)], ncol=2)
P("現在の事業所での勤務年数が１年未満の方は63人（19.9%）です。")
P("")
SUB("（５） 雇用形態と勤務年数の関係")
TBL(["雇用形態", "1年以上", "1年未満", "計", "1年未満の割合"],
    [["常勤", "206", "47", "253", "18.6%"],
     ["非常勤", "47", "16", "63", "25.4%"],
     ["計", "253", "63", "316", "19.9%"]],
    [3.6, 2.4, 2.4, 2.4, 3.2])
MBAND("f_koyou_kinzoku", "雇用形態別の勤務年数（n=316）",
      [("常勤", [206, 47]), ("非常勤", [47, 16])],
      ["1年以上", "1年未満"], ncol=2)
P("勤続１年未満の割合は、常勤18.6%に対し非常勤25.4%と非常勤の方が"
  "高くなっています。"
  "一方、勤続１年未満の63人のうち47人（74.6%）は常勤職員であり、"
  "定着の課題は非常勤に限られません。")
NOTE("雇用形態と勤務年数は、いずれも「1」が253人・「2」が63人と"
     "周辺度数が一致しますが、個票のレベルでは一致しません"
     "（上表のとおり）。同じ数値が並ぶことによる転記の誤りに留意します。")
P("")
SUB("（６） 過去１週間の勤務時間")
P("回答のあった315人の平均は35.1時間です。"
  "40時間を超える方が37人、０時間の方が５人います。")

H2("問３　前職について")
rows, tot = dist_rows(
    SK["直前の職場"],
    {"1": "現在の職場が初めての勤務先", "2": "介護以外の職場",
     "3": "介護老人福祉施設・介護老人保健施設・介護医療院・短期入所・"
          "認知症対応型共同生活介護・特定施設入居者生活介護",
     "4": "訪問介護・訪問入浴介護・夜間対応型訪問介護",
     "5": "小規模多機能型居宅介護・看護小規模多機能型居宅介護・"
          "定期巡回・随時対応型訪問介護看護",
     "6": "通所介護・通所リハビリテーション・認知症対応型通所介護",
     "7": "住宅型有料老人ホーム・サービス付き高齢者向け住宅"
          "（特定施設入居者生活介護以外）",
     "8": "その他の介護サービス", "9": "不明"},
    total=64)
TBL(["区分", "人数", "割合"], rows, [10.0, 2.4, 2.4])
HBAR("f_zenshoku", "直前の職場（n=64）",
     [("現在の職場が初めての勤務先", 19),
      ("特養・老健等の施設系", 17), ("介護以外の職場", 13),
      ("住宅型有料・サ高住", 6), ("その他の介護サービス", 3),
      ("不明", 2), ("通所介護・通所リハ等", 2),
      ("小規模多機能等", 1), ("訪問介護等", 1)], denom=64, unit="人")
NOTE("本設問は勤続１年未満の方等を対象とするもので、"
     "回答は64人です。「*」（対象外）253人を除いています。")
P("現在の職場が初めての勤務先という方が19人（29.7%）と最も多く、"
  "特養・老健等の施設系からの転職が17人（26.6%）、"
  "介護以外の職場からの転職が13人（20.3%）と続きます。"
  "初めての勤務先と介護以外からの転職を合わせると32人（50.0%）となり、"
  "介護の実務経験が浅い可能性のある層が半数を占めます。"
  "本設問は介護の経験年数を直接把握していないため"
  "（初めての勤務先の方にも研修・実習の経験があり得るほか、"
  "介護以外からの転職者にも過去の介護経験があり得ます）、"
  "研修・OJTの需要を検討する際の参考値として扱います。")

H2("問４　訪問介護員のサービス提供時間について")
h = collections.Counter()
for x in S.HOU:
    for k, v in x["内訳"].items():
        h[k] += v or 0
tot_h = sum(h.values())
rows = [[k, "%d" % v, "%.1f%%" % (v / tot_h * 100)]
        for k, v in h.most_common() if v > 0]
rows.append(["合計", "%d" % tot_h, "100.0%"])
TBL(["内容", "分", "割合"], rows, [8.0, 3.0, 3.0])
BAND("f_houmon", "訪問介護のサービス提供時間の内訳（n=26職員票・合計24,643分）",
     [("身体介護", 22103), ("生活援助（掃除）", 1880),
      ("生活援助（その他）", 660)], ncol=3)
P("訪問系の職員票26件によると、"
  "サービス提供時間の合計24,643分のうち、"
  "身体介護が22,103分（89.7%）を占めます。"
  "生活援助は掃除1,880分（7.6%）、その他660分（2.7%）で、"
  "買い物と調理・配膳の記録はありません。"
  "介護予防・総合事業による訪問の記録もありません。")
NOTE("訪問系は区域内13事業所のうち３事業所からの回答です。"
     "区域全体の訪問介護の内容を表すものではありません。")

# ================================================================== 第6章
doc.add_page_break()
H1("第６章　健康とくらしの調査")
NOTE("本章では、健康とくらしの調査の分析対象4,729票の要約を示します。"
     "回収は4,798票で、所定の除外を行った4,729票が分析対象です。"
     "設問により無回答の数が異なるため、分母（n）は設問ごとに異なります。"
     "本章では各指標にnを併記します。"
     "詳細は別冊「調査クロス集計・分析」（24シート）によります。"
     "本調査は要支援・要介護認定者を含まないため、"
     "第２章から第５章の各調査とは母集団が異なります。")

H2("問１　心身の状態について")
TBL(["指標", "該当数／n", "本広域連合", "同規模保険者40", "差"],
    [["フレイル該当割合", "903／4,719", "19.1%", "―", "―"],
     ["1年間の転倒あり", "1,651／4,656", "35.5%", "30.0%", "＋5.5pt"],
     ["口腔機能低下", "1,153／4,671", "24.7%", "21.8%", "＋2.9pt"]],
    [5.0, 3.4, 2.8, 3.0, 2.2])
GBAR("f_jages", "同規模保険者40との比較（設問別のnは本文のとおり）",
     ["1年間の転倒あり", "口腔機能低下", "友人知人と会う頻度が高い"],
     [("本広域連合", [35.5, 24.7, 63.0]),
      ("同規模保険者40", [30.0, 21.8, 71.2])], unit="%", ncol=2,
     ylabel="％")
P("フレイル該当割合は903人／4,719人（19.1%）です。"
  "町別では美瑛町21.6%、東神楽町17.0%と差がありますが、"
  "美瑛町は回答者に占める75歳以上の割合が高く、"
  "年齢調整後は差が縮まります。"
  "町間の比較は年齢調整後の値によることとします。")
P("転倒と口腔機能低下は、"
  "人口５万人未満の同規模保険者40と比べ、"
  "年齢調整を行わない記述比較では高い値となっています。"
  "転倒は65～69歳で同規模を9.6ポイント上回っています。")
NOTE("比較対象40保険者のうち31（77.5%）は要支援者を調査対象に含んでおり、"
     "本広域連合は含みません。母集団の定義が異なります。"
     "要支援者を含む集団の方が指標は不良側に出やすいため、"
     "本広域連合が上回る指標（転倒・口腔機能低下）は、"
     "要支援者の包含差だけでは高い値を説明しにくい可能性があります。"
     "ただし年齢・性別等を調整していないため、"
     "統計的な優劣は確定しません。")

H2("問２　社会参加について")
P("友人知人と会う頻度が高い方の割合は2,898人／4,602人（63.0%）で、"
  "同規模保険者40の71.2%を8.2ポイント下回ります。"
  "未調整の記述比較では、外出と社会参加が最も差の大きい領域です。")
P("通いの場の参加者は、独居30.3%（不参加16.7%）、"
  "生活動作の困りごとあり54.6%（同38.7%）と、"
  "支援の必要度が高い層の割合が高くなっています。"
  "通いの場は既にハイリスク層に届いており、"
  "量的な拡大だけでなく、"
  "届いている層と届いていない層を分けた施策が必要です。")

H2("問３　生活支援のニーズについて")
P("生活動作の困りごとのうち、除雪が1,002人／4,128人（24.3%）と"
  "最も多くなっています。"
  "除雪は介護保険の給付対象外であり、"
  "生活支援体制整備事業及び構成３町の施策で対応します。")
P("困りごとがある方のうち「解決できず、困っている」と回答した方は1.8%、"
  "回答者全体では0.6%です。"
  "この値は、生活上の困りごとが解決できていない層の規模を示すものです。"
  "本設問は困りごとの解決状況を尋ねたものであり、"
  "これを在宅生活の継続が困難な層の指標として用いるには、"
  "住み替えの意向、介護者の状況、サービスの利用状況との"
  "追加の検証を要します。")
P("何らかの介護・介助が必要だが現在は受けていないと回答した方は"
  "323人／4,645人（7.0%）です。"
  "認定を受けながらサービスを利用していない方とは別の集団であり、"
  "合算することはできません。")

H2("問４　支え手と意向について")
P("日常的に支障が生じた場合に世話をしてくれる人が「いない」と"
  "回答した方は401人／4,599人（8.7%）、独居の方では32.4%です。"
  "日常的な援助者がいない可能性のある層です。"
  "身元保証、入居・入所の支援、死後事務等の対象規模を推計するには、"
  "親族関係、連絡先、成年後見制度の利用、本人の意向の"
  "追加の把握を要します。")
P("介護が必要になった場合に自宅での介護を希望する方は"
  "2,279人／4,634人（49.2%）です。"
  "一方で、24時間対応の３サービスが区域内に存在しないことと"
  "合わせて考える必要があります。")
P("認知症に関する相談窓口を「知らない」と回答した方は"
  "2,981人／4,639人（64.3%）です。"
  "本人又は家族に認知症の症状がある方に限っても47.3%が知らないと"
  "回答しています。")

H2("問５　地域資源について")
P("ソーシャルキャピタルの連帯感は、"
  "美瑛町が参加74市町村中14位、東川町が15位と上位にあります。"
  "健康指標には課題がある一方、"
  "地域への信頼と連帯感は高い水準にあります。"
  "この強みを、通いの場の担い手確保及び生活支援体制整備の資源として"
  "位置づけます。")

# ================================================================== 第7章
doc.add_page_break()
H1("第７章　供給構造からみた調査結果の解釈")
NOTE("本章では、４つの調査を横断したクロス集計に代えて、"
     "北海道が公表する名簿及び地域包括ケア「見える化」システムにより"
     "区域内の供給構造を整理し、集計値の水準で調査結果を接続します。"
     "個票を突き合わせたものではないため、"
     "関連の指摘にとどめ、因果関係を述べるものではありません。")

H2("第１節　事業所数と定員")
SUB("（１） 介護保険の指定を受ける施設・居住系")
TBL(["区分", "事業所数", "定員", "調査の把握"],
    [["介護老人福祉施設（特別養護老人ホーム）", "3", "%d人" % CAP_TOKUYO,
      "2施設・110人"],
     ["地域密着型介護老人福祉施設", "3", "%d人" % CAP_CHITOKU,
      "2施設・42人"],
     ["介護老人保健施設", "3", "%d人" % CAP_ROKEN, "3施設・240人"],
     ["認知症対応型共同生活介護", "7", "%d人＋1事業所は未確認" % CAP_GH,
      "5事業所・81人"],
     ["特定施設入居者生活介護", "3", "%d人" % CAP_TOKUTEI, "2施設・58人"]],
    [6.4, 2.2, 4.0, 3.4])
SUB("（２） 介護保険の指定を受けない住まい")
TBL(["区分", "施設数", "定員・戸数", "調査の把握"],
    [["住宅型有料老人ホーム", "9", "%d人" % CAP_JUTAKU, "3施設・91人"],
     ["サービス付き高齢者向け住宅", "2", "%d戸" % CAP_SAKO, "1件・30戸"],
     ["軽費老人ホーム（ケアハウス）", "1", "%d人" % CAP_KEIHI, "―"],
     ["養護老人ホーム", "0", "―", "―"]],
    [6.4, 2.2, 4.0, 3.4])
HBAR("f_teiin", "区域内の施設・住まいの定員（N=31施設・北海道の名簿等による）",
     [("介護老人保健施設", CAP_ROKEN), ("住宅型有料老人ホーム", CAP_JUTAKU),
      ("介護老人福祉施設", CAP_TOKUYO),
      ("特定施設入居者生活介護", CAP_TOKUTEI),
      ("認知症対応型共同生活介護", CAP_GH),
      ("サービス付き高齢者向け住宅", CAP_SAKO),
      ("地域密着型介護老人福祉施設", CAP_CHITOKU),
      ("軽費老人ホーム", CAP_KEIHI)], unit="人")
NOTE("出典：北海道保健福祉部の各名簿（令和8年6月30日～7月1日現在）。"
     "介護老人保健施設及び認知症対応型共同生活介護の一部は"
     "居所変更実態調査及び介護サービス情報公表システムの"
     "個別公表画面によります。")
P("介護保険の指定を受けない住まいは12施設で、"
  "内訳は住宅型有料老人ホーム%d人、軽費老人ホーム%d人、"
  "サービス付き高齢者向け住宅%d戸です。"
  "定員・戸数の単純合計は%d（単位が異なるため参考）で、"
  "指定を受ける居住系（グループホーム・特定施設）の255人を上回ります。"
  "居所変更実態調査で把握できたのは４施設・定員121人にとどまり、"
  "区域内の住まいの３分の１です。"
  "これらの住まいの入居者は、要介護認定を受けている場合、"
  "訪問介護・通所介護等の在宅サービスを利用します。"
  % (CAP_JUTAKU, CAP_KEIHI, CAP_SAKO,
     CAP_JUTAKU + CAP_KEIHI + CAP_SAKO))

H2("第２節　サービスの通常の事業実施地域")
rows = []
for key, label in [("訪問介護", "訪問介護"), ("訪問看護", "訪問看護"),
                   ("通所介護", "通所介護"),
                   ("地域通所", "地域密着型通所介護"),
                   ("小規模居宅", "小規模多機能型居宅介護"),
                   ("短期生活", "短期入所生活介護"),
                   ("居宅支援", "居宅介護支援")]:
    recs = H.SHITEI.get(key, [])
    cov = [sum(1 for x in recs if t in (x["実施地域"] or "")) for t in TOWNS]
    rows.append([label, "%d" % len(recs)] + ["%d" % c for c in cov])
TBL(["サービス", "事業所数", "東川町", "美瑛町", "東神楽町"], rows,
    [6.4, 2.6, 2.4, 2.4, 2.4])
GBAR("f_jisshi", "町ごとに通常の事業実施地域とする事業所の数（指定事業所の届出による）",
     ["訪問介護", "訪問看護", "地域密着型通所介護", "小規模多機能",
      "居宅介護支援"],
     [(t, [sum(1 for x in H.SHITEI.get(k, [])
                if t in (x["実施地域"] or ""))
           for k in ["訪問介護", "訪問看護", "地域通所", "小規模居宅",
                     "居宅支援"]]) for t in TOWNS], unit="", ncol=3,
     ylabel="事業所数")
NOTE("単位：当該町を通常の事業実施地域とする事業所の数。"
     "出典：北海道「介護保険事業所一覧」（令和8年6月30日現在）。"
     "訪問介護13事業所のうち３事業所は実施地域欄が空欄です。")
P("訪問介護は美瑛町を実施地域とする事業所が４と最も少なくなっています。"
  "小規模多機能型居宅介護は５事業所すべてが美瑛町に所在し、"
  "東川町・東神楽町を実施地域とするのは１事業所のみです。"
  "地域密着型通所介護は４事業所のうち３事業所が美瑛町にあり、"
  "うち２事業所は美瑛町のみを実施地域としています。")
P("同じ広域連合の区域内であっても、"
  "町により選択できる事業所の数が異なります。"
  "在宅生活改善調査で最も必要とされた外出同行（51件・51.5%）は、"
  "こうした事業所の配置と併せて考える必要があります。")

H2("第３節　運営法人の集中")
rows = []
cum = 0
for h_, n in BYH.most_common(8):
    cum += n
    rows.append([h_, "%d" % n, "%.1f%%" % (n / N_JIG * 100),
                 "%.1f%%" % (cum / N_JIG * 100)])
TBL(["運営法人", "事業所数", "割合", "累積"], rows, [7.6, 2.4, 2.4, 2.4])
NOTE("上位8法人を示しています。7位・8位は同数（3事業所）です。"
     "本文の「上位6法人が39事業所（52.0%）」は表の累積欄によります。")
HBAR("f_hojin", "運営法人別の事業所数（上位8法人・N=75事業所）",
     BYH.most_common(8), denom=N_JIG, unit="")
P("区域内の実%d事業所を%d法人が運営しています。"
  "上位６法人が%d事業所（%.1f%%）を運営する一方、"
  "13法人は１事業所のみを運営しています。"
  % (N_JIG, N_HOJIN, TOP6N, TOP6N / N_JIG * 100))
P("特に、区域内の小規模多機能型居宅介護５事業所は"
  "すべて社会福祉法人美瑛慈光会が美瑛町で運営しており、"
  "東川町・東神楽町を通常の事業実施地域とするのは１事業所のみです。")
P("在宅生活改善調査の利用者票99票のうち36票（36.4%）が"
  "同一法人の小規模多機能２事業所からの提出であり、"
  "「より適切と思われるサービス」で小規模多機能を選択した44件のうち"
  "31件（70.5%）がこの２事業所からの回答でした。"
  "この偏りは、区域内の小規模多機能型居宅介護がすべて１法人であるという"
  "供給構造の集中と、提出元ごとの回答の集中の"
  "双方を反映している可能性があります。"
  "両者の寄与を分離することはできません。")
NOTE("法人の集中は、事業者の撤退が地域のサービス供給に与える影響が"
     "大きいこと、及び事業所間の比較による質の評価が成り立ちにくいことを"
     "意味します。整備の方針を検討する際の留意事項とします。")

H2("第４節　従事者数")
rows = []
for nm, cat, code in [("介護老人福祉施設", "特養", "M2a"),
                      ("地域密着型介護老人福祉施設", "地密特養", "M2k"),
                      ("介護老人保健施設", "老健", "M2b"),
                      ("通所リハビリテーション", "通所リハ", "M2g"),
                      ("訪問介護", "訪問介護", "M2e"),
                      ("認知症対応型共同生活介護", "GH", None),
                      ("特定施設入居者生活介護", "特定施設", None),
                      ("住宅型有料老人ホーム", "住宅型有料", None),
                      ("サービス付き高齢者向け住宅", "サ高住", None)]:
    mv = None
    if code:
        v = MK.M[code]["職種"].get("介護職員") \
            or MK.M[code]["職種"].get("訪問介護員", {})
        ks = [k for k in v if v[k] is not None]
        mv = int(v[ks[-1]]) if ks else None
    rows.append([nm, "%d" % BYCAT[cat],
                 "%d" % mv if mv is not None else "系列なし",
                 "%+d" % (BYCAT[cat] - mv) if mv is not None else "―"])
TBL(["サービス", "調査の介護職員", "見える化の従事者数", "差"], rows,
    [6.4, 3.0, 3.4, 2.2])
_M = []
for nm, cat, code in [("介護老人福祉施設", "特養", "M2a"),
                      ("地域密着型特養", "地密特養", "M2k"),
                      ("介護老人保健施設", "老健", "M2b"),
                      ("通所リハ", "通所リハ", "M2g"),
                      ("訪問介護", "訪問介護", "M2e")]:
    v = MK.M[code]["職種"].get("介護職員") \
        or MK.M[code]["職種"].get("訪問介護員", {})
    ks = [k for k in v if v[k] is not None]
    _M.append((nm, BYCAT[cat], int(v[ks[-1]])))
GBAR("f_juji",
     "調査と見える化システムの介護職員数の対照"
     "（サービス別・調査は令和7年4月1日、見える化は令和6年度）",
     [m[0] for m in _M],
     [("調査（令和7年4月1日）", [m[1] for m in _M]),
      ("見える化（令和6年度）", [m[2] for m in _M])], unit="人", ncol=2,
     ylabel="人")
MBAND("f_noser", "介護職員の把握手段（2ケースの対照）",
      [("【ケース①】重複を除かない（N=%d人）" % (N_SIS + N_HOU),
        [N_SIS + N_HOU - NOSER - DUP, NOSER + DUP]),
       ("【ケース②】重複を除く案（N=%d人）" % N_ALL,
        [N_ALL - NOSER, NOSER])],
      ["見える化で把握できる", "本調査が唯一の把握手段"], ncol=2)
NOTE("見える化システムのM2系列は介護サービス施設・事業所調査"
     "（各年10月1日現在）による職種別の実人員です。"
     "本調査は令和7年4月1日現在であり、半年のずれがあります。")
P("グループホーム、特定施設、住宅型有料老人ホーム及び"
  "サービス付き高齢者向け住宅の介護職員は、"
  "見える化システムに従事者数の系列がないため、"
  "本調査が唯一の把握手段です。"
  "該当する職員は、重複を除かないケース①で%d人（%.1f%%）、"
  "重複を除く案のケース②で%d人（%.1f%%）です。"
  "これらの住まいは区域内に22施設あり、"
  "従事者の把握を継続する必要があります。"
  % (NOSER + DUP, (NOSER + DUP) / (N_SIS + N_HOU) * 100,
     NOSER, NOSER / N_ALL * 100))
NOTE("介護老人保健施設は３施設すべてが回答して77人であるのに対し、"
     "見える化システムの令和6年度の介護職員は59人です。"
     "通所リハビリテーションも同様の差があります。"
     "特別養護老人ホームの差が未回答１施設分で説明できるのに対し、"
     "この２サービスは全施設が回答したうえで差が生じています。"
     "調査での「介護職員」の範囲の解釈を確認したうえで確定します"
     "（第８章第３節）。")

H2("第５節　供給と需要の対照")
TBL(["区分", "令和7年度実績", "令和11年度見込み", "区域内定員", "定員に対する割合"],
    [["施設サービス", "336人", "339人",
      "%d人" % (CAP_TOKUYO + CAP_CHITOKU + CAP_ROKEN),
      "%.1f%%" % (339 / (CAP_TOKUYO + CAP_CHITOKU + CAP_ROKEN) * 100)],
     ["居住系サービス", "144人", "145人", "%d人" % (CAP_GH + CAP_TOKUTEI),
      "%.1f%%" % (145 / (CAP_GH + CAP_TOKUTEI) * 100)]],
    [4.0, 3.0, 3.2, 3.0, 3.4])
GBAR("f_teiin_mikomi", "見込量と区域内定員の対照（令和11年度）",
     ["施設サービス", "居住系サービス"],
     [("令和11年度の見込み", [339, 145]),
      ("区域内定員", [CAP_TOKUYO + CAP_CHITOKU + CAP_ROKEN,
                    CAP_GH + CAP_TOKUTEI])], unit="人", ncol=2, ylabel="人")
NOTE("見込みは別冊「将来推計 第２段階　サービス見込量」によります。"
     "利用率と受給者１人当たりの利用日数・回数を"
     "令和7年度の値で固定した基本ケースの値です。")
P("施設・居住系の見込量はいずれも区域内定員の範囲内に収まります。"
  "ただし、区域内の施設には住所地特例により他の保険者の被保険者も"
  "入居しており、定員と見込量の差がそのまま空きを示すものではありません。"
  "特定施設は定員156人に対し入居者154人（98.7%）でほぼ満室です。")
P("")
TBL(["区分", "令和7年度の利用率", "認定者1,984人に適用した推計人数",
     "利用率の出典・時点"],
    [["在宅サービス", "51.2%", "1,016人", "見える化D45-a系列・令和7年度"],
     ["居住系サービス", "7.3%", "145人", "見える化D45-b系列・令和7年度"],
     ["施設サービス", "17.0%", "337人", "見える化D45-c系列・令和7年度"],
     ["いずれも利用していない", "24.5%", "486人",
      "100%から上記3区分を差し引いた残差"],
     ["合計", "100.0%", "1,984人", "認定者数は令和8年3月末"]],
    [3.6, 3.0, 4.6, 4.4])
HBAR("f_riyou",
     "認定者に令和7年度の利用率を適用した推計人数（n=1,984人）",
     [("在宅サービス", 1016), ("いずれも利用していない", 486),
      ("施設サービス", 337), ("居住系サービス", 145)],
     denom=1984, unit="人")
NOTE("本表・本図は、認定者数1,984人（令和8年3月末）に"
     "令和7年度の合計利用率（見える化D45系列）を乗じて算定した推計人数であり、"
     "各区分の実際の月平均利用者数を集計したものではありません。"
     "利用率は年度の平均、認定者数は年度末の値であり、時点が異なります。"
     "月平均の受給者実績（在宅1,013人、居住系144人、施設336人）とは"
     "算定の方法が異なるため、両者は一致しません。")
P("認定を受けている1,984人に令和7年度の利用率を適用すると、"
  "在宅サービスの利用者は1,016人（51.2%）、"
  "居住系は145人（7.3%）、施設は337人（17.0%）で、"
  "いずれも利用していない方が486人（24.5%）と推計されます。"
  "この486人は、健康とくらしの調査の"
  "「介護・介助が必要だが受けていない7.0%」とは母集団が異なるため、"
  "合算することも比率を比較することもできません。"
  "未利用の理由の内訳は照会中であり（資料依頼No.8）、"
  "受領後に要因を分解します。")

# ================================================================== 第8章
doc.add_page_break()
H1("第８章　調査結果の総括")

H2("第１節　調査結果からみえたこと")
P("【在宅生活の継続】")
P("本調査で抽出された課題事例99件でみると、"
  "在宅生活の継続を困難にしているのは本人の状態の変化だけではありません。"
  "介護者の不安・負担量の増大が50件（50.5%）と最も多く、"
  "支える側の状況が大きく関わっています。"
  "必要とされる生活支援は外出同行が51件（51.5%）と最も多く、"
  "移送サービスを含めると移動に関する支援が上位を占めます。"
  "見守り・声かけ、通いの場と合わせ、"
  "介護保険の給付では担いにくい支援が求められています。")
P("")
P("【施設・居住系への入退所】")
P("回答した18施設では、新規入所の57.0%、退去先の47.0%が"
  "病院・診療所です。"
  "施設・居住系への入退所の多くが医療機関との間で生じており、"
  "医療と介護の連携の重要性が示唆されます。"
  "医療依存度の高い方を受け入れられる施設は限られており、"
  "対応可能な医療処置がないと回答した施設が２施設あります。")
P("")
P("【介護人材】")
P("介護職員は348〜361人です（重複の取扱いにより変わります）。"
  "施設・通所系の職員個票317件のうち有効回答316件でみると、"
  "介護福祉士が67.1%を占め、資格水準は高い状況にあります。"
  "一方、勤続１年未満が19.9%、令和6年度の離職者が65〜67人であり、"
  "毎年２割近くが入れ替わっています。"
  "60歳以上の職員は13.3%です。"
  "勤務継続の意向は把握していないため、"
  "将来の退職見込みは別途確認を要します。")
P("")
P("【供給構造】")
P("区域内の実75事業所を28法人が運営し、"
  "上位６法人が39事業所（52.0%）を占めます。"
  "小規模多機能型居宅介護５事業所はすべて１法人が美瑛町で運営しており、"
  "東川町・東神楽町を実施地域とするのは１事業所のみです。"
  "町により選択できるサービスの幅が異なります。")
P("24時間対応の３サービス（定期巡回・随時対応型訪問介護看護、"
  "夜間対応型訪問介護、看護小規模多機能型居宅介護）は"
  "区域内に事業所が存在せず、受給率も0.0%です。"
  "利用者票では延べ26件がこれらを必要としています。")

H2("第２節　第10期計画への反映")
SUB("（１） 調査結果とサービス見込量の関係")
P("サービス見込量は、要介護度別の認定者数に給付実績から求めた利用率と"
  "受給者1人当たりの利用日数・回数を乗じて算定します。"
  "本報告書の調査結果は、この算定の入力値ではありません。"
  "調査で得られた件数をそのまま需要人数や必要整備量に"
  "読み替えることはしません。")
TBL(["段階", "用いる資料", "本報告書の位置づけ"],
    [["①認定者数の見込み", "見える化B3系列（要介護度別認定者数）と人口推計",
      "用いない"],
     ["②利用率", "見える化D45・D46系列（給付実績）", "用いない"],
     ["③1人当たり利用日数・回数", "見える化D系列（給付実績）", "用いない"],
     ["④見込量＝①×②×③", "上記の積", "用いない"],
     ["⑤施策の必要性の根拠", "本報告書の第2〜6章",
      "用いる（件数を需要量に換算しない）"],
     ["⑥見込量の上振れ・下振れの検討", "本報告書の第2・3・7章",
      "用いる（シナリオの材料）"],
     ["⑦代表KPIの基準値", "本報告書の第2・5・6章",
      "用いる（定義は本節（２）のとおり）"],
     ["⑧供給制約の確認", "本報告書の第7章（定員・実施地域）", "用いる"]],
    [4.6, 6.4, 6.0], num_from=99)
P("")
SUB("（２） KPIの定義")
P("本報告書の集計値をKPIの基準値として用いる場合は、"
  "次のとおり定義を明示します。"
  "目標値、担当課及び更新の頻度は計画側で確定します。")
TBL(["KPIの候補", "分子／分母", "基準値", "基準年", "データ源・更新方法"],
    [["認知症の相談窓口を知っている方の割合",
      "知っていると回答した方1,658人／設問の有効回答4,639人", "35.7%",
      "令和7年度", "健康とくらしの調査（次期調査で更新）"],
     ["友人・知人と会う頻度が高い方の割合",
      "頻度が高いと回答した方2,898人／設問の有効回答4,602人", "63.0%",
      "令和7年度", "健康とくらしの調査（次期調査で更新）"],
     ["介護福祉士の割合",
      "介護福祉士212人／職員個票の有効回答316人", "67.1%",
      "令和7年度",
      "介護人材実態調査の職員個票（施設・通所系24事業所に限る）"],
     ["24時間対応の3サービスの区域内事業所数",
      "定期巡回・随時対応型訪問介護看護、夜間対応型訪問介護及び"
      "看護小規模多機能型居宅介護の指定事業所数", "0事業所",
      "令和8年度", "介護サービス情報公表システム（毎年度更新）"]],
    [3.6, 5.0, 1.8, 1.6, 3.6], num_from=99)
NOTE("介護職員の採用率・離職率は、北海道の評価指標及び代表KPIが"
     "各年9月30日現在の在籍者数を分母とするのに対し、"
     "本調査は令和7年4月1日現在であり、分母の時点が異なります。"
     "本報告書の採用・離職の割合（第４章）は参考値であり、"
     "計画のKPIは9月30日現在の別データにより算定します。")
NOTE("介護福祉士の割合の分母は施設・通所系の職員個票に限られ、"
     "訪問系及び未回答の事業所の職員を含みません。"
     "基準値と同じ範囲で更新できる場合に限りKPIとして用います。")
P("")
SUB("（３） 需要シナリオの算定に要する補正")
P("調査で得られた件数を需要量として用いるには、"
  "次の補正が必要です。本報告書ではこの補正を行っていません。"
  "点検事項の取扱いが決まった後に、"
  "別冊「将来推計 需要3シナリオの感度表」の高位シナリオへ反映します。"
  "補正が済んだ場合も、算定式の外で人数を足すのではなく、"
  "利用率又は受給者1人当たりの利用日数・回数の置き方に反映します。")
TBL(["調査結果", "必要な補正", "補正に要する資料"],
    [["より適切と思われるサービス（小規模多機能44件等）",
      "回答事業所の構成による偏りの補正、重複の除去、"
      "現に利用しているサービスとの重なりの整理、利用意向の確認",
      "提出元別の集計（実施済み）、給付実績、"
      "町別の登録枠と稼働状況"],
     ["24時間対応の3サービス（延べ26件）",
      "実人数への換算（複数回答のため）、"
      "現在利用しているサービスとの重なり、夜間の要請頻度",
      "利用者票の再集計、給付実績、事業所への聞き取り"],
     ["特別養護老人ホームの待機者（82〜137人）",
      "個人単位の名寄せ、申込の継続意思の確認、"
      "施設入所中の待機者と在宅の待機者の区分",
      "匿名化した申込者の照合キー、申込日、現在の居所"],
     ["介護職員の需給",
      "重複の確定、勤務継続意向、職種別・サービス別の必要配置",
      "職員個票の照合、追加の意向調査、人員配置基準"]],
    [4.6, 6.4, 6.0], num_from=99)
NOTE("上記の補正が済むまでは、本報告書の件数を"
     "整備量の根拠として用いないでください。")
P("")
SUB("（４） 調査結果と計画項目の対応")
TBL(["調査結果", "計画への反映", "反映先"],
    [["施設・住まいの供給量を令和8年6〜7月時点の公表名簿により更新した"
      "（未確認1事業所を除き暫定確定）", "定員を名簿の数値に改める",
      "第2章第3節"],
     ["通常の事業の実施地域が判明した",
      "訪問・通所の供給制約が疑われる地域の候補抽出に用い、"
      "給付実績、受入可否、空き、人員、移動時間で確認する",
      "第1章第7節"],
     ["新規入所の57.0%が病院・診療所から",
      "医療と介護の連携の記述の根拠とする", "第6章第5節"],
     ["外出同行が最も必要とされる生活支援",
      "移動の確保を基本目標に位置づける", "第5章 基本目標2"],
     ["24時間対応の3サービスが存在しない",
      "確保方策を整備方針に位置づける", "第6章第4節"],
     ["介護職員の33.3%が見える化で把握できない",
      "把握範囲を明記したうえで記述する", "第3章第5節"],
     ["【参考・時点混在】特定施設の入居者の要介護度構成。"
      "要支援1〜要介護2は110人／154人（71.4%）。"
      "調査に回答した2施設のみでは30人／57人（52.6%）",
      "同一時点の3施設のデータへ更新するまで、"
      "整備の必要量の算定には使用しない。"
      "更新後に重度別構成の参考資料とする", "第6章第2節・第4節"],
     ["運営法人の集中",
      "撤退の影響と質の評価の留意事項とする", "第2章第3節・第6章第4節"],
     ["除雪が困りごとの最上位",
      "生活支援体制整備事業で対応する", "第5章 基本目標2"],
     ["外出と社会参加が他団体比較の最大の課題",
      "介護予防・社会参加の重点領域とする", "第5章 基本目標1・2"],
     ["認知症の相談窓口を64.3%が知らない",
      "窓口の周知を主な事業に位置づける", "第5章 基本目標2"]],
    [6.4, 6.0, 4.4], num_from=99)

H2("第３節　本報告書の限界と取扱いを確定していない事項")
SUB("（１） 本報告書の限界")
P("① ４つの調査は対象者が異なります。"
  "割合を相互に比較したり、人数を足し合わせたりすることはできません。")
P("② 在宅生活改善調査は、事業所が課題があると判断した利用者を"
  "抽出して回答する設計です。"
  "「在宅生活の維持が困難72人（有効回答98人の73.5%）」を"
  "区域内の在宅利用者全体の割合として読むことはできません。")
P("③ 在宅生活改善調査、居所変更実態調査及び介護人材実態調査は"
  "配布数の記録がないため、回収率を算定できません。")
P("④ 居所変更実態調査は、区域内31施設のうち13施設が未回答です"
  "（施設数ベースの把握率58.1%%）。"
  "介護保険の指定を受ける施設・居住系は19施設中14施設が回答しており"
  "（施設数ベース73.7%%）、"
  "回答施設の定員は%d人、区域内の既知定員%d人に対し%.1f%%です。"
  "ただし、グループホーム１事業所の定員が未確認のため、"
  "定員ベースの把握率は暫定値です。"
  % (CAP_SHITEI_ANS, CAP_SHITEI_KU,
     CAP_SHITEI_ANS / CAP_SHITEI_KU * 100))
P("　入所前の居場所335人・退去先349人は、回答した18施設の記述値であり、"
  "未回答13施設を含みません。"
  "とくに特定施設入居者生活介護は３施設のうち２施設の回答であり、"
  "回答施設の定員は58人、区域内定員は156人です。"
  "区域内最大の特定施設が未回答であるため、"
  "種別ごとの代表性には限界があります。")
P("⑤ 介護人材実態調査の訪問系は13事業所のうち３事業所からの回答です。"
  "訪問系を含む合計値を「区域内の介護職員数」と記述することはできません。")
P("⑥ 健康とくらしの調査は要支援者・要介護者を含みません。"
  "同規模保険者40との比較には、"
  "比較対象の77.5%が要支援者を調査対象に含むという偏りがあります。")
P("⑦ 介護職員数、採用者数、離職者数及び特別養護老人ホームの待機者数は、"
  "個人単位の照合ができておらず、重複の有無が確定していません。"
  "本報告書では範囲（職員348〜361人、採用83〜90人、離職65〜67人、"
  "待機者82〜137人）で示しています。")
P("⑧ 第７章の供給構造の分析は、個票を突き合わせたものではありません。"
  "調査結果と供給構造の対応は関連の指摘であって、"
  "因果関係を示すものではありません。")
P("⑨ 単数回答の設問は無回答を除いた有効回答数（n）を分母とし、"
  "複数回答の設問は回答票数を分母としています。"
  "同じ調査でも設問により分母が異なります。")
P("⑩ 居所変更実態調査の要介護度別の内訳は、"
  "問１の入所者数と一致しない種別があります。"
  "介護老人保健施設は206人に対し201人です。")
P("⑪ 北海道の名簿は令和8年6月30日から7月1日現在、"
  "調査は令和7年4月1日現在、"
  "見える化システムは令和6年度の値であり、時点が異なります。")
P("")
SUB("（２） 取扱いを確定していない事項")
TBL(["事項", "内容", "確定しないと決まらないこと", "計画での使用"],
    [["介護職員数の重複",
      "同一法人が施設・通所系と訪問系の両方に同一の職員13人を記入している。"
      "個人単位の照合ができていない",
      "介護職員の総数が348人か361人か", "範囲で記述"],
     ["採用者数・離職者数の重複",
      "上記と同じ2事業所が同一の採用7人・離職2人を計上している",
      "採用83〜90人・離職65〜67人の別", "範囲で記述"],
     ["利用者票の回答の偏り",
      "99票のうち36票が同一法人の小規模多機能2事業所からの提出であり、"
      "「小規模多機能がより適切」44件のうち31件を占める",
      "小規模多機能の需要の大きさ", "使用しない"],
     ["利用者の所在地区",
      "記入形式が9種類に分かれ、地区別の集計ができない",
      "地区別・日常生活圏域別の分析の可否", "使用しない"],
     ["老健・通所リハの介護職員数",
      "調査77人に対し見える化システムは59人。"
      "全施設が回答したうえで差が生じている",
      "従事者数の推移を見える化システムで記述できるか", "参考として記述"],
     ["居所変更実態調査の未回答13施設",
      "区域内31施設のうち13施設が未回答（施設数ベースの把握率58.1%）。"
      "とくに区域内最大の特定施設が回答していない。"
      "定員・戸数は北海道の名簿で一部補完できるが、"
      "入居者数・職員数は公表データで補えない施設があり、"
      "入所前の居場所と退去先は補完できない",
      "入所前の居場所・退去先の集計の代表性", "範囲を明記して記述"],
     ["特別養護老人ホームの待機者",
      "同一法人の2施設で同数が計上されており、重複を除くと137人が82人となる",
      "整備の必要量の判断", "範囲で記述"],
     ["認知症対応型共同生活介護1事業所の定員",
      "グループホームくるみの郷は介護サービス情報公表システムに掲載がなく、"
      "定員を公表資料で確認できない。"
      "区域内定員99人及び717人は同事業所を含まない",
      "定員ベースの把握率74.1%の確定", "暫定値と明記して記述"],
     ["特定施設の入居者の要介護度構成",
      "回答2施設（57人・令和7年4月1日）と公表画面1施設"
      "（97人・令和7年10月6日）で時点が異なる。"
      "要支援1〜要介護2は合算154人で110人（71.4%）、"
      "回答2施設のみでは57人で30人（52.6%）",
      "居住系の整備の必要量の判断", "使用しない"]],
    [3.0, 6.4, 4.8, 2.6], num_from=99)
NOTE("「計画での使用」は、確定するまでの間の取扱いです。"
     "「範囲で記述」は最小値と最大値を併記するもの、"
     "「参考として記述」は本文に数値を載せず注記にとどめるもの、"
     "「使用しない」は確定するまで計画本文に載せないものです。")
NOTE("各事項の確認主体、確認の期限、確認後の反映先及び版は、"
     "内部管理用の「実施済み3調査の受領点検と集計」"
     "（02 点検で見つかった事項）で管理しています。")

H2("第４節　次期調査に向けた申し送り")
P("４つの調査を横断したクロス集計は、"
  "個票を共通の単位（日常生活圏域）に割り付けられないため実施できません。"
  "次期の調査に向けて、次の３点を改善することで"
  "横断的な分析が可能になります。")
P("① 在宅生活改善調査の利用者票に、"
  "日常生活圏域の選択肢を具体的な地区名で設けること。"
  "現在の調査票は「1.○○　2.○○　3.○○　4.○○」のまま配布されており、"
  "記入形式が事業所ごとに異なっています。")
P("② 居所変更実態調査の入所前の居場所・退去先の選択肢を、"
  "「市内／市外」から町名の選択肢に改めること。"
  "現在は「市内」が町単位か広域連合単位かの定義が示されていません。")
P("③ 介護人材実態調査の事業所票に、事業所の所在町を記入する欄を設けること。")
NOTE("調査票の企画・設計は発注者が行うものであり、"
     "本項は受託者からの提案です。")

# ================================================================== 参考
doc.add_page_break()
H1("（参考）　用語集")
P("この用語解説は、本報告書で使用している主な用語を、"
  "できるだけわかりやすく説明したものです。"
  "わからない語句があった際の参考として活用してください。")
for head, items in [
    ("【か】", [
        ("介護サービス情報公表システム（かいごさーびすじょうほうこうひょうしすてむ）",
         "介護保険法に基づき、指定を受けた介護サービス事業所が"
         "年１回、事業所の情報を報告し、公表する仕組みです。"
         "職種別の従業者数や前年度の退職者数など、"
         "他の統計にはない項目が事業所ごとに公表されます。"),
        ("介護福祉士（かいごふくしし）",
         "介護系資格の中で唯一の国家資格です。"
         "身体介護や生活援助だけでなく、"
         "他の介護職員への指導などの役割も担います。"),
        ("介護職員初任者研修修了（かいごしょくいんしょにんしゃけんしゅうしゅうりょう）",
         "介護の基礎知識と技術を習得する、"
         "介護職のスタートラインとなる研修の修了者です。"
         "旧訪問介護員（ホームヘルパー）２級に相当します。"),
        ("居住系サービス（きょじゅうけいさーびす）",
         "認知症対応型共同生活介護（グループホーム）と"
         "特定施設入居者生活介護（介護付有料老人ホーム等）をいいます。"
         "施設サービスとは区別されます。"),
        ("軽費老人ホーム（けいひろうじんほーむ）",
         "低額な料金で高齢者を入所させ、"
         "食事の提供などの日常生活上の支援を行う施設です。"
         "ケアハウスはその一類型です。"),
    ]),
    ("【さ】", [
        ("住所地特例（じゅうしょちとくれい）",
         "他の市町村の施設等に入所・入居した場合に、"
         "入所前の市町村が引き続き保険者となる仕組みです。"
         "施設が所在する市町村に給付費の負担が偏らないようにするための"
         "制度です。"),
        ("住宅型有料老人ホーム（じゅうたくがたゆうりょうろうじんほーむ）",
         "食事などの生活支援サービスが付いた高齢者向けの住まいです。"
         "介護が必要になった場合は、"
         "外部の訪問介護等を利用します。"
         "介護保険の指定サービスではないため、"
         "介護保険の統計には事業所として現れません。"),
        ("小規模多機能型居宅介護（しょうきぼたきのうがたきょたくかいご）",
         "通い（デイサービス）を中心に、"
         "泊まりと訪問を組み合わせて利用できる地域密着型サービスです。"
         "利用者は登録した１つの事業所のサービスを利用します。"),
    ]),
    ("【た】", [
        ("地域密着型サービス（ちいきみっちゃくがたさーびす）",
         "原則としてその市町村（本広域連合の場合は区域内）の"
         "被保険者のみが利用できるサービスです。"
         "指定は市町村（広域連合）が行います。"),
        ("通常の事業の実施地域（つうじょうのじぎょうのじっしちいき）",
         "事業所が運営規程に定める、"
         "通常のサービス提供の範囲となる区域です。"
         "この区域を超えて利用する場合、"
         "交通費等の実費を負担することがあります。"
         "利用できないという意味ではありません。"),
        ("特定施設入居者生活介護（とくていしせつにゅうきょしゃせいかつかいご）",
         "有料老人ホーム等のうち、"
         "介護保険の指定を受けて施設の職員が介護を提供するものです。"
         "介護付有料老人ホームとも呼ばれます。"),
    ]),
    ("【は】", [
        ("フレイル",
         "加齢に伴って心身の活力が低下し、"
         "介護が必要になる手前の状態をいいます。"
         "適切な対応により回復が期待できる段階です。"),
        ("見える化システム（みえるかしすてむ）",
         "地域包括ケア「見える化」システムのことで、"
         "厚生労働省が全国の保険者の介護保険に関する指標を"
         "比較できる形で提供するものです。"),
    ]),
]:
    P(head, bold=True, space_after=3)
    for term, desc in items:
        P("○　" + term, bold=True, space_after=2)
        P(desc, space_after=6)

P("")
P("以上", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

doc.save(OUT)
print("saved:", OUT)
print("段落 %d / 表 %d" % (len(doc.paragraphs), len(doc.tables)))
