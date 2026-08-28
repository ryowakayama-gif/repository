# -*- coding: utf-8 -*-
"""
川崎町 ニーズ調査結果報告書 修正スクリプト（案B：分母は現行のまま維持）

方針
  1. 概要表（主な項目／割合）の行ごとの (n=…) 表記を削除し、母数は注記で一括して示す
  2. 概要表の直下に母数の定義注記を追加、３章冒頭に母数の相違の注記を追加
  3. 転記漏れ（問8 幸福度の11段階分布）を補完
  4. 1表に2ブロックが同居している表を分割（身長・体重／社会参加）
  5. 体裁の統一（助詞・全半角・小数桁・括弧・セル内改行・並び順）
  6. グラフ73点の書体とデータラベル書式を本文に合わせる

割合は一切変更しない。原本の要素を複製する方式で書式を保持する。
"""
import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(__file__).resolve().parent.parent
SRC = BASE / "04_調査・入力・分析/R8調査データ/R8.8.28受領版/川崎町_ニーズ調査結果報告書_R8.8.28版.docx"
OUT = BASE / "01_第10期_最新版成果品/川崎町_ニーズ調査結果報告書_R8.9.2版.docx"

SUMMARY_TABLES = [1, 8, 19, 30, 49, 60, 63, 72, 80]   # 概要表（主な項目／割合）
CROSS_TABLES = list(range(83, 93))                    # ３－1〜３－10
T_HAPPY, T_HAPPY_TPL, T_HOBBY = 74, 73, 99
SPLITS = [(49, 9), (20, 3)]                           # (表番号, 分割する行)

NOTE_SUMMARY = (
    "※ 割合は、設問ごとに無回答を除いた有効回答数を分母として算出している"
    "（調査全体の母数は有効票576件）。設問別の有効回答数は、本節末尾の単純集計表に記載している。"
)
NOTE_CROSS = (
    "※ 本章の割合は、両設問（又は属性と設問）に有効回答があった方を分母としているため、"
    "第２章の単純集計とは母数が異なる。"
)
HAPPY = [
    ("0（とても不幸）", 2, "0.4%"), ("1", 5, "0.9%"), ("2", 4, "0.7%"),
    ("3", 11, "2.0%"), ("4", 19, "3.5%"), ("5", 104, "19.1%"),
    ("6", 49, "9.0%"), ("7", 102, "18.7%"), ("8", 107, "19.6%"),
    ("9", 44, "8.1%"), ("10（とても幸せ）", 98, "18.0%"),
]

log = []


# ---------------------------------------------------------------- ユーティリティ
def tables_of(doc):
    return [Table(c, doc) for c in doc.element.body.iterchildren() if c.tag == qn("w:tbl")]


def node_text(el):
    return "".join(n.text or "" for n in el.iter(qn("w:t")))


def node_sub(el, pattern, repl):
    """w:p / w:tc の配下の w:t を通しテキストとして扱って置換する（run分割に強い）。"""
    ts = [t for t in el.iter(qn("w:t"))]
    if not ts:
        return False
    whole = "".join(t.text or "" for t in ts)
    new = re.sub(pattern, repl, whole)
    if new == whole:
        return False
    ts[0].text = new
    ts[0].set(qn("xml:space"), "preserve")
    for t in ts[1:]:
        t.text = ""
    return True


def el_sub(el, pattern, repl):
    """w:tc なら段落単位、w:p ならその段落で置換する（セル内の段落構成を壊さない）。"""
    ps = el.findall(qn("w:p"))
    if not ps:
        return node_sub(el, pattern, repl)
    return any([node_sub(p, pattern, repl) for p in ps])


def iter_tc(tbl):
    for tr in tbl.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            yield tc


def find_par(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise KeyError(needle)


def clone_par(tpl_par, text):
    """注記段落のテンプレートを複製し、本文だけ差し替える。"""
    p = copy.deepcopy(tpl_par._p)
    par = Paragraph(p, tpl_par._parent)
    runs = par.runs
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    runs[0].text = text
    return p


def blank_par(tpl_par):
    p = copy.deepcopy(tpl_par._p)
    for r in Paragraph(p, tpl_par._parent).runs:
        r._element.getparent().remove(r._element)
    return p


# ---------------------------------------------------------------- 1. 概要表の母数表記
def strip_n_from_summary(doc):
    tabs = tables_of(doc)
    total = 0
    for ti in SUMMARY_TABLES:
        cnt = sum(el_sub(tc, r"[（(]\s*n\s*=\s*\d+\s*[）)]", "") for tc in iter_tc(tabs[ti]._tbl))
        total += cnt
        log.append(f"[1] 表{ti}：行別の(n=…)を{cnt}箇所削除")
    log.append(f"[1] 概要表9表 計{total}箇所を削除")


# ---------------------------------------------------------------- 2. 注記
def add_notes(doc, note_tpl):
    tabs = tables_of(doc)
    for ti in SUMMARY_TABLES:
        tbl = tabs[ti]._tbl
        nxt = tbl.getnext()
        if nxt is not None and nxt.tag == qn("w:p") and node_text(nxt).startswith("※ 割合は、設問ごとに"):
            continue
        tbl.addnext(clone_par(note_tpl, NOTE_SUMMARY))
    log.append(f"[2] 概要表{len(SUMMARY_TABLES)}表の直下に母数の注記を追加")

    anchor = find_par(doc, "ただし、列に複数回答の設問を取った")
    anchor._p.addnext(clone_par(note_tpl, NOTE_CROSS))
    log.append("[2] ３章冒頭に母数の相違に関する注記を追加")


# ---------------------------------------------------------------- 3. 幸福度分布の転記
def add_happiness_rows(doc):
    tabs = tables_of(doc)
    tbl = tabs[T_HAPPY]._tbl
    if "とても不幸" in node_text(tbl) and "0.4%" in node_text(tbl):
        log.append("[3] 幸福度分布は掲載済み（スキップ）")
        return
    src_rows = tabs[T_HAPPY_TPL]._tbl.findall(qn("w:tr"))
    odd, even = src_rows[2], src_rows[3]          # ゼブラの明細行テンプレート
    tail = tbl.findall(qn("w:tr"))[-1]            # 「平均値: 7.1 …」の行
    for i, (label, cnt, pct) in enumerate(HAPPY):
        tr = copy.deepcopy(odd if i % 2 == 0 else even)
        for tc, val in zip(tr.findall(qn("w:tc")), [label, "―", str(cnt), pct]):
            ts = list(tc.iter(qn("w:t")))
            ts[0].text = val
            ts[0].set(qn("xml:space"), "preserve")
            for t in ts[1:]:
                t.text = ""
        tail.addprevious(tr)
    log.append(f"[3] 表{T_HAPPY}に幸福度の11段階分布（n=545・計545件）を追加")


# ---------------------------------------------------------------- 5. 体裁
def fix_typography(doc):
    # F-1 助詞
    if node_sub(find_par(doc, "を合わせる48.2")._p, "を合わせる48.2", "を合わせると48.2"):
        log.append("[5] F-1 「合わせる48.2％」→「合わせると48.2％」")

    tabs = tables_of(doc)

    # F-2 表内の全角％を半角%に
    n = sum(el_sub(tc, r"(\d)％", r"\1%") for t in tabs for tc in iter_tc(t._tbl))
    log.append(f"[5] F-2 表内の全角％を半角%に統一（{n}セル）")

    # F-3/F-4 小数点のない百分率を小数第1位に
    n = sum(el_sub(tc, r"(?<![\d.])(\d+)%", r"\1.0%") for t in tabs for tc in iter_tc(t._tbl))
    log.append(f"[5] F-3/F-4 小数点なしの%を小数第1位に統一（{n}セル）")

    # F-5 件数の括弧を全角に統一（３－1〜３－10）
    n = sum(el_sub(tc, r"\s*\((\d+)件\)", r"（\1件）") for ti in CROSS_TABLES for tc in iter_tc(tabs[ti]._tbl))
    log.append(f"[5] F-5 件数の括弧を全角に統一（{n}セル）")

    # F-6 趣味カテゴリ表：セル内の分割段落を1段落にまとめ、件数の前を半角スペースに
    t = tabs[T_HOBBY]
    fixed = 0
    for row in t.rows:
        pars = row.cells[0].paragraphs
        if len(pars) <= 1:
            continue
        merged = " ".join(p.text.strip() for p in pars if p.text.strip())
        merged = re.sub(r"\s*(\d+件)", r" \1", merged).strip()
        node_sub(pars[0]._p, r"^.*$", lambda m: merged)
        for p in pars[1:]:
            p._p.getparent().remove(p._p)
        fixed += 1
    log.append(f"[5] F-6 趣味カテゴリ表のセル内改段落を整理（{fixed}行）")

    # F-6 件数の降順に並べ替え
    tbl = t._tbl
    rows = tbl.findall(qn("w:tr"))
    head, body = rows[0], rows[1:]

    def cnt_of(tr):
        m = re.search(r"(\d+)件", node_text(tr.findall(qn("w:tc"))[0]))
        return int(m.group(1)) if m else 0

    order = sorted(body, key=cnt_of, reverse=True)
    if [id(x) for x in order] != [id(x) for x in body]:
        for tr in body:
            tbl.remove(tr)
        prev = head
        for tr in order:
            prev.addnext(tr)
            prev = tr
        log.append("[5] F-6 趣味カテゴリ表を件数の降順に並べ替え")


# ---------------------------------------------------------------- 4. 表の分割
def split_tables(doc, blank_tpl):
    for ti, at_row in SPLITS:
        tabs = tables_of(doc)
        tbl = tabs[ti]._tbl
        rows = tbl.findall(qn("w:tr"))
        new_tbl = copy.deepcopy(tbl)
        for tr in rows[at_row:]:
            tbl.remove(tr)
        for tr in new_tbl.findall(qn("w:tr"))[:at_row]:
            new_tbl.remove(tr)
        tbl.addnext(new_tbl)
        tbl.addnext(blank_par(blank_tpl))   # 表どうしが結合しないよう空段落を挟む
        log.append(f"[4] 表{ti}を第{at_row}行の直前で2表に分割")


# ---------------------------------------------------------------- 6. グラフ
def fix_charts(path):
    tmp = path.with_suffix(".tmp.docx")
    n_font = n_fmt = n_chart = 0
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if re.match(r"word/charts/chart\d+\.xml$", item.filename):
                x = data.decode("utf-8")
                n_font += x.count("ＭＳ Ｐゴシック")
                n_fmt += x.count('<c:numFmt formatCode="0%" sourceLinked="0"/>')
                y = x.replace("ＭＳ Ｐゴシック", "BIZ UDPゴシック").replace(
                    '<c:numFmt formatCode="0%" sourceLinked="0"/>',
                    '<c:numFmt formatCode="0.0%" sourceLinked="0"/>',
                )
                if y != x:
                    n_chart += 1
                data = y.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)
    log.append(f"[6] グラフ{n_chart}点：書体を{n_font}箇所置換／データラベル書式を{n_fmt}箇所 0.0% に変更")


# ---------------------------------------------------------------- 実行
def main():
    doc = Document(str(SRC))
    note_tpl = find_par(doc, "※ 転倒予防・移動手段の確保は")

    strip_n_from_summary(doc)
    add_happiness_rows(doc)
    fix_typography(doc)
    add_notes(doc, note_tpl)
    split_tables(doc, note_tpl)     # 表番号がずれるため最後に実施

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    fix_charts(OUT)

    print("\n".join(log))
    print(f"\n出力: {OUT}")


if __name__ == "__main__":
    main()
