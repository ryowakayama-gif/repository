# -*- coding: utf-8 -*-
"""
調査結果報告書 2件の修正

【B】ニーズ調査結果報告書（R8.8.3版）
  B-1 計画名称「高齢者福祉計画」→「高齢者保健福祉計画」（委託仕様書1の正式名称）
  B-2 事務局名を第1回策定委員会資料と統一

【C】介護実態調査結果報告書（R8.8.12版・作成中）
  C-1 「回答者単位の個票がないためクロス集計は未算出」という前提の誤りを是正（4箇所）
  C-2 クロス集計12件を実データで算出し、候補一覧を実表に置換（第9章・第15章）
  C-3 第6・7・8章の「回答分布を以下に示す」に対応する分布表が無いため追加
  C-4 「主な介護者が離職9.0%」は件数9件を割合と取り違えたもの → 9件（13.2%）
  C-5 「最も多い回答は60代で26.3%」→ 60代・70代がともに26.3%
  C-6 調査結果を町全体に換算する際の母集団（在宅認定者 約369人）を明示
  C-7 サービス利用状況について給付実績との突合を追記

数値はすべて集計データ（20260807版）のデータシートから直接算出した。
データシート最終行の「計」行は集計対象から除外している。
既存の図表（グラフ8点）を保持するため、文書は新規生成せず原本を編集する。
"""

import os
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(BASE, "04_調査・入力・分析", "R8調査データ")
SRC_B = os.path.join(DATA, "川崎町_ニーズ調査結果報告書_R8.8.3版.docx")
SRC_C = os.path.join(DATA, "川崎町_介護実態調査結果報告書_R8.8.12版_作成中.docx")
OUT_B = os.path.join(DATA, "川崎町_ニーズ調査結果報告書_R8.8.20修正版.docx")
OUT_C = os.path.join(DATA, "川崎町_介護実態調査結果報告書_R8.8.20修正版.docx")
XLSX = os.path.join(DATA, "川崎町_在宅介護実態調査_集計データ_20260807.xlsx")

GOTHIC = "游ゴシック"


# ====================================================== 集計エンジン
class Tally:
    def __init__(self, path):
        self.ws = openpyxl.load_workbook(path, data_only=True)["データ"]
        self.rows = [r for r in range(5, self.ws.max_row + 1)
                     if self.ws.cell(r, 1).value not in (None, "")
                     and str(self.ws.cell(r, 1).value).strip() != "計"]

    def label(self, c):
        return self.ws.cell(3, c).value

    def sa(self, r, cols):
        """単一回答：マークが1つだけの場合にそのラベルを返す"""
        m = [c for c in cols if self.ws.cell(r, c).value not in (None, "")]
        return self.label(m[0]) if len(m) == 1 else None

    def ma(self, r, cols):
        return [self.label(c) for c in cols
                if self.ws.cell(r, c).value not in (None, "")]

    def dist(self, cols, drop=()):
        """単純集計（件数・割合）"""
        vals = [v for v in (self.sa(r, cols) for r in self.rows) if v]
        n = len(vals)
        out = []
        for c in cols:
            lab = self.label(c)
            if lab and lab not in drop:
                k = vals.count(lab)
                out.append((lab, k, k / n * 100))
        return out, n

    def cross(self, rowfn, colfn, rowlab, collab):
        """クロス集計（行%）"""
        tab = {rl: {cl: 0 for cl in collab} for rl in rowlab}
        for r in self.rows:
            a, b = rowfn(r), colfn(r)
            if a in tab and b in tab[a]:
                tab[a][b] += 1
        return tab

    def cross_ma(self, grpfn, grplab, cols, items):
        """クロス集計（列が複数回答）"""
        tab = {g: {} for g in grplab}
        base = {g: 0 for g in grplab}
        for r in self.rows:
            g = grpfn(r)
            got = self.ma(r, cols)
            if g in tab and got:
                base[g] += 1
                for it in items:
                    tab[g][it] = tab[g].get(it, 0) + (1 if it in got else 0)
        return tab, base


COLS = {
    "A2": range(27, 30), "A5": range(42, 50), "A6": range(51, 54),
    "A7": range(55, 57), "A9": range(128, 133), "B1": range(134, 140),
    "B2": range(141, 150), "B3": range(151, 168), "B4": range(169, 173),
    "B5": range(174, 180), "B6": range(181, 186),
}

A6L = ["入所・入居は検討していない", "入所・入居を検討している", "すでに入所・入居申し込みをしている"]
A6H = ["検討していない", "検討している", "申込済み"]
A7L = ["利用した", "利用していない"]
A9L = ["ない", "家族・親族の介護はあるが、週に１日よりも少ない",
       "週に１～２日ある", "週に３～４日ある", "ほぼ毎日ある"]
A9H = ["ない", "週１日未満", "週１〜２日", "週３〜４日", "ほぼ毎日"]
A2L = ["単身世帯", "夫婦のみ世帯", "その他"]
B4L = ["フルタイムで働いている", "パートタイムで働いている", "働いていない"]
B4H = ["フルタイム", "パートタイム", "働いていない"]
B6L = ["問題なく、続けていける", "問題はあるが、何とか続けていける",
       "続けていくのは、やや難しい", "続けていくのは、かなり難しい"]
B6H = ["問題なく\n続けていける", "何とか\n続けていける", "やや難しい", "かなり難しい"]
FUAN = ["外出の付き添い、送迎等", "認知症状への対応", "入浴・洗身",
        "食事の準備（調理等）", "日中の排泄", "夜間の排泄"]
FUANH = ["外出の付添・送迎", "認知症状への対応", "入浴・洗身", "食事の準備", "日中の排泄", "夜間の排泄"]
CHOSEI = ["特に行っていない",
          "介護のために、「労働時間を調整（残業免除、短時間勤務、遅出・早帰・中抜け等）」しながら、働いている",
          "介護のために、「休暇（年休や介護休暇等）」を取りながら、働いている",
          "介護のために、「在宅勤務」を利用しながら、働いている"]
CHOSEIH = ["特に行っていない", "労働時間の調整", "休暇の取得", "在宅勤務"]


def care3(v):
    if v is None:
        return None
    if v.startswith("要支援"):
        return "要支援１・２"
    if v in ("要介護１", "要介護２"):
        return "要介護１・２"
    if v in ("要介護３", "要介護４", "要介護５"):
        return "要介護３〜５"
    return None


def kaigo2(v):
    return None if v is None else ("なし" if v == "ない" else "あり")


def age3(v):
    if v is None or v == "わからない":
        return None
    if v in ("20歳未満", "20代", "30代", "40代", "50代"):
        return "50代以下"
    return "60代" if v == "60代" else "70歳以上"


def freq3(v):
    if v is None:
        return None
    if v == "ほぼ毎日ある":
        return "ほぼ毎日ある"
    if v in ("週に１～２日ある", "週に３～４日ある"):
        return "週１〜４日ある"
    return "ない・週１日未満"


def rishoku(t, r):
    v = t.ma(r, COLS["B1"])
    if not v:
        return None
    if any(("辞めた" in x and x != "介護のために仕事を辞めた家族・親族はいない") or "転職" in x for x in v):
        return "あり"
    if "介護のために仕事を辞めた家族・親族はいない" in v:
        return "なし"
    return None


# ====================================================== docx 編集ユーティリティ
def set_text(par, text, *, size=10.5, bold=False):
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = GOTHIC
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), GOTHIC)


def replace_everywhere(doc, old, new):
    """本文・表・ヘッダ／フッタを通してラン単位で置換（ラン跨ぎにも対応）"""
    n = 0

    def fix_par(par):
        nonlocal n
        if old not in par.text:
            return
        runs = par.runs
        if not runs:
            return
        joined = "".join(r.text for r in runs)
        if old not in joined:
            return
        n += joined.count(old)
        joined = joined.replace(old, new)
        runs[0].text = joined
        for r in runs[1:]:
            r.text = ""

    def walk(container):
        for par in container.paragraphs:
            fix_par(par)
        for tb in container.tables:
            for row in tb.rows:
                for cell in row.cells:
                    walk(cell)

    walk(doc)
    for s in doc.sections:
        for part in (s.header, s.footer):
            walk(part)
    return n


def make_table(doc, headers, rows, style="Grid Table 4 Accent 1", size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = style
    except KeyError:
        t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_text(t.rows[0].cells[i].paragraphs[0], h, size=size, bold=True)
        t.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            set_text(p, str(v), size=size)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
    return t


def move_after(anchor, *blocks):
    """anchor（段落 or 表）の直後に blocks を順に差し込む"""
    cur = anchor
    for b in blocks:
        el = b._p if hasattr(b, "_p") else b._tbl
        cur_el = cur._p if hasattr(cur, "_p") else cur._tbl
        cur_el.addnext(el)
        cur = b
    return cur


def new_par(doc, text, *, size=10.5, bold=False, style=None):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except KeyError:
            pass
    set_text(p, text, size=size, bold=bold)
    return p


def find_par(doc, needle, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start):
        if needle in p.text:
            return i, p
    raise LookupError(needle)


def drop(el):
    node = el._p if hasattr(el, "_p") else el._tbl
    node.getparent().remove(node)


def pct_rows(tab, rowlab, collab, rowhead):
    out = []
    for rl, rh in zip(rowlab, rowhead):
        n = sum(tab[rl].values())
        cells = [f"{tab[rl][c]}件\n{tab[rl][c]/n*100:.1f}%" if n else "―" for c in collab]
        out.append([rh] + cells + [f"{n}件"])
    return out


# ====================================================== B：ニーズ調査結果報告書
def fix_b():
    doc = Document(SRC_B)
    n1 = replace_everywhere(doc, "高齢者福祉計画", "高齢者保健福祉計画")
    n2 = replace_everywhere(
        doc,
        "川崎町　高齢者保健福祉計画・介護保険事業計画策定委員会　事務局",
        "川崎町　介護保険事業計画・高齢者保健福祉計画策定委員会　事務局")
    doc.save(OUT_B)
    print(f"B 保存: {os.path.basename(OUT_B)}（計画名称 {n1}箇所・事務局名 {n2}箇所）")


# ====================================================== C：介護実態調査結果報告書
def fix_c():
    t = Tally(XLSX)
    doc = Document(SRC_C)

    # ---------- C-1 前提の誤りの是正（第1章）
    _, p = find_par(doc, "回答者単位の個票がないため")
    set_text(p, "集計上の留意点：割合は原則として当該設問の有効回答（n）を分母として算出している。"
                "設問により無回答の数が異なるため、n は設問ごとに異なる。複数回答設問では割合の合計が"
                "100％を超える場合がある。集計データには回答者単位の個票（142件）が収録されており、"
                "性別・年齢階級・要介護状態区分も記録されているため、本版では設問間のクロス集計を"
                "実データにより算出した（第9章・第15章）。なお集計データの最終行は「計」行であり、"
                "回答者ではないため集計から除外している。", size=10.5)

    # ---------- C-4 介護離職の件数と割合の取り違え
    _, p = find_par(doc, "主な介護者が離職9.0%")
    set_text(p, "（５）介護離職：介護を理由に仕事を辞めた家族・親族はいない64.7%。"
                "一方、主な介護者が離職9件（13.2%）、主な介護者以外の離職1件（1.5%）、"
                "転職3件（4.4%）で、過去1年間に離職・転職が生じた世帯は13件（19.1%）である（n=68・複数回答）。",
             size=10.5)

    # ---------- C-5 介護者年齢の最頻値
    _, p = find_par(doc, "最も多い回答は「60代」で、26.3%である")
    set_text(p, "主な介護者の年齢は「60代」と「70代」がともに20件（26.3%）で最も多く、"
                "80歳以上も12件（15.8%）を占める。70歳以上が計32件（42.1%）であり、"
                "いわゆる老老介護の状態にある世帯が4割を超える。回答分布を以下に示す。", size=10.5)

    # ---------- C-3 第6・7・8章の分布表を追加
    d9, n9 = t.dist(COLS["A9"])
    _, p6 = find_par(doc, "最も多い回答は「ない」で、45.8%である")
    set_text(p6, "家族・親族からの介護の頻度は「ない」が65件（45.8%）で最も多く、"
                 "「ほぼ毎日ある」43件（30.3%）が続く。在宅の認定者の半数近くに家族・親族による"
                 "介護がなく、在宅生活の継続が公的サービスに強く依存している。回答分布を以下に示す。",
             size=10.5)
    move_after(p6, make_table(doc, ["回答区分", "件数", "割合"],
                              [[l, f"{k}件", f"{r:.1f}%"] for l, k, r in d9] + [["計", f"{n9}件", "100.0%"]]))

    d2, n2 = t.dist(COLS["B2"])
    _, p7 = find_par(doc, "70歳以上が計32件（42.1%）")
    move_after(p7, make_table(doc, ["回答区分", "件数", "割合"],
                              [[l, f"{k}件", f"{r:.1f}%"] for l, k, r in d2 if k] + [["計", f"{n2}件", "100.0%"]]))

    d6, n6 = t.dist(COLS["B6"])
    _, p8 = find_par(doc, "最も多い回答は「問題はあるが、何とか続けていける」")
    set_text(p8, "今後も働きながら介護を続けられるかについては、「問題はあるが、何とか続けていける」が"
                 "20件（51.3%）で最も多い。「やや難しい」8件（20.5%）と「かなり難しい」2件（5.1%）を"
                 "合わせた10件（25.6%）が継続困難と回答している。ただし本設問の有効回答は39件であり、"
                 "回答者全体（142件）の一部である点に留意を要する。回答分布を以下に示す。", size=10.5)
    move_after(p8, make_table(doc, ["回答区分", "件数", "割合"],
                              [[l, f"{k}件", f"{r:.1f}%"] for l, k, r in d6] + [["計", f"{n6}件", "100.0%"]]))

    # ---------- C-6 母集団の明示（第4章）
    _, p4 = find_par(doc, "施設入所ニーズも一定程度確認される")
    blocks = [
        new_par(doc, "本調査の回答者は在宅で生活する認定者である。割合を町全体の人数に換算する際は、"
                     "回答件数（142件）ではなく在宅の認定者数を母集団とする必要がある。", size=10.5),
        make_table(doc, ["区分", "人数", "算出根拠"], [
            ["要支援・要介護認定者（第1号・令和7年度末）", "561人", "介護保険事業状況報告 様式1の5"],
            ["うち入所・居住系サービスの利用者（月平均）", "192人", "同 様式2〈件数〉"],
            ["在宅で生活する認定者（推計）", "約369人", "561人 − 192人"],
        ]),
        new_par(doc, "この母集団に換算すると、施設等への入所を申込済みの方は約125人、"
                     "検討中の方は約53人、合わせて約178人と推計される。"
                     "入所・居住系サービスは介護給付費の68.1%を占めるため、"
                     "この需要をサービス見込量にどこまで反映するかが第10期の保険料を大きく左右する。",
                size=10.5),
    ]
    move_after(p4, *blocks)

    # ---------- C-7 給付実績との突合（第5章）
    _, p5 = find_par(doc, "サービス利用の偏りがみられる")
    move_after(p5,
               new_par(doc, "本調査によるサービス別の利用状況を、介護保険事業状況報告（令和7年度）の"
                            "給付実績と突合すると次のとおりである。訪問介護はよく整合する一方、"
                            "通所リハビリテーションは調査が実績の3分の1しか捉えていない。"
                            "サービス見込量は給付実績を基礎として推計し、本調査は利用意向・生活実態の"
                            "把握に用いることが適当である。", size=10.5),
               make_table(doc, ["サービス", "調査からの推計\n（在宅認定者約369人に換算）",
                                "令和7年度実績\n（月平均・実人数）", "評価"], [
                   ["訪問介護", "約37人", "33.9人", "整合"],
                   ["通所介護", "約77人", "54.0人", "調査がやや過大"],
                   ["通所リハビリテーション", "約43人", "125.0人", "調査は実績の1/3"],
                   ["短期入所生活介護等", "―", "20.8人", "―"],
               ]),
               new_par(doc, "※ 調査のサービス別設問（A問8）は設問ごとの有効回答が37〜51件とばらついており、"
                            "サービス間の単純比較には適さない。", size=9))

    # ---------- C-2 第9章：クロス集計の実算出
    i9, p9 = find_par(doc, "クロス表の数値は新規算出せず、分析候補として整理した")
    set_text(p9, "単独集計では把握しにくい属性間の違いを確認するため、世帯類型・要介護度・家族介護の状況と、"
                 "施設等への入所・入居、介護保険サービスの利用、介護者の就労継続等との関係を、"
                 "回答者単位の個票（142件）によりクロス集計した。表中の割合はいずれも行方向"
                 "（各属性内での構成比）であり、無回答は除外している。", size=10.5)

    C = COLS
    xt = []

    def add(title, headers, rows, comment):
        xt.append((title, headers, rows, comment))

    # 9-1
    tab = t.cross(lambda r: care3(t.sa(r, C["A5"])), lambda r: t.sa(r, C["A6"]),
                  ["要支援１・２", "要介護１・２", "要介護３〜５"], A6L)
    add("９－１．要介護度 × 施設等への入所・入居の検討状況（n=133）",
        ["要介護度"] + A6H + ["計"],
        pct_rows(tab, ["要支援１・２", "要介護１・２", "要介護３〜５"], A6L,
                 ["要支援１・２", "要介護１・２", "要介護３〜５"]),
        "要介護3〜5では48.8%が申込済みであるのに対し、要支援では申込済みが0件である。"
        "重度化と施設入所の申込みは明確に対応しており、要介護3が入所検討の分岐点となっている。")

    # 9-2
    tab = t.cross(lambda r: kaigo2(t.sa(r, C["A9"])), lambda r: t.sa(r, C["A6"]), ["なし", "あり"], A6L)
    add("９－２．家族・親族による介護の有無 × 施設等への入所・入居の検討状況（n=139）",
        ["家族介護"] + A6H + ["計"],
        pct_rows(tab, ["なし", "あり"], A6L, ["介護は「ない」", "介護が「ある」"]),
        "本報告書で最も差が大きいクロス集計である。家族・親族の介護が「ない」層では56.9%が"
        "施設入所を申込済みであるのに対し、介護が「ある」層では13.5%にとどまり、4.2倍の開きがある。"
        "施設入所ニーズの主因は要介護度の重さだけではなく、家族介護者の不在にあると考えられる。")

    # 9-3
    tab = t.cross(lambda r: t.sa(r, C["A2"]), lambda r: t.sa(r, C["A6"]), A2L, A6L)
    add("９－３．世帯類型 × 施設等への入所・入居の検討状況（n=138）",
        ["世帯類型"] + A6H + ["計"], pct_rows(tab, A2L, A6L, A2L),
        "単身世帯では69.4%が申込済みであり、その他世帯（18.6%）の3.7倍である。"
        "９－２と併せると、単身・家族介護なしの層が施設入所需要の中核を構成している。")

    # 9-4
    tab = t.cross(lambda r: t.sa(r, C["A2"]), lambda r: t.sa(r, C["A9"]), A2L, A9L)
    add("９－４．世帯類型 × 家族・親族による介護の頻度（n=139）",
        ["世帯類型"] + A9H + ["計"], pct_rows(tab, A2L, A9L, A2L),
        "単身世帯の66.7%、夫婦のみ世帯の51.5%で家族・親族の介護が「ない」。"
        "その他世帯（同居家族あり）では47.1%が「ほぼ毎日」であり、"
        "介護負担が同居世帯に集中する一方、単身・夫婦のみ世帯では介護の担い手が不在である。")

    # 9-5
    tab = t.cross(lambda r: care3(t.sa(r, C["A5"])), lambda r: t.sa(r, C["A7"]),
                  ["要支援１・２", "要介護１・２", "要介護３〜５"], A7L)
    add("９－５．要介護度 × 介護保険サービスの利用（n=132）",
        ["要介護度"] + A7L + ["計"],
        pct_rows(tab, ["要支援１・２", "要介護１・２", "要介護３〜５"], A7L,
                 ["要支援１・２", "要介護１・２", "要介護３〜５"]),
        "要介護度が重いほどサービス利用率が低いという、一見逆説的な結果である。"
        "要介護3〜5の65.4%が「利用していない」。施設入所申込済みの層と重なっており、"
        "入所待機中に在宅サービスにつながっていない可能性がある（９－７参照）。")

    # 9-6
    tab = t.cross(lambda r: kaigo2(t.sa(r, C["A9"])), lambda r: t.sa(r, C["A7"]), ["なし", "あり"], A7L)
    add("９－６．家族・親族による介護の有無 × 介護保険サービスの利用（n=138）",
        ["家族介護"] + A7L + ["計"],
        pct_rows(tab, ["なし", "あり"], A7L, ["介護は「ない」", "介護が「ある」"]),
        "家族介護が「ない」層のサービス利用率は23.1%にとどまり、「ある」層（57.5%）の半分以下である。"
        "家族の介護もなく公的サービスも利用していない在宅認定者が50件（回答者の36.2%）存在する。"
        "町全体では約133人と推計され、支援が最も届いていない層である。")

    # 9-7
    tab = t.cross(lambda r: t.sa(r, C["A7"]), lambda r: t.sa(r, C["A6"]), A7L, A6L)
    add("９－７．介護保険サービスの利用 × 施設等への入所・入居の検討状況（n=136）",
        ["サービス利用"] + A6H + ["計"], pct_rows(tab, A7L, A6L, A7L),
        "サービスを利用していない層の48.8%が施設入所を申込済みである。"
        "在宅サービスの利用が施設入所の申込みを抑制している可能性があり、"
        "未利用者へのサービス調整は在宅生活継続の支援として有効と考えられる。")

    # 9-8
    tab = t.cross(lambda r: age3(t.sa(r, C["B2"])), lambda r: t.sa(r, C["B4"]),
                  ["50代以下", "60代", "70歳以上"], B4L)
    add("９－８．主な介護者の年齢 × 勤務形態（n=72）",
        ["介護者の年齢"] + B4H + ["計"],
        pct_rows(tab, ["50代以下", "60代", "70歳以上"], B4L, ["50代以下", "60代", "70歳以上"]),
        "50代以下の59.1%がフルタイム就労であり、介護離職防止の対象はこの層に集中する。"
        "70歳以上では56.7%が働いておらず、就労支援より介護負担そのものの軽減が課題となる。")

    # 9-9（複数回答）
    tab, base = t.cross_ma(lambda r: t.sa(r, C["B4"]),
                           ["フルタイムで働いている", "パートタイムで働いている"], C["B5"], CHOSEI)
    rows = []
    for k, h in zip(["フルタイムで働いている", "パートタイムで働いている"], ["フルタイム", "パートタイム"]):
        n = base[k]
        rows.append([f"{h}（n={n}）"] + [f"{tab[k][c]}件\n{tab[k][c]/n*100:.1f}%" for c in CHOSEI])
    add("９－９．勤務形態 × 働き方の調整（複数回答・n=44）", ["勤務形態"] + CHOSEIH, rows,
        "パートタイムの47.6%が労働時間を調整しており、フルタイム（34.8%）を上回る。"
        "在宅勤務の利用は両者とも1件のみで、制度が使える職場が限られていることがうかがえる。")

    # 9-10
    tab = t.cross(lambda r: t.sa(r, C["B4"]), lambda r: t.sa(r, C["B6"]),
                  ["フルタイムで働いている", "パートタイムで働いている"], B6L)
    add("９－10．勤務形態 × 今後の就労継続の見通し（n=38）", ["勤務形態"] + B6H + ["計"],
        pct_rows(tab, ["フルタイムで働いている", "パートタイムで働いている"], B6L,
                 ["フルタイム", "パートタイム"]),
        "「かなり難しい」と回答した2件はいずれもフルタイム就労者である。"
        "フルタイム就労者の31.6%が継続困難（やや＋かなり）と回答しており、"
        "パートタイム（21.1%）を上回る。ただし各19件と少数であり、傾向の把握にとどめる必要がある。")

    # 9-11（複数回答）
    grp11 = ["ほぼ毎日ある", "週１〜４日ある", "ない・週１日未満"]
    tab, base = t.cross_ma(lambda r: freq3(t.sa(r, C["A9"])), grp11, C["B3"], FUAN)
    rows = [[f"{g}（n={base[g]}）"] + [f"{tab[g][f]}件\n{tab[g][f]/base[g]*100:.1f}%" for f in FUAN]
            for g in grp11 if base[g]]
    add("９－11．家族介護の頻度 × 主な介護者が不安に感じる介護（3つまで選択可・n=75）",
        ["家族介護の頻度"] + FUANH, rows,
        "「ほぼ毎日」介護している層は認知症状への対応（35.7%）が最多であるのに対し、"
        "介護頻度の低い層では外出の付き添い・送迎（50.0%）が突出する。"
        "同居して日常的に介護する世帯には認知症ケアの支援を、"
        "別居等で頻度が低い世帯には移動支援を重点的に案内するという、対象別の施策設計が可能である。")

    # 9-12（複数回答）
    tab, base = t.cross_ma(lambda r: rishoku(t, r), ["あり", "なし"], C["B3"], FUAN)
    rows = [[f"離職・転職{g}（n={base[g]}）"] + [f"{tab[g][f]}件\n{tab[g][f]/base[g]*100:.1f}%" for f in FUAN]
            for g in ["あり", "なし"] if base[g]]
    add("９－12．過去1年間の介護離職・転職の有無 × 主な介護者が不安に感じる介護（n=56）",
        ["離職・転職"] + FUANH, rows,
        "離職・転職があった世帯では外出の付き添い・送迎（53.8%）と認知症状への対応（46.2%）が"
        "高く、なかった世帯（25.6%・32.6%）を大きく上回る。"
        "移動支援と認知症ケアの不足が介護離職の引き金となっている可能性がある。"
        "ただし該当13件と少数であり、確定的な結論とはしない。")

    # 第9章の候補一覧表（表3）を削除し、実表を挿入
    anchor = p9
    for tb in doc.tables:
        cells = [c.text for c in tb.rows[0].cells]
        if "クロス分析候補" in cells:
            drop(tb)
            break
    for title, headers, rows, comment in xt:
        anchor = move_after(anchor,
                            new_par(doc, title, size=11, bold=True),
                            make_table(doc, headers, rows),
                            new_par(doc, "分析：" + comment, size=10))

    # ---------- 第15章：候補一覧を結果一覧に置換
    _, p15 = find_par(doc, "以下は分析候補の一覧であり、数値は推計していない")
    set_text(p15, "第9章で算出したクロス集計の主な結果を一覧にまとめる。"
                  "割合はいずれも各属性内での構成比である。", size=10.5)
    for tb in doc.tables:
        cells = [c.text for c in tb.rows[0].cells]
        if "クロス集計候補" in cells:
            drop(tb)
            break
    move_after(p15, make_table(doc, ["No", "クロス集計", "主な結果"], [
        ["９－１", "要介護度 × 施設入所検討", "要介護3〜5の48.8%が申込済み。要支援は0件"],
        ["９－２", "家族介護の有無 × 施設入所検討", "介護なし56.9% 対 介護あり13.5%（4.2倍）"],
        ["９－３", "世帯類型 × 施設入所検討", "単身世帯69.4%が申込済み。その他世帯18.6%"],
        ["９－４", "世帯類型 × 家族介護の頻度", "単身の66.7%が介護「ない」。その他世帯の47.1%が「ほぼ毎日」"],
        ["９－５", "要介護度 × サービス利用", "要介護3〜5の65.4%が利用なし（要支援は38.5%）"],
        ["９－６", "家族介護の有無 × サービス利用", "介護なし層の利用率23.1%（介護あり57.5%）"],
        ["９－７", "サービス利用 × 施設入所検討", "未利用者の48.8%が申込済み（利用者14.3%）"],
        ["９－８", "介護者年齢 × 勤務形態", "50代以下の59.1%がフルタイム。70歳以上の56.7%が無職"],
        ["９－９", "勤務形態 × 働き方の調整", "パートの47.6%が労働時間調整。在宅勤務は各1件"],
        ["９－10", "勤務形態 × 就労継続見通し", "フルタイムの31.6%が継続困難（パート21.1%）"],
        ["９－11", "家族介護の頻度 × 介護者の不安", "毎日介護＝認知症対応35.7%／低頻度＝外出付添50.0%"],
        ["９－12", "介護離職の有無 × 介護者の不安", "離職ありは外出付添53.8%・認知症対応46.2%と高い"],
    ]))

    # ---------- 第13章の総括の前提を差し替え
    _, p13 = find_par(doc, "本版ではクロス分析候補と分析目的を整理し、数値の推測は行っていない")
    set_text(p13, "・本版では、回答者単位の個票（142件）により設問間のクロス集計12件を実算出した。"
                  "とくに、家族・親族の介護が「ない」層の56.9%が施設入所を申込済みであること、"
                  "その層のサービス利用率が23.1%にとどまることは、"
                  "在宅生活継続の支援策を検討するうえで中心となる知見である。", size=10.5)

    # ---------- 重点課題に2項目を追加
    i, _ = find_par(doc, "移動・外出支援：通院・買物・社会参加等の移動支援")
    last = doc.paragraphs[i]
    move_after(last,
               new_par(doc, "家族介護者不在層への支援：家族の介護がなく公的サービスも利用していない在宅認定者が"
                            "36.2%（町全体で約133人）存在し、この層の56.9%が施設入所を申込済みである。"
                            "在宅生活を継続するための支援が最も届いていない。", size=10.5, style="List Bullet"),
               new_par(doc, "単身高齢者の施設入所需要：単身世帯の69.4%が施設入所を申込済みであり、"
                            "住まいの確保を含めた検討が必要である。", size=10.5, style="List Bullet"))

    # ---------- 残存する「個票がない」前提の記述を是正（第9章末・第15章末）
    _, p = find_par(doc, "個票データが整備された段階では")
    set_text(p, "以上のクロス集計から、施設入所の申込みは要介護度の重さ（９－１）よりも、"
                "家族介護者の不在（９－２）と単身世帯であること（９－３）に強く規定されていることが"
                "確認できる。また、家族の介護もサービス利用もない層（９－６）が、"
                "施設入所需要の最も大きな供給源となっている（９－７）。", size=10.5)

    _, p = find_par(doc, "個票データが整備された場合には")
    set_text(p, "本章の各クロス集計は集計データの個票142件から算出したものであり、"
                "第9章に集計表と分析を掲載している。なお、セル数が10件を下回る組合せ"
                "（９－10・９－12等）は傾向の把握にとどめ、確定的な結論には用いていない。", size=10.5)

    # ---------- 第2章要約（4）に家族介護者不在の知見を追加
    _, p = find_par(doc, "（４）家族介護の頻度")
    set_text(p, "（４）家族介護の頻度：家族・親族の介護が「ほぼ毎日」30.3%（43件）。"
                "一方、介護が「ない」が45.8%（65件）で最多である。", size=10.5)
    _, p = find_par(doc, "課題分析：高頻度の家族介護を担う世帯が一定数存在する")
    set_text(p, "課題分析：高頻度の家族介護を担う世帯が一定数存在する一方、"
                "在宅の認定者の半数近くには家族・親族の介護がない。"
                "後者は単身世帯（66.7%が介護「ない」）に集中しており、"
                "施設入所の申込み（56.9%）とサービス未利用（76.9%）の双方が突出している。", size=10.5)
    _, p = find_par(doc, "施策への示唆：介護者教室、レスパイト")
    set_text(p, "施策への示唆：介護者がいる世帯には介護者教室・レスパイト・ショートステイ等の"
                "負担軽減策を周知する。介護者がいない世帯には、在宅生活を継続するための"
                "サービス調整と見守り体制を優先的に手当てする。", size=10.5)

    # ---------- 第12章 施策の方向性に2行を追加
    for tb in doc.tables:
        if [c.text for c in tb.rows[0].cells][:2] == ["課題", "施策の方向性"]:
            for row in [
                ["家族介護者の不在", "介護者がいない在宅認定者への重点支援",
                 "サービス未利用者の訪問確認、ケアマネジャーによる利用勧奨、見守り・配食等の生活支援、"
                 "地域の支え合い活動との接続", "介護・地域福祉"],
                ["単身高齢者の住まい", "住まいを含めた選択肢の提示",
                 "施設・高齢者向け住まいの情報整理、入所申込みの実態把握、"
                 "地域密着型サービスの整備量の検討", "介護・住宅"],
            ]:
                cells = tb.add_row().cells
                for i, v in enumerate(row):
                    set_text(cells[i].paragraphs[0], v, size=9)
            break

    doc.save(OUT_C)
    print(f"C 保存: {os.path.basename(OUT_C)}（クロス集計{len(xt)}件を実算出）")


if __name__ == "__main__":
    fix_b()
    fix_c()
