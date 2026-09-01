# -*- coding: utf-8 -*-
"""
計画素案 docx のパッケージ修復

word/media 配下の画像が拡張子 ".undefined" で格納されており、[Content_Types].xml に
対応する Default 宣言がないため、OOXML として不正な状態になっている。
（Word は開ける場合があるが、python-docx をはじめ厳密に検証する処理系では失敗する）

実体はいずれも PNG であるため、拡張子を .png に改め、リレーションシップの
Target を書き換える。画像そのものは一切変更しない。
"""
import re
import shutil
import zipfile
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
TARGETS = [
    BASE / "01_第10期_最新版成果品/川崎町_計画書素案_v1.6_本文充実版.docx",
    BASE / "02_計画素案/川崎町_計画素案_概要版.docx",
    BASE / "01_第10期_最新版成果品/川崎町_資料編_構成案・素案.docx",
]

SIG = {
    b"\x89PNG\r\n\x1a\n": "png",
    b"\xff\xd8\xff": "jpg",
    b"GIF8": "gif",
    b"BM": "bmp",
}


def sniff(data):
    for sig, ext in SIG.items():
        if data.startswith(sig):
            return ext
    return None


def repair(path):
    if not path.exists():
        return f"{path.name}: ファイルなし（スキップ）"
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        bad = [n for n in names if n.startswith("word/media/") and n.endswith(".undefined")]
        if not bad:
            return f"{path.name}: 修復不要"
        rename = {}
        for n in bad:
            ext = sniff(z.read(n)[:16])
            if ext is None:
                return f"{path.name}: 形式を判別できない画像がある（{n}）。手作業での確認が必要"
            rename[n] = n[: -len("undefined")] + ext
        blobs = {n: z.read(n) for n in names}

    exts = sorted({v.rsplit(".", 1)[1] for v in rename.values()})
    tmp = path.with_suffix(".repair.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as out:
        for n, data in blobs.items():
            name = rename.get(n, n)
            if n.endswith(".rels") or n == "[Content_Types].xml":
                x = data.decode("utf-8")
                for old, new in rename.items():
                    x = x.replace(old.split("/")[-1], new.split("/")[-1])
                if n == "[Content_Types].xml":
                    # 必要な Default 宣言が欠けていれば補う
                    for ext in exts:
                        if f'Extension="{ext}"' not in x:
                            ct = {"png": "image/png", "jpg": "image/jpeg",
                                  "gif": "image/gif", "bmp": "image/bmp"}[ext]
                            x = x.replace("<Types ", f'<Types ', 1)
                            x = re.sub(r"(<Types[^>]*>)",
                                       rf'\1<Default ContentType="{ct}" Extension="{ext}"/>', x, count=1)
                data = x.encode("utf-8")
            out.writestr(name, data)
    shutil.move(tmp, path)
    return f"{path.name}: 画像{len(rename)}点の拡張子を .{'/.'.join(exts)} に修復"


def main():
    for p in TARGETS:
        print(" ", repair(p))
    print("\n--- 読み込み確認 ---")
    from docx import Document
    for p in TARGETS:
        if not p.exists():
            continue
        try:
            d = Document(str(p))
            print(f"  OK  {p.name}  段落{len(d.paragraphs)} 表{len(d.tables)} "
                  f"図{d.element.xml.count('<w:drawing>')}")
        except Exception as e:
            print(f"  NG  {p.name}  {e}")


if __name__ == "__main__":
    main()
