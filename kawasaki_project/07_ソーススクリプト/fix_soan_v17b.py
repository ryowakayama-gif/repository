# -*- coding: utf-8 -*-
"""
計画書素案 v1.7 仕上げ

  1. 「介護保険運営委員会」の残りを「策定委員会」に統一（第9期の記述は当時の名称を維持）
  2. 表紙のサブタイトルを 4-1 の記載に揃える
  3. 3-1 の重複した出典注記を整理
"""
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

BASE = Path(__file__).resolve().parent.parent
DOC = BASE / "01_第10期_最新版成果品/川崎町_計画書素案_v1.7_第9期体系是正版.docx"

# 第9期当時の名称として残す記述
KEEP = "第9期策定では、ニーズ調査・在宅介護実態調査、パブリックコメント"

log = []


def sub(el, old, new):
    ts = list(el.iter(qn("w:t")))
    if not ts:
        return False
    whole = "".join(t.text or "" for t in ts)
    out = re.sub(re.escape(old), lambda m: new, whole)
    if out == whole:
        return False
    ts[0].text = out
    ts[0].set(qn("xml:space"), "preserve")
    for t in ts[1:]:
        t.text = ""
    return True


def find(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise KeyError(needle)


def main():
    doc = Document(str(DOC))

    # 1. 委員会名称の残り
    n = 0
    for p in doc.paragraphs:
        if KEEP in p.text:
            continue
        if sub(p._p, "介護保険運営委員会", "策定委員会"):
            n += 1
    for t in [Table(c, doc) for c in doc.element.body.iterchildren() if c.tag == qn("w:tbl")]:
        for tc in t._tbl.iter(qn("w:tc")):
            for para in tc.findall(qn("w:p")):
                if KEEP in "".join(x.text or "" for x in para.iter(qn("w:t"))):
                    continue
                if sub(para, "介護保険運営委員会", "策定委員会"):
                    n += 1
    left = [p.text[:36] for p in doc.paragraphs if "介護保険運営委員会" in p.text]
    log.append(f"[No.6] 残りの「介護保険運営委員会」を「策定委員会」に統一（{n}箇所）／"
               f"第9期の記述として維持 {len(left)}箇所")

    # 2. 表紙のサブタイトル
    p = doc.paragraphs[2]
    if "認知症になっても住み慣れた地域で安心して暮らせるまちづくり" in p.text:
        sub(p._p, "〜 認知症になっても住み慣れた地域で安心して暮らせるまちづくり 〜",
            "〜 認知症になっても誰もが自分らしく暮らせる地域共生社会の実現 〜")
        log.append("[整合] 表紙のサブタイトルを 4-1 の記載（第10期で新たに付加するサブタイトル）に統一")

    # 3. 3-1 の重複注記を削除
    p = find(doc, "※ 以下は、第9期計画書（原本42〜44頁「第3章 計画の基本理念」）の記載によります。")
    p._p.getparent().remove(p._p)
    log.append("[整理] 3-1 の重複した出典注記を削除（直下の「出典：…」に一本化）")

    doc.save(str(DOC))
    print("\n".join(log))


if __name__ == "__main__":
    main()
